#!/usr/bin/env python3
"""
proposal_v5_engine.py — Proposal V5: Narrative & Style Engine (PF2-010)

V5 changes the Proposal Factory from a document assembler into a professional
proposal author. No placeholders. No pending sections. No AI-draft markers.
No Oracle feature lists. No capability catalogue verbatim.

Writing standard: experienced Oracle proposal manager.
Every section answers four questions:
  1. What is the customer's challenge?
  2. What does APPSolve propose?
  3. Why is this approach better?
  4. What business outcome will the customer receive?

Outputs:
  CLIENT_PROPOSAL_PLENNEGY_V5.docx
  CLIENT_PROPOSAL_PLENNEGY_V5.md
  (Style guide, reports: written separately)

Usage:
    python3 proposal_v5_engine.py --tender PLENNEGY-HCM-001
"""

import argparse, os, re, sys
from datetime import datetime
from typing import List, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)
from docx_renderer import (
    Seg, parse_md, add_inline, render_table, render_ulist, render_olist,
    setup_styles, _page_break, _toc_field, _cell_bg, _cell_borders,
    NAVY, BLUE, GREY, LGREY, WHITE, ORANGE,
    HEX_NAVY, HEX_BLUE, HEX_WHITE, PROPOSALS_DIR
)
from proposal_shaper import (
    group_v2, classify_v2, _page_setup, _headers_footers, _ps, _ul, _tbl, _heading,
)

# ── Authoring helpers ─────────────────────────────────────────────────────────

def _md_heading(text: str, level: int = 2) -> str:
    return f"\n{'#' * level} {text}\n"

def _md_para(text: str) -> str:
    return f"\n{text}\n"

def _md_bullet(items: List[str]) -> str:
    return "\n" + "\n".join(f"- {item}" for item in items) + "\n"

def _md_table(rows: List[str]) -> str:
    return "\n" + "\n".join(rows) + "\n"

# ── Section 1 — Executive Summary ─────────────────────────────────────────────

def _v5_exec_summary() -> Tuple[List[Seg], str]:
    segs = [
        _ps("Plennegy Group is at an inflection point. The organisation has grown to a scale where "
            "the informal HR processes and disconnected systems that supported smaller operations can "
            "no longer carry the weight of a large, multi-entity agribusiness group. HR is spending "
            "its time reconciling records and rebuilding data rather than managing people. That needs "
            "to change."),
        _ps("This proposal sets out how APPSolve will implement Oracle Fusion HCM Cloud to give "
            "Plennegy a single, unified human capital management capability across all operating "
            "entities — one employee record, one recruitment process, one learning platform, and "
            "one source of truth for workforce decisions."),
        _heading("Plennegy's Business Objectives", 3),
        _ul([
            "Establish a single employee record across all legal entities that HR and management "
            "can rely on — no reconciliation, no duplication, no version conflict",
            "Automate the flow of employee and position data to Plennegy's payroll provider, "
            "eliminating manual rekeying and the payroll errors that come with it",
            "Give hiring managers a structured, trackable recruitment process that reduces "
            "time-to-hire and produces consistent candidate experiences",
            "Build a compliance training capability that tracks completion across employment "
            "categories and produces audit-ready evidence",
            "Provide executive leadership with real-time workforce visibility across all entities "
            "without requiring HR to build reports manually each time",
        ]),
        _heading("APPSolve's Proposal", 3),
        _ps("APPSolve proposes a two-phase implementation of Oracle Fusion HCM Cloud, designed for "
            "a multi-entity South African agribusiness operation. Phase 1 delivers the core: Oracle "
            "Fusion HCM Core, Oracle Recruiting Cloud, and Oracle Integration Cloud connecting "
            "Oracle to Plennegy's payroll provider. Phase 2 extends the platform with Oracle "
            "Learning Cloud and Oracle Talent Management once the Phase 1 foundation is stable "
            "and adopted."),
        _ps("This phasing is intentional. A multi-entity HCM implementation carries meaningful "
            "organisational change. APPSolve manages that change by delivering value in the first "
            "phase — visible, tangible HR improvement before extending to talent and learning — "
            "rather than launching twelve modules simultaneously and overwhelming Plennegy's "
            "team with change."),
        _heading("Why APPSolve", 3),
        _ps("APPSolve has implemented Oracle Fusion HCM Cloud in South African enterprise "
            "organisations across gaming, retail, agribusiness, and financial services. We "
            "understand South African HR compliance — BCEA leave rules, employment equity "
            "reporting, bargaining council agreements — as standard delivery practice, not "
            "research. Our team of more than 50 Senior Consultants brings that experience "
            "directly to Plennegy's project."),
        _ps("We are an Oracle Level 1 Partner, recognised with Oracle's Business Impact Award "
            "(EMEA 2024) for measurable implementation outcomes. We run every engagement through "
            "a named Business Unit Lead — a senior practitioner accountable for delivery quality "
            "from kick-off to go-live, not a project co-ordinator who escalates problems."),
        _heading("Expected Business Outcomes", 3),
        _ul([
            "**HR productivity:** HR business partners spend their time on people strategy, not "
            "data reconciliation across disconnected entity records",
            "**Payroll accuracy:** Payroll errors arising from manual data transfer are structurally "
            "eliminated — the integration carries the data, not a spreadsheet",
            "**Recruitment speed:** Time-to-hire reduces as hiring managers work from a single "
            "system with clear pipeline visibility and structured approval workflows",
            "**Compliance confidence:** Mandatory training completion is tracked automatically, "
            "evidence is available on demand, and employment equity reporting draws from "
            "Oracle's built-in data",
            "**Leadership visibility:** Management reporting on headcount, absence, positions, and "
            "cost is available in real time, without HR rebuilding it from scratch each period",
        ]),
        _ps("APPSolve is ready to mobilise within two weeks of contract signature. We ask Plennegy "
            "to confirm scope and nominate a project sponsor. The next step is yours to take."),
    ]

    md = _md_heading("1. Executive Summary")
    md += _md_para("Plennegy Group is at an inflection point. The organisation has grown to a scale where "
                   "the informal HR processes and disconnected systems that supported smaller operations can "
                   "no longer carry the weight of a large, multi-entity agribusiness group. HR is spending "
                   "its time reconciling records and rebuilding data rather than managing people. That needs to change.")
    md += _md_para("This proposal sets out how APPSolve will implement Oracle Fusion HCM Cloud to give "
                   "Plennegy a single, unified human capital management capability across all operating "
                   "entities — one employee record, one recruitment process, one learning platform, and "
                   "one source of truth for workforce decisions.")
    md += _md_heading("Plennegy's Business Objectives", 3)
    md += _md_bullet([
        "Establish a single employee record across all legal entities",
        "Automate payroll data flow — eliminating manual rekeying and payroll errors",
        "Give hiring managers a structured, trackable recruitment process",
        "Build compliance training capability with audit-ready evidence",
        "Provide real-time workforce visibility for executive leadership",
    ])
    md += _md_heading("APPSolve's Proposal", 3)
    md += _md_para("APPSolve proposes a two-phase Oracle Fusion HCM Cloud implementation. Phase 1: "
                   "HCM Core + Recruiting + OIC payroll integration. Phase 2: Learning + Talent Management.")
    md += _md_heading("Why APPSolve", 3)
    md += _md_para("South African Oracle HCM experience across gaming, retail, agribusiness, and financial "
                   "services. More than 50 Senior Consultants. Oracle Level 1 Partner. Business Impact Award "
                   "EMEA 2024. Named BU Lead accountable for every engagement.")
    md += _md_heading("Expected Business Outcomes", 3)
    md += _md_bullet([
        "HR productivity: time on people strategy, not data reconciliation",
        "Payroll accuracy: structural elimination of manual rekeying errors",
        "Recruitment speed: reduced time-to-hire with pipeline visibility",
        "Compliance confidence: tracked completion, audit-ready evidence",
        "Leadership visibility: real-time workforce reporting",
    ])
    md += _md_para("APPSolve is ready to mobilise within two weeks of contract signature.")
    return segs, md

# ── Section 2 — Understanding Plennegy's Requirements ─────────────────────────

def _v5_understanding() -> Tuple[List[Seg], str]:
    segs = [
        _ps("Plennegy operates across multiple agribusiness entities with a workforce that spans "
            "permanent, seasonal, and contract employment categories. Each entity manages its own "
            "HR records. When employees transfer between entities, information moves by hand — or "
            "does not move at all. The result is a group-level HR function that cannot see its own "
            "workforce without assembling the picture manually."),
        _ps("The consequences are practical and immediate. Payroll requires manual reconciliation "
            "at every period-end because employee and position data are not flowing automatically "
            "to the payroll provider. Absence records are unreliable across entities because each "
            "entity manages them differently. HR cannot provide accurate headcount or employment "
            "equity numbers to leadership without rebuilding the data from source each time."),
        _heading("What Plennegy Needs from This Project", 3),
        _ul([
            "**A single employee record.** One source of truth for employment data, position, "
            "absence, leave, and contract history — accessible across all Plennegy entities, "
            "updated in real time, reliable for payroll and reporting without manual consolidation.",
            "**Automated payroll integration.** Oracle Fusion HCM data flowing directly to "
            "Plennegy's payroll provider at each pay period. Confirmed pay results flowing back "
            "for reconciliation. Exceptions surfacing automatically — not discovered in the "
            "reconciliation report after the fact.",
            "**Structured recruitment.** A consistent, trackable recruitment process across all "
            "entities. Every vacancy, every candidate, every hiring decision — in one system with "
            "one workflow and one audit trail.",
            "**Compliance training at scale.** Mandatory training completion tracked per "
            "employee, per employment category, per entity. Evidence available on demand. No "
            "manual tracking spreadsheets.",
            "**Performance and talent management.** A structured performance cycle across all "
            "entities — goals, reviews, calibration — replacing ad hoc and inconsistent processes "
            "with a uniform approach that HR and management can rely on.",
        ]),
        _heading("How APPSolve Has Interpreted This Requirement", 3),
        _ps("APPSolve has scoped this as a two-phase delivery. Phase 1 establishes the foundation "
            "that Plennegy needs most urgently: HCM Core for the single employee record, Oracle "
            "Recruiting for structured talent acquisition, and Oracle Integration Cloud for the "
            "payroll provider interface. Phase 2 adds Learning and Talent Management once the "
            "Phase 1 platform is stable, adopted, and producing the outcomes described above."),
        _ps("This sequencing is not conservative. It is accurate. A single-phase deployment of "
            "all five modules simultaneously would place excessive organisational change demands "
            "on Plennegy's HR team and management during a period when they are also learning a "
            "new system. Phase 1 creates the visible wins — a working payroll integration, a "
            "reliable employee record, a consistent recruitment process — that build the "
            "confidence and capability for Phase 2."),
        _ps("APPSolve's South African HCM experience means we bring configuration templates for "
            "BCEA leave types, employment equity categories, and South African payroll provider "
            "integration patterns to this engagement from day one. These are not design decisions "
            "Plennegy's project team will need to guide us through. They are solved problems we "
            "bring to Plennegy's table."),
    ]

    md = _md_heading("2. Understanding Plennegy's Requirements")
    md += _md_para("Plennegy operates across multiple agribusiness entities with a workforce spanning "
                   "permanent, seasonal, and contract employment categories. Each entity manages its "
                   "own HR records. When employees transfer between entities, information moves by "
                   "hand — or does not move at all.")
    md += _md_heading("What Plennegy Needs", 3)
    md += _md_bullet([
        "A single employee record across all entities — reliable for payroll and reporting",
        "Automated payroll integration — no manual rekeying, no period-end reconciliation",
        "Structured recruitment — one process, one system, one audit trail across entities",
        "Compliance training tracking — per employee, per category, per entity",
        "Structured performance management — uniform process across all entities",
    ])
    md += _md_heading("APPSolve's Interpretation", 3)
    md += _md_para("Two-phase delivery: Phase 1 (HCM Core + Recruiting + OIC) establishes the "
                   "foundation. Phase 2 (Learning + Talent) extends it once the foundation is stable "
                   "and adopted. South African compliance configuration (BCEA, employment equity, "
                   "payroll integration) is pre-built — not designed during the project.")
    return segs, md

# ── Section 3 — Proposed Oracle HCM Solution ──────────────────────────────────

def _v5_proposed_solution() -> Tuple[List[Seg], str]:
    segs = [
        _heading("Where Plennegy Is Today", 3),
        _ps("Plennegy's HR function manages multiple sets of employee records across its "
            "operating entities. Position management, leave tracking, and payroll inputs are "
            "handled differently by different entities. There is no single view of the group's "
            "workforce. Every consolidated HR report requires someone to build it from source."),
        _heading("Where Plennegy Will Be After Implementation", 3),
        _ps("After implementation, every Plennegy HR business partner works from the same "
            "employee record. Position transfers happen in Oracle — the record follows the "
            "employee. Leave approvals happen in Oracle — managers approve on mobile without "
            "routing paper through HR. Payroll inputs flow automatically to the payroll provider "
            "at period-end. Hiring managers see their recruitment pipeline in real time. "
            "Leadership sees headcount, absence, and position data without waiting for HR "
            "to consolidate it."),
        _heading("How Oracle Fusion HCM Makes This Possible", 3),
        _ps("Oracle Fusion HCM Cloud is not a collection of modules. It is a single data model "
            "that manages everything about an employee — their employment record, their position, "
            "their absence, their learning history, their performance — in one place, in real time, "
            "across all of Plennegy's entities. The modules are different windows onto the same "
            "data, not separate systems that need to exchange information."),
        _heading("Phase 1 — The Core Platform", 3),
        _tbl([
            "| What | Why it matters to Plennegy |",
            "|---|---|",
            "| **Oracle Fusion HCM Core** | Single employee record for all entities. Position management, absence management, and employment contract history in one place. HR self-service for managers eliminates approval paper trails. |",
            "| **Oracle Recruiting Cloud** | End-to-end recruitment from vacancy creation to accepted offer. Every hiring decision is tracked. Employment equity data is captured at application, not reconstructed after hiring. |",
            "| **Oracle Integration Cloud** | Payroll provider interface. Employee and position data flows automatically at period-end. Pay results flow back for reconciliation. Exceptions surface in the integration monitor. |",
        ]),
        _heading("Phase 2 — Platform Extension", 3),
        _tbl([
            "| What | Why it matters to Plennegy |",
            "|---|---|",
            "| **Oracle Learning Cloud** | Compliance training tracked per employee, per employment category, per entity. Mandatory completions evidenced automatically. No spreadsheet tracking. |",
            "| **Oracle Talent Management** | Structured performance cycle across all entities. Goal setting, mid-year check-in, year-end review — uniform and consistent. Succession planning and talent pipeline visibility for executive leadership. |",
        ]),
        _heading("How APPSolve Delivers It", 3),
        _ps("APPSolve configures Oracle to work for Plennegy's operating model — your entity "
            "structure, your position hierarchies, your employment categories, your leave policies, "
            "your approval workflows. The configuration decisions are made with Plennegy in "
            "solution design workshops, not by APPSolve working in isolation. Your HR team "
            "participates in every design decision because they are the ones who will manage "
            "the system after APPSolve leaves."),
        _heading("What This Delivers for Plennegy", 3),
        _ul([
            "**Period-end payroll:** Plennegy's payroll team processes payroll from Oracle data "
            "without rekeying. Errors surface in the integration monitor before they reach the "
            "payroll provider.",
            "**HR reporting:** Leadership gets workforce reports from Oracle directly — "
            "headcount, absence rates, position vacancies — without HR assembling them manually.",
            "**Employment equity:** Equity reporting draws from Oracle's live employee record, "
            "not from manually consolidated spreadsheets.",
            "**Recruitment transparency:** Every open vacancy, every candidate in process, every "
            "pending offer — visible to hiring managers and HR in real time.",
        ]),
        _heading("What Is Not in This Scope", 3),
        _ul([
            "Oracle Fusion Payroll (native) — APPSolve proposes Oracle Integration Cloud to connect "
            "to Plennegy's existing payroll provider rather than replacing it",
            "Oracle HR Help Desk — available as a Phase 3 extension if required",
            "Custom software development beyond Oracle's Low-Code/No-Code extension framework",
            "Oracle Analytics Cloud — separately licensed; available as a Phase 3 addition",
        ]),
    ]

    md = _md_heading("3. Proposed Oracle HCM Solution")
    md += _md_heading("Where Plennegy Is Today", 3)
    md += _md_para("Multiple employee record sets across entities. No single workforce view. "
                   "Every consolidated report built manually from source.")
    md += _md_heading("Where Plennegy Will Be After Implementation", 3)
    md += _md_para("One employee record across all entities. Payroll inputs flow automatically. "
                   "Managers approve leave on mobile. Hiring managers see recruitment pipelines in real time. "
                   "Leadership gets workforce reports from Oracle — not from HR.")
    md += _md_heading("Phase 1 — Core Platform", 3)
    md += _md_table([
        "| Module | Why it matters to Plennegy |",
        "|---|---|",
        "| Oracle Fusion HCM Core | Single employee record; position management; absence; self-service |",
        "| Oracle Recruiting Cloud | End-to-end recruitment; equity data at point of application |",
        "| Oracle Integration Cloud | Payroll provider interface; automated period-end data flow |",
    ])
    md += _md_heading("Phase 2 — Platform Extension", 3)
    md += _md_table([
        "| Module | Why it matters to Plennegy |",
        "|---|---|",
        "| Oracle Learning Cloud | Compliance training tracking; automated completion evidence |",
        "| Oracle Talent Management | Structured performance cycle; succession planning |",
    ])
    md += _md_heading("Not in Scope", 3)
    md += _md_bullet([
        "Oracle Fusion Payroll (native) — OIC integration to existing payroll provider instead",
        "Oracle HR Help Desk — Phase 3 extension available",
        "Custom development beyond Low-Code/No-Code framework",
    ])
    return segs, md

# ── Section 4 — Why APPSolve ───────────────────────────────────────────────────

def _v5_why_appsolve() -> Tuple[List[Seg], str]:
    segs = [
        _ps("Oracle implementations succeed or fail on the implementation partner. Oracle provides "
            "the technology. The configuration decisions, delivery discipline, change management, "
            "knowledge transfer, and long-term support are the partner's responsibility. These are "
            "the factors that determine whether Plennegy realises the full value of Oracle "
            "Fusion HCM Cloud in years one, three, and five."),
        _heading("We Have Solved Plennegy's Problems Before", 3),
        _ps("APPSolve has implemented Oracle Fusion HCM Cloud in South African enterprise "
            "organisations facing the same challenges Plennegy faces today: multi-entity HR "
            "complexity, manual payroll data transfer, disconnected recruitment processes, and "
            "the need to build group-wide HR capability without disrupting operations that "
            "cannot stop."),
        _ps("We understand South African HR as a configuration challenge, not a research project. "
            "BCEA leave types, employment equity categories, bargaining council agreement rules, "
            "and South African payroll provider integration patterns are solved problems for "
            "APPSolve's team. They will not be designed during Plennegy's project — they will be "
            "brought to Plennegy's workshops from day one."),
        _heading("Our Oracle Capability", 3),
        _ul([
            "More than 50 Senior Consultants across Oracle HCM, ERP, Integration Cloud, Database "
            "Administration, and Cloud Infrastructure",
            "Oracle Level 1 Partner — the Oracle Partner Network recognition for consistent "
            "cloud delivery capability",
            "Oracle Business Impact Award (EMEA 2024) — awarded for measurable business outcomes "
            "from Oracle HCM implementations, not for licence volume",
            "Oracle ECEMEA Innovation Award (2024) — awarded for innovation in Oracle Cloud "
            "delivery in our region",
            "Delivery footprint across South Africa, Botswana, Namibia, Mozambique, Zambia, "
            "Tanzania, and internationally (USA, France, Abu Dhabi)",
        ]),
        _heading("Named Accountability — Not Escalation Chains", 3),
        _ps("Every APPSolve engagement is run by a named Business Unit Lead. This is a senior "
            "Oracle practitioner — not a project manager or account executive — who takes "
            "personal responsibility for delivery quality from kick-off to go-live. If the "
            "project encounters a problem, the BU Lead resolves it. There is no escalation path "
            "that routes to someone who does not know Plennegy's system."),
        _ps("This matters for a multi-entity implementation where configuration decisions made "
            "in Mobilise have downstream consequences in Validate. The person who made those "
            "decisions needs to be available, accountable, and present throughout the project — "
            "not cycling through to the next engagement."),
        _heading("Knowledge Transfer by Design", 3),
        _ps("APPSolve designs every engagement so that the client can operate and evolve the "
            "system without the implementation partner after go-live. Plennegy's HR and IT teams "
            "will participate in every design workshop, shadow every configuration decision, and "
            "receive structured training before hypercare ends. Three years from now, Plennegy "
            "should not need APPSolve to change a leave policy or update an approval workflow."),
        _heading("Long-Term Support Through AMS", 3),
        _ps("The majority of APPSolve's clients continue with our Application Managed Services "
            "practice after go-live. Oracle updates its cloud platform every quarter — new "
            "functionality, changed interfaces, configuration impacts. APPSolve's AMS practice "
            "manages those updates, tests impacts to Plennegy's configuration, and handles the "
            "quarterly release cycle. Plennegy's team does not need to track Oracle's release "
            "notes. We track them on your behalf."),
        _ps("Client references are available on request, subject to individual client approval. "
            "Contact your APPSolve Account Manager."),
    ]

    md = _md_heading("4. Why APPSolve")
    md += _md_para("Oracle implementations succeed or fail on the partner. APPSolve has solved "
                   "Plennegy's specific problems — multi-entity HR complexity, manual payroll data "
                   "transfer, disconnected recruitment — in South African enterprise organisations "
                   "across gaming, retail, agribusiness, and financial services.")
    md += _md_heading("Oracle Capability", 3)
    md += _md_bullet([
        "50+ Senior Consultants across Oracle HCM, ERP, OIC, DBA, and Cloud Infrastructure",
        "Oracle Level 1 Partner",
        "Oracle Business Impact Award EMEA 2024 (measurable HCM outcomes)",
        "Oracle ECEMEA Innovation Award 2024",
        "Delivery across South Africa and Sub-Saharan Africa",
    ])
    md += _md_heading("Named Accountability", 3)
    md += _md_para("Every engagement has a named Business Unit Lead — a senior practitioner who is "
                   "personally accountable for delivery quality. No escalation chain. No handover "
                   "to someone who doesn't know your project.")
    md += _md_heading("Knowledge Transfer and AMS", 3)
    md += _md_para("APPSolve's engagements are designed so that the client can operate the system "
                   "independently after go-live. Oracle quarterly updates managed by APPSolve's AMS "
                   "practice — Plennegy's team does not need to track Oracle's release notes.")
    return segs, md

# ── Section 5 — How We Deliver ────────────────────────────────────────────────

def _v5_how_we_deliver() -> Tuple[List[Seg], str]:
    segs = [
        _ps("APPSolve delivers Oracle Fusion HCM Cloud using Oracle's Modern Best Practices and "
            "Oracle Success Navigator — the adoption framework Oracle built for SaaS cloud "
            "implementations. This is not a traditional ERP methodology applied to a cloud "
            "engagement. That distinction matters."),
        _ps("Oracle Fusion HCM is a continuously evolving cloud platform. It is not a system "
            "APPSolve installs and maintains on Plennegy's infrastructure. Our job is to configure "
            "Oracle to reflect how Plennegy actually operates — your entity structure, your "
            "position hierarchies, your leave policies, your approval workflows — not to recreate "
            "Plennegy's current processes in new software. Where current processes can be improved "
            "by adopting Oracle's standard approach, we will make that recommendation during "
            "solution design."),
        _heading("Our Delivery Principles", 3),
        _ul([
            "**Configure, do not customise.** Oracle's standard configuration meets the "
            "requirements of most South African organisations. Custom code adds cost, upgrade "
            "risk, and long-term maintenance burden. APPSolve works within Oracle's Low-Code/"
            "No-Code extension framework when configuration alone is insufficient.",
            "**Decisions are the deliverable.** Oracle HCM implementation requires Plennegy to "
            "make design decisions — about organisation structure, security, approval hierarchies, "
            "position naming, leave policy — during the Scope and Design phase. APPSolve runs "
            "structured workshops to drive those decisions. Deferred decisions become project risk.",
            "**Data quality is a shared responsibility.** The quality of Plennegy's Oracle "
            "system on day one is determined by the quality of the data migrated into it. "
            "APPSolve provides templates, tools, and migration guidance. Plennegy's team owns "
            "the data quality process — that is the correct allocation of responsibility.",
            "**Go-live is the beginning, not the end.** Every APPSolve engagement plans for "
            "long-term success from the first workshop. Hypercare, AMS handover, and knowledge "
            "transfer are designed in, not bolted on.",
        ]),
        _heading("The Six Phases of Delivery", 3),
        _tbl([
            "| Phase | Purpose | Your Involvement | Outcome |",
            "|---|---|---|---|",
            "| **1 — Mobilise** (4 weeks) | Establish the foundation | Confirm project team, sign charter, attend kick-off | Project chartered; Oracle environments provisioned; governance active |",
            "| **2 — Scope & Design** (8–10 weeks) | Build the configuration blueprint | Participate in solution design workshops; approve design documents | Solution design approved; CRP-1 demonstrated; integration design confirmed |",
            "| **3 — Prototype** (6–8 weeks) | Build and demonstrate the configured system | Review the prototype in CRP-2; provide structured feedback | CRP-2 signed off; integrations built and tested in development |",
            "| **4 — Validate** (6–8 weeks) | Confirm the system is ready for your users | Lead user acceptance testing; approve data migration; confirm go-live | UAT signed off; migration rehearsed; go-live decision made |",
            "| **5 — Deploy** (2–4 weeks) | Go live | Support your users during hypercare | Oracle is live; APPSolve on-site during hypercare; issues resolved within agreed SLA |",
            "| **6 — Evolve** (ongoing) | Optimise and sustain | Raise service requests; participate in quarterly update reviews | Oracle quarterly updates managed by APPSolve AMS; configuration continuously optimised |",
        ]),
        _heading("Change Management", 3),
        _ps("Oracle HCM implementations do not fail because of configuration errors. They fail "
            "because people revert to the old way of working after go-live. APPSolve builds "
            "change management into every phase. Stakeholder engagement begins in Mobilise. "
            "End-user training runs throughout Validate. Hypercare in Deploy puts APPSolve "
            "consultants alongside Plennegy's users during the first weeks on the new system. "
            "Change management is not a workshop before go-live. It is continuous activity."),
        _heading("Reporting and Transparency", 3),
        _ul([
            "Weekly project status report covering RAG status per workstream, upcoming milestones, "
            "open risks, and decisions pending from Plennegy",
            "RAID Register (Risks, Assumptions, Issues, Dependencies) maintained and reviewed weekly",
            "Formal Change Request process for every scope change — cost and timeline impact "
            "assessed before approval; no surprises",
            "Go-live readiness review: formal sign-off criteria reviewed with Plennegy's sponsor "
            "before any go-live decision is confirmed",
        ]),
    ]

    md = _md_heading("5. How We Deliver")
    md += _md_para("APPSolve uses Oracle's Modern Best Practices and Oracle Success Navigator — "
                   "Oracle's SaaS cloud adoption framework. Not a traditional ERP methodology. "
                   "Configure to Plennegy's operating model; do not recreate existing processes.")
    md += _md_heading("Delivery Principles", 3)
    md += _md_bullet([
        "Configure, do not customise — stay within Oracle's standard framework",
        "Design decisions are the deliverable in Phase 2 — deferred decisions become project risk",
        "Data quality is a shared responsibility — APPSolve provides tools; Plennegy owns the data",
        "Go-live is the beginning — hypercare, AMS, and knowledge transfer are designed in from day one",
    ])
    md += _md_heading("Six Phases", 3)
    md += _md_table([
        "| Phase | Duration | Outcome |",
        "|---|---|---|",
        "| Mobilise | 4 weeks | Project chartered; Oracle environments provisioned |",
        "| Scope & Design | 8–10 weeks | Solution design approved; CRP-1 complete |",
        "| Prototype | 6–8 weeks | CRP-2 signed off; integrations built |",
        "| Validate | 6–8 weeks | UAT signed off; migration rehearsed; go-live confirmed |",
        "| Deploy | 2–4 weeks | Go-live; hypercare active |",
        "| Evolve | Ongoing | AMS active; Oracle quarterly updates managed |",
    ])
    return segs, md

# ── Section 6 — Project Governance ────────────────────────────────────────────

def _v5_governance() -> Tuple[List[Seg], str]:
    segs = [
        _ps("Good governance does not slow a project down. It prevents the decisions, delays, "
            "and misalignments that slow a project down from happening in the first place. "
            "APPSolve's governance model gives Plennegy's leadership clear visibility of project "
            "progress and the mechanisms to intervene early when something needs attention."),
        _tbl([
            "| Level | Forum | Frequency | Purpose |",
            "|---|---|---|---|",
            "| **Strategic** | Steering Committee | Monthly | Scope authority, commercial decisions, strategic risk — attended by Plennegy executive sponsor and APPSolve BU Lead |",
            "| **Operational** | Project Governance | Weekly | Progress, risks, open decisions, upcoming milestones — attended by Plennegy PM and APPSolve PM |",
            "| **Workstream** | Module Reviews | Bi-weekly | Design decisions, CRP planning, testing sign-off — attended by Plennegy functional leads and APPSolve consultants |",
        ]),
        _heading("What Plennegy Can Expect", 3),
        _ul([
            "A weekly project status report: RAG by workstream, milestone progress, open risks, "
            "and decisions pending from Plennegy — sent to all stakeholders every Friday",
            "A RAID Register maintained by APPSolve and reviewed at every operational governance meeting",
            "A formal Change Request for every scope change, with timeline and cost impact assessed "
            "before Plennegy approves or declines — no unauthorised scope additions",
            "A go-live readiness review at the end of Validate, with formal sign-off criteria, "
            "before any go-live decision is confirmed",
        ]),
        _ps("Governance templates — charter, RAID Register, Change Request form, meeting agenda — "
            "are provided by APPSolve during Mobilise and adapted to Plennegy's requirements."),
    ]

    md = _md_heading("6. Project Governance")
    md += _md_table([
        "| Level | Forum | Frequency | Purpose |",
        "|---|---|---|---|",
        "| Strategic | Steering Committee | Monthly | Scope authority; commercial decisions; strategic risk |",
        "| Operational | Project Governance | Weekly | Progress; risks; pending decisions |",
        "| Workstream | Module Reviews | Bi-weekly | Design decisions; testing sign-off |",
    ])
    md += _md_bullet([
        "Weekly status report: RAG by workstream, milestones, risks, pending decisions",
        "RAID Register reviewed weekly",
        "Formal Change Request for every scope change",
        "Go-live readiness sign-off before any go-live is confirmed",
    ])
    return segs, md

# ── Section 7 — Implementation Roadmap ────────────────────────────────────────

def _v5_roadmap() -> Tuple[List[Seg], str]:
    segs = [
        _ps("The roadmap below shows indicative phase timelines based on a standard Oracle HCM "
            "Phase 1 implementation. Actual durations depend on confirmed scope — entity count, "
            "headcount, number of integrations — and Plennegy's resource availability for design "
            "workshops and testing. APPSolve will provide a milestone-level project plan with named "
            "resource assignments during Mobilise."),
        _tbl([
            "| Phase | Name | Indicative Duration | What Plennegy Commits |",
            "|---|---|---|---|",
            "| 1 | Mobilise | 4 weeks | Project sponsor nominated; project team available at kick-off |",
            "| 2 | Scope & Design | 8–10 weeks | Functional leads available for design workshops; decisions made within 48 hours of workshop |",
            "| 3 | Prototype | 6–8 weeks | CRP-2 attended by functional leads; structured feedback provided |",
            "| 4 | Validate | 6–8 weeks | Named UAT team with confirmed availability; data migration approved before go-live |",
            "| 5 | Deploy | 2–4 weeks | Management visible and present during first week on Oracle |",
            "| 6 | Evolve | 12+ months | AMS service requests raised through agreed channel; quarterly update reviews attended |",
        ]),
        _ul([
            "**Kick-off target:** within two weeks of contract signature and Oracle environment provisioning",
            "**Go-live target:** confirmed at end of Validate once UAT is signed off and migration is rehearsed",
            "Oracle environment provisioning is Oracle Corporation's responsibility — APPSolve's "
            "timeline commences from confirmed environment access",
            "Phase 2 (Learning and Talent) commences three to six months after Phase 1 go-live, "
            "once the Phase 1 platform is stable and the team is operating confidently",
        ]),
    ]

    md = _md_heading("7. Implementation Roadmap")
    md += _md_table([
        "| Phase | Duration | Plennegy Commitment |",
        "|---|---|---|",
        "| Mobilise | 4 weeks | Project sponsor nominated; team at kick-off |",
        "| Scope & Design | 8–10 weeks | Functional leads in workshops; decisions within 48 hours |",
        "| Prototype | 6–8 weeks | CRP-2 attended; structured feedback provided |",
        "| Validate | 6–8 weeks | Named UAT team; data migration approved |",
        "| Deploy | 2–4 weeks | Management visible during first week on Oracle |",
        "| Evolve | 12+ months | AMS service requests; quarterly update reviews |",
    ])
    md += _md_bullet([
        "Kick-off within two weeks of contract signature",
        "Go-live confirmed at end of Validate — not before",
        "Phase 2 commences 3–6 months after Phase 1 go-live",
    ])
    return segs, md

# ── Section 8 — Key Commercial Assumptions ────────────────────────────────────

def _v5_assumptions() -> Tuple[List[Seg], str]:
    segs = [
        _ps("This proposal rests on a set of commercial assumptions that define the boundaries "
            "of the scope, the timeline, and the price. APPSolve's formal Statement of Work will "
            "include a complete Assumption Schedule. The five assumptions below are the most "
            "material — they carry the most consequence if they prove to be different from "
            "what APPSolve has understood."),
        _tbl([
            "| Assumption | What it means for the engagement |",
            "|---|---|",
            "| **Platform:** Oracle provisions Oracle Fusion HCM Cloud SaaS environments. APPSolve's project timeline begins from confirmed environment access. | Delays in Oracle environment provisioning affect project timelines. APPSolve manages the provisioning process but cannot accelerate Oracle Corporation's delivery. |",
            "| **Scope:** Implementation scope, timeline, and price are based on entity count, headcount, modules, and go-live date confirmed at the start of Mobilise. | Any additions after Mobilise require a formal Change Request. APPSolve will assess the impact on timeline and cost and present it for Plennegy's approval before proceeding. |",
            "| **Data:** Employee data extraction, cleansing, and validation are Plennegy's responsibility. APPSolve provides templates, tools, and migration guidance. | The quality of Plennegy's Oracle system on go-live day is determined by the quality of the migration data. APPSolve cannot guarantee outcomes driven by data quality it does not control. |",
            "| **Integration:** Each integration point is separately scoped and priced. Plennegy's payroll provider participates in design and testing. | Payroll providers who do not participate in integration design and testing introduce project risk that APPSolve cannot manage. |",
            "| **Resource:** Plennegy provides a nominated project team with confirmed availability for design workshops, testing, and go-live decisions. | Decision delays caused by resource unavailability extend project timelines. APPSolve will log all pending decisions in the RAID Register and escalate to Steering Committee when decisions are overdue. |",
        ]),
        _ps("The complete Assumption Schedule will be attached to APPSolve's Statement of Work. "
            "No item excluded from the Assumption Schedule is in scope unless explicitly added "
            "by approved Change Request."),
    ]

    md = _md_heading("8. Key Commercial Assumptions")
    md += _md_table([
        "| Assumption | Consequence |",
        "|---|---|",
        "| Oracle provisions SaaS environments | Timeline starts from environment access date |",
        "| Scope confirmed at Mobilise | Post-Mobilise additions require Change Request |",
        "| Plennegy owns data migration quality | APPSolve provides templates; Plennegy owns execution |",
        "| Each integration separately scoped | Payroll provider participates in design and testing |",
        "| Plennegy resource available | Decision delays extend timelines |",
    ])
    md += _md_para("Full Assumption Schedule in the Statement of Work. Items not in the schedule are out of scope.")
    return segs, md

# ── Section 9 — Key Delivery Risks ────────────────────────────────────────────

def _v5_risks() -> Tuple[List[Seg], str]:
    segs = [
        _ps("Every Oracle HCM implementation carries delivery risks. The five below are the "
            "most material for an engagement of this type. APPSolve has managed all five before — "
            "in South African multi-entity organisations with the same profile as Plennegy. The "
            "mitigations described are practices APPSolve applies on every engagement, not "
            "plans written in response to risk identification."),
        _tbl([
            "| Risk | Probability | Impact | How APPSolve manages it |",
            "|---|---|---|---|",
            "| **Scope ambiguity** — scope not confirmed before Mobilise begins | Medium | High | APPSolve will not start Mobilise without written scope confirmation from Plennegy's sponsor: entity count, headcount, modules, payroll provider, go-live target, and preferred commercial model. |",
            "| **Data quality** — migration data provided late, incomplete, or with errors | High | High | APPSolve issues data migration templates in Scope and Design. Data validation runs in Prototype. Final migration sign-off is part of the go-live readiness criteria in Validate. Go-live does not proceed without approved migration results. |",
            "| **Resource availability** — Plennegy's project team unavailable for workshops or decisions | Medium | High | Named Plennegy project team with confirmed availability is a condition of Mobilise. Availability commitments are documented in the project charter. Pending decisions are tracked in the RAID Register and escalated to Steering Committee when overdue. |",
            "| **Integration complexity** — undocumented APIs or unresponsive third-party vendors | Medium | Medium | Each integration is separately scoped before build begins. The payroll provider's participation in design and testing is confirmed before integration development starts. APPSolve's integration timeline includes contingency for vendor responsiveness. |",
            "| **Change adoption** — end users revert to previous systems after go-live | Medium | High | Change management runs across all phases from Mobilise. End-user training is delivered in Validate. APPSolve consultants are on-site during hypercare in Deploy to support users in real time. Oracle self-service adoption rates are tracked as a hypercare KPI. |",
        ]),
        _ps("A full Risk Register, with individual risk owners, mitigation actions, review "
            "cadence, and residual risk ratings, will be established during Mobilise and maintained "
            "throughout the engagement."),
    ]

    md = _md_heading("9. Key Delivery Risks and Mitigations")
    md += _md_table([
        "| Risk | Probability | Impact | Mitigation |",
        "|---|---|---|---|",
        "| Scope ambiguity | Medium | High | Written scope confirmation before Mobilise begins |",
        "| Data quality | High | High | Templates in Phase 2; validation in Phase 3; sign-off in Phase 4 |",
        "| Resource availability | Medium | High | Named team with confirmed availability in project charter |",
        "| Integration complexity | Medium | Medium | Each integration separately scoped; vendor participation confirmed |",
        "| Change adoption | Medium | High | Change management in all phases; on-site hypercare in Deploy |",
    ])
    return segs, md

# ── Section 10 — Commercial Position ──────────────────────────────────────────

def _v5_commercial() -> Tuple[List[Seg], str]:
    segs = [
        _ps("APPSolve's commercial proposal is a separate document, issued by our Commercial "
            "Director once scope is confirmed. The price we agree with Plennegy will be built "
            "from the confirmed entity count, headcount, module list, integration points, and "
            "go-live target — not from a template. Pricing conversations that happen before "
            "scope is confirmed produce numbers that change when scope is confirmed. "
            "APPSolve avoids that cycle."),
        _heading("Commercial Models", 3),
        _tbl([
            "| Model | How it works | When it suits Plennegy |",
            "|---|---|---|",
            "| **Fixed Price** | Agreed scope at an agreed cost. Changes are priced through Change Request. | When scope is well-defined and Plennegy wants certainty on implementation cost. |",
            "| **Time and Materials** | Agreed daily rates. No fixed ceiling. | When scope is expected to evolve as the project progresses. |",
            "| **Monthly Recurring Invoice** | Fixed monthly cost. Capacity carries over. | When Plennegy wants predictable monthly spend and a sustained partnership model. APPSolve's preferred model for Phase 2 and AMS. |",
        ]),
        _heading("Getting to a Commercial Proposal", 3),
        _ul([
            "**Confirm scope** with APPSolve's BU Lead: modules, entities, headcount, "
            "payroll provider, go-live target",
            "**Choose a commercial model** — Fixed Price, Time and Materials, or Monthly Recurring Invoice",
            "**Receive the commercial proposal** — typically within five to ten business days of "
            "scope confirmation",
            "The commercial proposal, Statement of Work, and compliance certificates are "
            "submitted as a complete package — APPSolve does not issue them separately",
        ]),
        _ps("APPSolve is committed to a commercial arrangement that works for Plennegy's budget "
            "and planning cycle, not just for APPSolve's revenue cycle. We are happy to discuss "
            "phased payment structures that align cash outflow to project milestones."),
    ]

    md = _md_heading("10. Commercial Position")
    md += _md_para("APPSolve's commercial proposal is issued separately, after scope is confirmed. "
                   "Price is built from confirmed scope — not from a template.")
    md += _md_table([
        "| Model | When it suits Plennegy |",
        "|---|---|",
        "| Fixed Price | Well-defined scope; cost certainty needed |",
        "| Time and Materials | Scope expected to evolve |",
        "| Monthly Recurring Invoice | Predictable monthly spend; APPSolve's preferred model for AMS |",
    ])
    md += _md_para("To receive APPSolve's commercial proposal: confirm scope (modules, entities, "
                   "headcount, payroll provider, go-live date) and preferred commercial model. "
                   "Typical turnaround: 5–10 business days.")
    return segs, md

# ── Section 11 — Your Next Steps ──────────────────────────────────────────────

def _v5_next_steps() -> Tuple[List[Seg], str]:
    segs = [
        _ps("Moving from proposal to implementation requires a small number of decisions and "
            "confirmations from Plennegy. APPSolve is ready to move as soon as these are in "
            "place. The actions below are sequenced by the order in which they unlock progress."),
        _tbl([
            "| Priority | Your Action | What it unlocks |",
            "|---|---|---|",
            "| **1** | Nominate a Plennegy Project Sponsor with authority over scope and commercial decisions | APPSolve's BU Lead confirms engagement and begins Mobilise planning |",
            "| **1** | Confirm Oracle HCM scope: modules, legal entities, headcount, payroll provider, target go-live date | APPSolve issues commercial proposal within 5–10 business days |",
            "| **2** | Choose a commercial model: Fixed Price, Time and Materials, or Monthly Recurring Invoice | APPSolve structures commercial proposal accordingly |",
            "| **2** | Schedule scope confirmation call with APPSolve BU Lead and Account Manager | Enables commercial proposal and Mobilise planning to run in parallel |",
            "| **3** | Confirm any client references Plennegy wishes to contact | APPSolve arranges reference conversations through the relevant client account team |",
            "| **3** | Confirm B-BBEE compliance documentation requirements for this submission | APPSolve provides current compliance certificates as part of the commercial package |",
        ]),
        _ps("APPSolve's Account Manager is the first point of contact for all of the above. "
            "Once scope is confirmed and the commercial proposal is issued, APPSolve's BU Lead "
            "joins the conversation directly to align on kick-off timing, project team, and "
            "Mobilise planning."),
        _ps("We look forward to working with Plennegy."),
    ]

    md = _md_heading("11. Your Next Steps")
    md += _md_table([
        "| Priority | Action | What it unlocks |",
        "|---|---|---|",
        "| 1 | Nominate Plennegy Project Sponsor | Mobilise planning begins |",
        "| 1 | Confirm Oracle HCM scope | Commercial proposal in 5–10 days |",
        "| 2 | Choose commercial model | Proposal structured accordingly |",
        "| 2 | Schedule scope confirmation call | Commercial + Mobilise planning in parallel |",
        "| 3 | Confirm reference requirements | APPSolve arranges reference conversations |",
    ])
    md += _md_para("Contact: APPSolve Account Manager. BU Lead joins once scope is confirmed.")
    return segs, md

# ── Section builder ───────────────────────────────────────────────────────────

_V5_BUILDERS = [
    _v5_exec_summary, _v5_understanding, _v5_proposed_solution,
    _v5_why_appsolve, _v5_how_we_deliver, _v5_governance,
    _v5_roadmap, _v5_assumptions, _v5_risks, _v5_commercial, _v5_next_steps,
]

def _build_v5() -> Tuple[List[Tuple], List[Tuple]]:
    sections_docx = []; sections_md = []
    for fn in _V5_BUILDERS:
        segs, md = fn()
        heading = segs[0].lines[0] if segs[0].kind == "heading" else "Section"
        heading = fn.__name__.replace("_v5_", "").replace("_", " ").title()
        # Reconstruct proper heading from the function's first paragraph context
        sections_docx.append((None, segs)); sections_md.append((None, md))
    return sections_docx, sections_md

_V5_HEADINGS = [
    "1. Executive Summary",
    "2. Understanding Plennegy's Requirements",
    "3. Proposed Oracle HCM Solution",
    "4. Why APPSolve",
    "5. How We Deliver",
    "6. Project Governance",
    "7. Implementation Roadmap",
    "8. Key Commercial Assumptions",
    "9. Key Delivery Risks and Mitigations",
    "10. Commercial Position",
    "11. Your Next Steps",
]

# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _render_seg(doc, seg: Seg):
    from docx_renderer import render_callout, classify_comment
    if seg.kind == "heading":
        h = doc.add_heading(level=min(seg.level, 5))
        h.paragraph_format.keep_with_next = True
        add_inline(h, seg.lines[0] if seg.lines else "",
                   sz={1:22,2:16,3:13,4:11,5:10}.get(seg.level, 11))
    elif seg.kind == "table":   render_table(doc, seg.lines)
    elif seg.kind == "ulist":   render_ulist(doc, seg.lines)
    elif seg.kind == "olist":   render_olist(doc, seg.lines)
    elif seg.kind == "paragraph":
        text = " ".join(seg.lines).strip()
        if text:
            p = doc.add_paragraph(style="Normal"); add_inline(p, text)
    elif seg.kind == "comment":
        render_callout(doc, [seg.meta or ""])

def _cover(doc, subtitle: str, tender_id: str):
    doc.add_paragraph().paragraph_format.space_before = Pt(60)
    for txt, sz, bold, colour in [
        ("APPSolve (Pty) Ltd", 32, True, NAVY),
        ("Enterprise Technology Services  |  Oracle  |  Acumatica  |  BeBanking", 11, False, BLUE),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "Enterprise" in txt: p.paragraph_format.space_after = Pt(48)
        r = p.add_run(txt); r.font.name="Calibri"; r.font.size=Pt(sz)
        r.font.bold=bold; r.font.color.rgb=colour
    for txt, sz in [("ORACLE FUSION HCM CLOUD", 18), (subtitle, 14)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(txt); r.font.name="Calibri"; r.font.size=Pt(sz)
        r.font.bold=True; r.font.color.rgb=GREY
    pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.space_after = Pt(40)
    rc = pc.add_run("Plennegy Group"); rc.font.name="Calibri"; rc.font.size=Pt(16)
    rc.font.bold=True; rc.font.color.rgb=NAVY
    tbl = doc.add_table(rows=0, cols=2); tbl.style="Table Grid"
    for k, v in [
        ("Prepared For","Plennegy Group"),("Platform","Oracle Fusion HCM Cloud"),
        ("Prepared By","APPSolve (Pty) Ltd"),("Version","V5 — Narrative & Style Engine"),
        ("Date",datetime.now().strftime("%Y-%m-%d")),("Tender ID",tender_id),
    ]:
        row=tbl.add_row(); kc=row.cells[0]; vc=row.cells[1]
        _cell_bg(kc, HEX_NAVY); _cell_borders(kc, HEX_NAVY, 4)
        _cell_bg(vc, "F2F5FB"); _cell_borders(vc, "D0D0D0", 4)
        for cell, text, bold in [(kc,k,True),(vc,v,False)]:
            cp=cell.paragraphs[0]; cp.paragraph_format.space_before=Pt(4)
            cp.paragraph_format.space_after=Pt(4)
            cr=cp.add_run(text); cr.font.name="Calibri"; cr.font.size=Pt(11)
            cr.font.bold=bold; cr.font.color.rgb = WHITE if cell is kc else GREY
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before=Pt(48)
    fr=fp.add_run("COMMERCIAL IN CONFIDENCE  |  Proposal V5 — Narrative & Style Engine")
    fr.font.name="Calibri"; fr.font.size=Pt(10); fr.font.bold=True
    from proposal_shaper import LGREY2; fr.font.color.rgb=LGREY2
    _page_break(doc)

# ── DOCX renderer ─────────────────────────────────────────────────────────────

def render_docx_v5(sections: List[Tuple[str, List[Seg]]], out_path: str, tender_id: str):
    doc = Document(); setup_styles(doc); _page_setup(doc)
    _cover(doc, "IMPLEMENTATION PROPOSAL", tender_id)
    h = doc.add_heading("Table of Contents", level=1)
    h.paragraph_format.space_before = Pt(0)
    _toc_field(doc); _page_break(doc)
    for heading, segs in sections:
        _page_break(doc)
        doc.add_heading(heading, level=2).paragraph_format.keep_with_next = True
        for seg in segs: _render_seg(doc, seg)
    _headers_footers(doc, "APPSolve (Pty) Ltd",
                     "Oracle HCM Cloud Proposal — Plennegy Group  |  COMMERCIAL IN CONFIDENCE",
                     tender_id)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    sz = os.path.getsize(out_path)
    print(f"[V5 DOCX]  {out_path.split('/')[-1]}  ({sz//1024} KB)  Sections={len(sections)}")

# ── Markdown renderer ─────────────────────────────────────────────────────────

def render_md_v5(sections: List[Tuple[str, str]], out_path: str, tender_id: str):
    lines = [
        "---",
        "title: Oracle Fusion HCM Cloud Implementation Proposal",
        "client: Plennegy Group",
        "prepared_by: APPSolve (Pty) Ltd",
        f"date: {datetime.now().strftime('%Y-%m-%d')}",
        "version: V5",
        "engine: Proposal V5 Narrative & Style Engine (PF2-010)",
        "classification: COMMERCIAL IN CONFIDENCE",
        "---",
        "",
        "# Oracle Fusion HCM Cloud — Implementation Proposal",
        "## Plennegy Group",
        "",
        f"*Prepared by APPSolve (Pty) Ltd | Proposal V5 | {datetime.now().strftime('%Y-%m-%d')}*",
        "",
        "*This document is Commercial in Confidence. It is prepared exclusively for Plennegy Group.*",
        "",
        "---",
        "",
    ]
    for heading, md_text in sections:
        lines.append(f"## {heading}")
        lines.append(md_text.strip())
        lines.append("")
        lines.append("---")
        lines.append("")
    lines += [
        "---",
        "",
        "*Oracle Fusion HCM Cloud Proposal | Plennegy Group | APPSolve (Pty) Ltd*",
        f"*Generated by Proposal V5 Narrative & Style Engine | {datetime.now().strftime('%Y-%m-%d')}*",
    ]
    content = "\n".join(lines)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as fh: fh.write(content)
    print(f"[V5 MD]    {out_path.split('/')[-1]}  ({len(content)//1024} KB)  Sections={len(sections)}")

# ── Orchestrator ──────────────────────────────────────────────────────────────

def run_v5(md_path: str, out_dir: str, tender_id: str) -> dict:
    print(f"[V5] Tender: {tender_id}")
    with open(md_path, "r", encoding="utf-8") as fh: md_text = fh.read()
    segs_all = parse_md(md_text)
    v2_secs = group_v2(segs_all)
    classify_v2(v2_secs)
    print(f"[V5] V2 source: {len(v2_secs)} secs — read as source context only")

    built_segs = []; built_md = []
    for fn in _V5_BUILDERS:
        segs, md_section = fn(); built_segs.append(segs); built_md.append(md_section)

    sections_docx = list(zip(_V5_HEADINGS, built_segs))
    sections_md   = list(zip(_V5_HEADINGS, built_md))

    oum_check = sum(
        len(re.findall(r'Oracle\s+Unified\s+Method|\bOUM\b', " ".join(
            " ".join(s.lines) for s in segs if s.lines
        ), re.IGNORECASE))
        for segs in built_segs
    )
    print(f"[V5] GOV-PAE-001 check: OUM occurrences in authored content = {oum_check}")

    prefix = "PLENNEGY" if "PLENNEGY" in tender_id.upper() else tender_id
    os.makedirs(out_dir, exist_ok=True)
    docx_path = os.path.join(out_dir, f"CLIENT_PROPOSAL_{prefix}_V5.docx")
    md_out    = os.path.join(out_dir, f"CLIENT_PROPOSAL_{prefix}_V5.md")

    render_docx_v5(sections_docx, docx_path, tender_id)
    render_md_v5(sections_md, md_out, tender_id)

    print(f"\n[V5] Complete — outputs in {out_dir}")
    assert oum_check == 0, f"GOV-PAE-001 VIOLATION: {oum_check} OUM occurrences found"
    return {"docx": docx_path, "md": md_out}

# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Proposal V5 Narrative & Style Engine")
    parser.add_argument("--tender", metavar="TENDER_ID")
    parser.add_argument("--md",  metavar="MD_PATH")
    parser.add_argument("--out", metavar="OUT_DIR")
    args = parser.parse_args()
    if args.tender:
        tid = args.tender
        td_dir = os.path.join(PROPOSALS_DIR, tid)
        md_path = os.path.join(td_dir, f"PROPOSAL_RENDERED_{tid}.md")
        out_dir = td_dir
    elif args.md and args.out:
        tid = os.path.basename(args.out); md_path = args.md; out_dir = args.out
    else:
        parser.print_help(); sys.exit(1)
    if not os.path.exists(md_path):
        print(f"[V5] ERROR: {md_path} not found", file=sys.stderr); sys.exit(1)
    run_v5(md_path, out_dir, tid)

if __name__ == "__main__":
    main()
