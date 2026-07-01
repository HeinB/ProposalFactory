#!/usr/bin/env python3
"""
proposal_narrative_engine.py — Proposal Narrative Authoring Engine v1.0 (PNE / PF2-011)

PNE transforms governed source content into a concise, persuasive customer proposal.
Knowledge assets are source material — they are not proposal sections.

For every selected asset, PNE creates a customer-facing proposal view:
  • Business outcome  • APPSolve value  • Delivery confidence
  • Customer benefit  • Relevant evidence

Outputs:
  CLIENT_PROPOSAL_[Tender]_VNEXT.docx        — 25–35 pages, authored narrative
  INTERNAL_REVIEW_PACK_[Tender]_VNEXT.docx   — All sections + action register
  TRACEABILITY_REPORT_[Tender]_VNEXT.docx    — Source mapping + readiness

Usage:
    python3 proposal_narrative_engine.py --tender PLENNEGY-HCM-001
"""

import argparse, glob, os, re, sys
from dataclasses import dataclass, field
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)
from docx_renderer import (
    Seg, parse_md, add_inline, render_table, render_callout,
    render_ulist, render_olist, setup_styles,
    _page_break, _hr, _toc_field, _cell_bg, _cell_borders,
    NAVY, BLUE, GREY, LGREY, WHITE, ORANGE,
    HEX_NAVY, HEX_BLUE, HEX_WHITE, classify_comment, PROPOSALS_DIR
)
from proposal_shaper import (
    group_v2, classify_v2, _page_setup, _headers_footers,
    _ps, _ul, _tbl, _heading, LGREY2,
)

# ── Traceability types ────────────────────────────────────────────────────────
@dataclass
class SectionTrace:
    id: str; heading: str; status: str
    segs: int; kb_sources: List[str]; notes: str = ""

# ── Authored section content ──────────────────────────────────────────────────

def _authored_exec_summary() -> Tuple[List[Seg], SectionTrace]:
    segs = [
        Seg("comment", meta=(
            "PLACEHOLDER — Executive Summary. Authored by BU Lead and Account Manager. "
            "Must include: Plennegy-specific win theme; APPSolve positioning statement; "
            "3 key differentiators for this client; clear call to action. Maximum 2 pages. "
            "Context: Plennegy is an agribusiness group seeking Oracle HCM Cloud to unify "
            "HR across multiple operating entities."
        )),
    ]
    trace = SectionTrace("1", "Executive Summary", "PENDING", 0,
                         [], "Requires BU Lead authoring with Plennegy win themes")
    return segs, trace

def _authored_understanding() -> Tuple[List[Seg], SectionTrace]:
    segs = [
        Seg("comment", meta=(
            "PLACEHOLDER — Understanding of Plennegy's Requirements. Authored by Account Manager "
            "from RFP or qualification call notes. Must demonstrate APPSolve has understood "
            "Plennegy's specific context — not a generic statement. Include: business context, "
            "HR technology objectives, key challenges, what success looks like for Plennegy. "
            "Maximum 2 pages."
        )),
    ]
    trace = SectionTrace("2", "Understanding Plennegy's Requirements", "PENDING", 0,
                         [], "Requires Account Manager input from qualification")
    return segs, trace

def _authored_hcm_solution() -> Tuple[List[Seg], SectionTrace]:
    segs = [
        _ps("APPSolve proposes a phased Oracle Fusion HCM Cloud implementation that gives "
            "Plennegy Group a unified human capital management capability across all operating "
            "entities — from a candidate's first application to an employee's final day. "
            "Phase 1 delivers the core platform: HCM Core, Recruiting, and Payroll Integration. "
            "Phase 2 delivers Learning and Talent Management."),
        _heading("Proposed Scope", 3),
        _tbl([
            "| Oracle Module | Business Capability | Phase |",
            "|---|---|---|",
            "| **Oracle Fusion HCM Core** | Single employee record across all entities — hire to retire | 1 |",
            "| **Oracle Recruiting Cloud** | End-to-end recruitment — from vacancy to onboarded employee | 1 |",
            "| **Oracle Integration Cloud** | Payroll provider interface; third-party system connections | 1 + 2 |",
            "| **Oracle Learning Cloud** | Structured learning, certifications, compliance training | 2 |",
            "| **Oracle Talent Management** | Performance reviews, goals, succession planning | 2 |",
        ]),
        _heading("Oracle Fusion HCM Core — Your Workforce Foundation", 3),
        _ps("Plennegy's HR function manages employee data across multiple entities and systems. "
            "The compliance burden of absence management, position control, and employment "
            "contract management grows with every entity added to the group. Oracle Fusion "
            "HCM Core resolves this by establishing a single, authoritative employee record "
            "that the whole organisation works from."),
        _ps("APPSolve configures Oracle HCM Core to reflect how Plennegy actually operates — "
            "your organisation structure, position hierarchies, employment categories, and "
            "absence policies — not a generic template. Your HR team gains real-time visibility "
            "of headcount, absence status, and employment data across all entities from day one. "
            "Managers gain self-service capability for leave approvals, position changes, and "
            "routine HR requests without HR as an intermediary."),
        _ul([
            "**Business outcome:** HR spends less time reconciling data and more time managing "
            "people. Payroll accuracy improves because Oracle HCM data flows directly to your "
            "payroll provider — no manual rekeying.",
            "**Delivery confidence:** APPSolve has implemented Oracle HCM Core in South African "
            "enterprise organisations across gaming, retail, and agribusiness. We bring tested "
            "configuration templates and implementation runbooks to every engagement.",
        ]),
        _heading("Oracle Recruiting Cloud — Structured Talent Acquisition", 3),
        _ps("Oracle Recruiting Cloud replaces informal, email-based recruitment with a "
            "structured, trackable pipeline. Every hiring decision is logged, every candidate "
            "experience is consistent, and every hiring manager sees their pipeline in real "
            "time — from vacancy creation through to offer acceptance and onboarding."),
        _ul([
            "**Business outcome:** Time-to-hire reduces. Candidate quality improves because "
            "your sourcing channels are tracked and your recruitment process is consistent. "
            "Employment equity reporting is supported by Oracle's built-in reporting.",
            "**Delivery confidence:** APPSolve has implemented Oracle Recruiting Cloud in "
            "large South African organisations. We know which configuration choices matter most "
            "and where adoption risk is highest.",
        ]),
        _heading("Oracle Integration Cloud — Connected Operations", 3),
        _ps("Oracle Integration Cloud (OIC) connects Oracle Fusion HCM to Plennegy's payroll "
            "provider and any third-party HR systems. APPSolve's proven South African payroll "
            "integration pattern delivers bidirectional data flow — employee and position data "
            "from Oracle HCM to payroll processing, confirmed pay results back to Oracle — "
            "with exception handling and real-time monitoring built in."),
        _ul([
            "**Business outcome:** Your payroll team processes payroll from Oracle HCM data "
            "without manual rekeying. Errors surface immediately via the integration monitor, "
            "not in the payroll reconciliation report.",
            "**Customer benefit:** Payroll accuracy is a structural outcome, not a manual "
            "quality check.",
        ]),
        _heading("Oracle Learning Cloud — Structured Learning at Scale", 3),
        _ps("A multi-entity organisation like Plennegy faces increasing compliance training "
            "requirements across varied employment categories. Oracle Learning Cloud brings "
            "mandatory compliance training, skills development programmes, and certifications "
            "into a single platform — within Oracle HCM, without additional software licences."),
        _ul([
            "**Business outcome:** Compliance training requirements are met with automated "
            "evidence. Administrative overhead of tracking completion across entities is "
            "eliminated. Your learning investment is visible and measurable.",
        ]),
        _heading("Oracle Talent Management — People Development at Scale", 3),
        _ps("Oracle Talent Management gives Plennegy a structured, consistent performance cycle "
            "across all entities — goal setting, mid-year check-ins, year-end reviews — in a "
            "single system. Succession planning, talent profiling, and calibration tools give "
            "HR and executive leadership visibility of the people pipeline."),
        _ul([
            "**Business outcome:** Your performance review process is consistent and data-driven "
            "across all entities. Executive leadership has real-time talent pipeline visibility. "
            "Succession decisions are based on structured data, not informal knowledge.",
        ]),
        _heading("What Is Not in This Scope", 3),
        _ul([
            "Oracle Fusion Payroll (native) — APPSolve proposes OIC integration to Plennegy's existing payroll provider",
            "Oracle HR Help Desk — available as a Phase 3 extension if required",
            "Oracle Analytics Cloud — separately licensed Oracle add-on",
            "Custom software development beyond Oracle's Low-Code/No-Code extension framework",
        ]),
        _ps("Final scope, commercial proposal, and payment schedule will be confirmed in "
            "APPSolve's Commercial Proposal — issued separately once scope inputs are confirmed "
            "by Plennegy."),
    ]
    sources = [
        "Oracle Fusion HCM Capability (KB asset — synthesised)",
        "Oracle OIC / Integration Capability (KB asset — synthesised)",
        "Oracle Fusion Learning Cloud (KB asset — synthesised)",
        "Oracle Fusion Talent Management (KB asset — synthesised)",
    ]
    trace = SectionTrace("3", "Proposed Oracle HCM Solution", "AUTHORED",
                         sum(1 for s in segs if s.kind in ("paragraph","table","ulist","olist")),
                         sources, "Full KB assets (69+ pages) condensed to authored 6-page narrative")
    return segs, trace

def _authored_why_appsolve() -> Tuple[List[Seg], SectionTrace]:
    segs = [
        _ps("Oracle implementations succeed or fail on the quality of the implementation "
            "partner. The technology is Oracle's. The configuration decisions, the delivery "
            "discipline, the knowledge transfer, and the long-term support are the partner's. "
            "These are the factors that determine whether Plennegy realises the full value "
            "of Oracle HCM Cloud."),
        _heading("Our Oracle Capability", 3),
        _ps("Oracle is not a product we sell alongside others. It is the core of what "
            "APPSolve does. Our team of more than 50 Senior Consultants carries Oracle "
            "HCM, ERP, Integration Cloud, Database Administration, and Cloud Infrastructure "
            "expertise built across more than two decades and hundreds of engagements."),
        _ul([
            "**Oracle Level 1 Partner** — the foundational Oracle Partner Network tier "
            "recognising consistent delivery capability across Oracle Cloud applications",
            "**Oracle Business Impact Award EMEA 2024** — awarded for measurable business "
            "outcomes from our Oracle HCM implementations",
            "**Oracle ECEMEA Innovation Award 2024** — awarded for innovation in Oracle Cloud "
            "delivery in our region",
        ]),
        _heading("Our South African Enterprise Track Record", 3),
        _ps("We have implemented Oracle HCM Cloud in South African enterprise organisations "
            "across gaming, retail, agribusiness, and financial services. We understand the "
            "specific requirements of South African HCM — employment equity, SARS compliance, "
            "BCEA leave rules, and bargaining council agreements — and how they translate into "
            "Oracle configuration. We are not learning these requirements on Plennegy's time."),
        _ps("Our delivery footprint extends to Sub-Saharan Africa (Botswana, Namibia, "
            "Mozambique, Zambia, Tanzania) and internationally (USA, France, Abu Dhabi). "
            "For organisations operating across borders, we understand the multi-entity, "
            "multi-jurisdiction complexity that comes with scale."),
        _heading("Our Accountability Model", 3),
        _ul([
            "**A named Business Unit Lead** assigned to every engagement — a senior practitioner "
            "who takes personal responsibility for delivery quality and client satisfaction, "
            "not a project manager who escalates",
            "**Knowledge transfer by design** — every APPSolve engagement includes structured "
            "transfer so that Plennegy's HR and IT teams can operate and evolve the system "
            "independently after go-live",
            "**Long-term partnership** — our Application Managed Services practice provides "
            "post-go-live support, Oracle quarterly update management, and continuous "
            "optimisation; the majority of our clients continue with APPSolve for years after go-live",
        ]),
        _ps("Reference project details are available on request, subject to individual client "
            "approval. Contact your APPSolve Account Manager."),
    ]
    sources = [
        "Company Overview (KB asset — synthesised)",
        "Company History (KB asset — synthesised)",
        "Awards and Recognition (KB asset — synthesised)",
        "Oracle Partnership (KB asset — synthesised)",
        "Key Differentiators (KB asset — synthesised)",
        "Delivery Model (KB asset — synthesised)",
    ]
    trace = SectionTrace("4", "Why APPSolve", "AUTHORED",
                         sum(1 for s in segs if s.kind in ("paragraph","table","ulist","olist")),
                         sources, "7 KB company sections (76 content segs) condensed to 3-page narrative")
    return segs, trace

def _authored_how_we_deliver() -> Tuple[List[Seg], SectionTrace]:
    segs = [
        _ps("APPSolve delivers Oracle Fusion HCM Cloud using Oracle's Modern Best Practices "
            "and Oracle Success Navigator — the adoption framework Oracle has designed "
            "specifically for SaaS cloud implementations. We do not apply a traditional ERP "
            "implementation methodology to an Oracle Cloud engagement."),
        _ps("The distinction matters. Oracle Fusion HCM is a continuously evolving cloud "
            "platform — not a custom system. Our delivery philosophy is to configure Oracle "
            "to work for Plennegy, not to rebuild Plennegy's current processes in new "
            "technology. This approach produces faster go-lives, lower configuration risk, "
            "and lower long-term maintenance cost."),
        _heading("Our Delivery Philosophy", 3),
        _ul([
            "**Configure, not customise.** Oracle's platform is rich. Custom code adds cost, "
            "upgrade risk, and maintenance burden. We work within Oracle's standard configuration "
            "and Low-Code/No-Code extension framework.",
            "**Decisions enable delivery.** Oracle HCM requires Plennegy to make design "
            "decisions about organisation structure, approval hierarchies, security, and workflow. "
            "We structure workshops to drive decisions — not defer them.",
            "**Data quality is shared.** Your data quality determines your system quality. "
            "APPSolve provides templates, tools, and guidance. Plennegy's team owns the data "
            "quality process.",
            "**Go-live is the beginning.** Every engagement includes structured hypercare and "
            "handover to our Application Managed Services practice. We plan for long-term "
            "success from day one.",
        ]),
        _heading("The Six Delivery Phases", 3),
        _tbl([
            "| Phase | Name | Duration | What Plennegy Experiences |",
            "|---|---|---|---|",
            "| 1 | Mobilise | 4 weeks | Governance established; team confirmed; Oracle environments provisioned; kick-off |",
            "| 2 | Scope & Design | 8–10 weeks | Solution Design Workshops; Oracle configured to your organisation; CRP-1 |",
            "| 3 | Prototype | 6–8 weeks | Your configured Oracle system demonstrated in CRP-2; integrations built |",
            "| 4 | Validate | 6–8 weeks | Your team tests; data migration rehearsed; go-live readiness confirmed |",
            "| 5 | Deploy | 2–4 weeks | Go-live; APPSolve on-site hypercare support |",
            "| 6 | Evolve | 12+ months | Oracle quarterly updates managed; configuration optimised; AMS active |",
        ]),
        _heading("What Plennegy Can Expect", 3),
        _ul([
            "A named APPSolve Project Manager and Business Unit Lead for the full engagement",
            "Weekly project status reports: RAG by workstream, upcoming milestones, open risks, pending decisions",
            "A formal Change Request process for any scope changes — no surprises to timeline or cost",
            "Fully documented system at go-live: configuration workbooks, user guides, operational runbooks",
            "A structured hypercare period with direct APPSolve support before AMS handover",
            "Oracle quarterly cloud updates managed by APPSolve's AMS team after go-live",
        ]),
        _heading("Change Management", 3),
        _ps("Oracle HCM implementations succeed or fail on change management — the degree "
            "to which your people adopt the new system. APPSolve builds change management "
            "into every phase: stakeholder engagement in Mobilise, end-user training "
            "programme in Validate, structured support during Deploy. Change management is "
            "not a workshop before go-live. It is a continuous activity throughout the project."),
    ]
    sources = [
        "Implementation Methodology (KB asset — 140 content segs condensed to 4-page narrative)",
        "Project Governance (KB asset — synthesised)",
    ]
    trace = SectionTrace("5", "How We Deliver", "AUTHORED",
                         sum(1 for s in segs if s.kind in ("paragraph","table","ulist","olist")),
                         sources, "Methodology (140 segs ≈ 17.5 pages) condensed to Oracle Modern Best Practices narrative")
    return segs, trace

def _authored_project_governance() -> Tuple[List[Seg], SectionTrace]:
    segs = [
        _ps("APPSolve's project governance model provides clear accountability and "
            "transparent risk management at three levels: strategic, operational, and "
            "workstream. The model is established during the Mobilise phase and documented "
            "in the project charter."),
        _tbl([
            "| Level | Forum | Frequency | Led By | Purpose |",
            "|---|---|---|---|---|",
            "| Strategic | Steering Committee | Monthly | Executive Sponsors | Scope authority, commercial decisions, strategic risk |",
            "| Operational | Project Management | Weekly | Client PM + APPSolve PM | Progress, risks, actions, pending decisions |",
            "| Workstream | Module Reviews | Bi-weekly | Functional Leads | Design decisions, CRP planning, testing sign-off |",
        ]),
        _heading("Reporting Commitments", 3),
        _ul([
            "Weekly project status report to all stakeholders: RAG status by workstream, upcoming milestones, open risks, pending client decisions",
            "RAID Register (Risks, Assumptions, Issues, Dependencies) maintained and reviewed weekly",
            "Formal Change Request for every scope change — assessed for cost and timeline impact before approval",
            "Go-live readiness review: formal sign-off criteria before any go-live is confirmed",
        ]),
        _ps("Governance details, RAID Register template, Change Control procedure, and "
            "Cutover Plan are developed with Plennegy's Project Sponsor during the "
            "Mobilise phase."),
    ]
    trace = SectionTrace("6", "Project Governance", "AUTHORED",
                         sum(1 for s in segs if s.kind in ("paragraph","table","ulist","olist")),
                         ["Project Governance (KB asset — 140 segs condensed to 2-page governance table)"],
                         "Governance (140 segs ≈ 17.5 pages) condensed to governance structure table")
    return segs, trace

def _authored_roadmap() -> Tuple[List[Seg], SectionTrace]:
    segs = [
        _ps("The following roadmap provides indicative phase timelines. Durations are "
            "subject to confirmed scope, Oracle environment provisioning, and Plennegy "
            "resource availability. A detailed project plan with named milestones and "
            "resource assignments is developed in Mobilise."),
        _tbl([
            "| Phase | Name | Indicative Duration | Client Milestone |",
            "|---|---|---|---|",
            "| 1 | Mobilise | 4 weeks | Project charter signed; project team confirmed |",
            "| 2 | Scope & Design | 8–10 weeks | Solution design documents approved; CRP-1 complete |",
            "| 3 | Prototype | 6–8 weeks | CRP-2 demonstrated and signed off |",
            "| 4 | Validate | 6–8 weeks | UAT signed off; data migration approved; go-live confirmed |",
            "| 5 | Deploy | 2–4 weeks | Go-live; hypercare begins |",
            "| 6 | Evolve | 12+ months | AMS active; Oracle quarterly updates managed |",
        ]),
        _ul([
            "**Kick-off target:** within 2 weeks of contract signature",
            "**Go-live target:** to be confirmed with Plennegy once entity count, headcount, and module scope are confirmed",
            "Oracle environment provisioning is Oracle Corporation's responsibility and is outside APPSolve's control",
            "Phase 2 (Learning and Talent) commences after Phase 1 go-live is stable — typically 3–6 months post go-live",
        ]),
    ]
    trace = SectionTrace("7", "Implementation Roadmap", "AUTHORED",
                         sum(1 for s in segs if s.kind in ("paragraph","table","ulist","olist")),
                         ["Project Plan / Timeline (KB asset — synthesised)"],
                         "Phase timeline table with client milestones")
    return segs, trace

def _authored_key_assumptions() -> Tuple[List[Seg], SectionTrace]:
    segs = [
        _ps("This proposal is governed by a comprehensive assumption schedule that defines "
            "the contractual basis for scope, timeline, and cost. The following five "
            "commercial assumptions are the most material — they have the greatest impact "
            "on scope and pricing. The full assumption schedule is available on request and "
            "will form part of the Statement of Work."),
        _tbl([
            "| Commercial Assumption | Implication |",
            "|---|---|",
            "| **Platform:** Oracle provisions Oracle Fusion HCM Cloud SaaS environments; APPSolve's timeline commences from environment access date | Oracle environment provisioning timelines are outside APPSolve's control |",
            "| **Scope:** Implementation scope, timeline, and pricing are based on entity count, headcount, and module list confirmed at Mobilise | Any additions after Mobilise require a formal Change Request |",
            "| **Data:** Employee data extraction, cleansing, and validation are Plennegy's responsibility | APPSolve provides templates and guidance; Plennegy's team owns data quality |",
            "| **Integration:** Each third-party integration point is separately scoped and priced; payroll provider participates in design and testing | Undocumented APIs or legacy system access are out of scope |",
            "| **Resource:** Plennegy provides a nominated project team with confirmed availability for design workshops, testing, and sign-off decisions | Decision delays extend project timelines |",
        ]),
        _ps("A detailed Assumption Schedule covering all commercial, technical, functional, "
            "and operational assumptions will form part of APPSolve's Statement of Work. "
            "No scope item excluded from the Assumption Schedule is in scope unless "
            "explicitly added by Change Request."),
    ]
    trace = SectionTrace("8", "Key Commercial Assumptions", "AUTHORED",
                         sum(1 for s in segs if s.kind in ("paragraph","table","ulist","olist")),
                         ["Key Assumptions (KB asset — 497 segs condensed to 5-row table)"],
                         "341+ individual assumptions condensed to 5 commercial statements")
    return segs, trace

def _authored_key_risks() -> Tuple[List[Seg], SectionTrace]:
    segs = [
        _ps("The following risks are the most material delivery risks for this engagement. "
            "Each has been identified from APPSolve's experience with similar Oracle HCM "
            "implementations in South African enterprise environments. A full Risk Register "
            "will be maintained throughout the engagement."),
        _tbl([
            "| Risk | Probability | Impact | APPSolve Mitigation |",
            "|---|---|---|---|",
            "| **Scope ambiguity** — scope not confirmed before Mobilise | Medium | High | Formal scope confirmation session at contract signature; written confirmation of entity count, headcount, modules, and go-live date before kick-off |",
            "| **Data quality** — migration data provided late or with errors | High | High | APPSolve provides data migration templates in Scope & Design; data validation rehearsed in Prototype; final migration approved in Validate |",
            "| **Resource availability** — client team not available for workshops or decisions | Medium | High | Nominated client project team confirmed before Mobilise; availability commitments documented in project charter |",
            "| **Integration complexity** — undocumented APIs or unresponsive third-party vendors | Medium | Medium | Each integration separately scoped and vendor participation confirmed before build; contingency in integration timeline |",
            "| **Change adoption** — end users revert to previous systems after go-live | Medium | High | Change management built into all phases from Mobilise; end-user training delivered in Validate; hypercare support in Deploy |",
        ]),
        _ps("A full Risk Register with detailed risk ratings, individual risk owners, "
            "mitigation actions, and residual risk assessments will be established during "
            "Mobilise and maintained throughout the engagement."),
    ]
    trace = SectionTrace("9", "Key Delivery Risks and Mitigations", "AUTHORED",
                         sum(1 for s in segs if s.kind in ("paragraph","table","ulist","olist")),
                         ["Risk Register (KB asset — synthesised and expanded)"],
                         "5 material delivery risks with APPSolve mitigations")
    return segs, trace

def _authored_commercial() -> Tuple[List[Seg], SectionTrace]:
    segs = [
        _ps("APPSolve's formal commercial proposal will be issued by our Commercial Director "
            "once scope inputs are confirmed. The commercial proposal constitutes the binding "
            "offer and must be counter-signed before work commences. No pricing, rate, or "
            "commercial estimate in this document is binding."),
        _heading("Commercial Models Available", 3),
        _tbl([
            "| Model | Description | Best Suited To |",
            "|---|---|---|",
            "| Fixed Price | Agreed scope at a fixed cost; changes priced separately | Defined scope with low likelihood of change |",
            "| Time and Materials | Agreed rates per day; no fixed ceiling | Scope expected to evolve during delivery |",
            "| Monthly Recurring Invoice | Fixed monthly cost; capacity carries over | Predictable cost management; APPSolve's preferred model for sustained engagements |",
        ]),
        _ul([
            "**To receive APPSolve's Commercial Proposal:** confirm scope (modules, entities, headcount, go-live date) and preferred commercial model",
            "**Typical cycle:** 5–10 business days from scope confirmation to commercial proposal",
            "**Inclusions:** Commercial Proposal, Statement of Work, and compliance certificates submitted as a complete package",
        ]),
    ]
    trace = SectionTrace("10", "Commercial Position", "AUTHORED",
                         sum(1 for s in segs if s.kind in ("paragraph","table","ulist","olist")),
                         ["Commercial Assumptions (KB section — synthesised)"],
                         "Commercial framework — no pricing; scope confirmation required")
    return segs, trace

def _authored_next_steps() -> Tuple[List[Seg], SectionTrace]:
    segs = [
        _ps("APPSolve is ready to advance this engagement. The following actions are "
            "Plennegy's to complete:"),
        _tbl([
            "| Priority | Your Action | Purpose |",
            "|---|---|---|",
            "| **1 — Urgent** | Confirm Oracle HCM scope: modules, legal entities, headcount, payroll provider, go-live target | Required before commercial proposal can be issued |",
            "| **1 — Urgent** | Nominate a Plennegy Project Sponsor with authority over scope and commercial approval | APPSolve requires an identified decision-maker before mobilisation |",
            "| **2** | Confirm preferred commercial model: Fixed Price, Time and Materials, or Monthly Recurring Invoice | APPSolve structures commercial proposal accordingly |",
            "| **2** | Approve any client references for inclusion in this submission | Each reference requires individual client sign-off — advise your APPSolve Account Manager |",
            "| **2** | Confirm target go-live date | Required to plan phase timelines and resource availability |",
            "| **3** | Schedule scope confirmation call with APPSolve BU Lead and Account Manager | Enables commercial proposal within 5–10 business days |",
        ]),
    ]
    trace = SectionTrace("11", "Your Next Steps", "AUTHORED",
                         sum(1 for s in segs if s.kind in ("paragraph","table","ulist","olist")),
                         [],
                         "Client-actions-only next steps — no APPSolve internal tasks")
    return segs, trace

# ── Section map ───────────────────────────────────────────────────────────────

def _build_sections() -> Tuple[List[Tuple], List[SectionTrace]]:
    builders = [
        _authored_exec_summary,
        _authored_understanding,
        _authored_hcm_solution,
        _authored_why_appsolve,
        _authored_how_we_deliver,
        _authored_project_governance,
        _authored_roadmap,
        _authored_key_assumptions,
        _authored_key_risks,
        _authored_commercial,
        _authored_next_steps,
    ]
    sections = []; traces = []
    for fn in builders:
        segs, trace = fn()
        sections.append((trace.heading, segs))
        traces.append(trace)
    return sections, traces

# ── Shared DOCX helpers ───────────────────────────────────────────────────────

def _render_seg(doc, seg: Seg):
    if seg.kind == "heading":
        h = doc.add_heading(level=min(seg.level, 5)); h.paragraph_format.keep_with_next = True
        add_inline(h, seg.lines[0] if seg.lines else "",
                   sz={1:22,2:16,3:13,4:11,5:11}.get(seg.level,11))
    elif seg.kind == "table":   render_table(doc, seg.lines)
    elif seg.kind == "ulist":   render_ulist(doc, seg.lines)
    elif seg.kind == "olist":   render_olist(doc, seg.lines)
    elif seg.kind == "blockquote": render_callout(doc, seg.lines)
    elif seg.kind == "hr":
        from docx_renderer import _hr; _hr(doc)
    elif seg.kind == "paragraph":
        text = " ".join(seg.lines).strip()
        if text: p = doc.add_paragraph(style="Normal"); add_inline(p, text)
    elif seg.kind == "comment":
        kind = classify_comment(seg.meta)
        if kind in ("placeholder","ai_draft"):
            label = "Action Required" if kind == "placeholder" else "AI-Draft Required"
            render_callout(doc, [f"**{label}:** {seg.meta}"])

def _pending_box(doc, label: str):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    c = tbl.cell(0,0); _cell_bg(c,"F7F7F7"); _cell_borders(c,"AAAAAA",6)
    p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"[ {label} — content to be authored before submission ]")
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = LGREY2
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def _cover(doc, subtitle: str, tender_id: str, profile: str):
    doc.add_paragraph().paragraph_format.space_before = Pt(60)
    for txt, sz, bold, colour in [
        ("APPSolve (Pty) Ltd", 32, True, NAVY),
        ("Enterprise Technology Services  |  Oracle  |  Acumatica  |  BeBanking", 11, False, BLUE),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "Enterprise" in txt: p.paragraph_format.space_after = Pt(48)
        r = p.add_run(txt); r.font.name="Calibri"; r.font.size=Pt(sz); r.font.bold=bold; r.font.color.rgb=colour
    for txt, sz in [("ORACLE FUSION HCM CLOUD", 18), (subtitle, 14)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)
        r = p.add_run(txt); r.font.name="Calibri"; r.font.size=Pt(sz); r.font.bold=True; r.font.color.rgb=GREY
    pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER; pc.paragraph_format.space_after = Pt(40)
    rc = pc.add_run("Plennegy Group"); rc.font.name="Calibri"; rc.font.size=Pt(16); rc.font.bold=True; rc.font.color.rgb=NAVY
    tbl = doc.add_table(rows=0, cols=2); tbl.style = "Table Grid"
    for k, v in [("Prepared For","Plennegy Group"),("Platform","Oracle Fusion HCM Cloud"),
                 ("Prepared By","APPSolve (Pty) Ltd"),("Engine",f"PNE v1.0 — {profile}"),
                 ("Date",datetime.now().strftime("%Y-%m-%d")),("Tender ID",tender_id)]:
        row=tbl.add_row(); kc=row.cells[0]; vc=row.cells[1]
        _cell_bg(kc,HEX_NAVY); _cell_borders(kc,HEX_NAVY,4)
        _cell_bg(vc,"F2F5FB"); _cell_borders(vc,"D0D0D0",4)
        for cell, text, bold in [(kc,k,True),(vc,v,False)]:
            cp=cell.paragraphs[0]; cp.paragraph_format.space_before=Pt(4); cp.paragraph_format.space_after=Pt(4)
            cr=cp.add_run(text); cr.font.name="Calibri"; cr.font.size=Pt(11); cr.font.bold=bold
            cr.font.color.rgb = WHITE if cell is kc else GREY
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; fp.paragraph_format.space_before=Pt(48)
    fr=fp.add_run("COMMERCIAL IN CONFIDENCE  |  Proposal Narrative Engine v1.0")
    fr.font.name="Calibri"; fr.font.size=Pt(10); fr.font.bold=True; fr.font.color.rgb=LGREY2
    _page_break(doc)

# ── CLIENT PROPOSAL renderer ──────────────────────────────────────────────────

def render_client(sections: List[Tuple], out_path: str, tender_id: str):
    doc = Document(); setup_styles(doc); _page_setup(doc)
    _cover(doc, "IMPLEMENTATION PROPOSAL", tender_id, "CLIENT")
    h = doc.add_heading("Table of Contents", level=1); h.paragraph_format.space_before = Pt(0)
    _toc_field(doc); _page_break(doc)
    n_authored = 0; n_pending = 0
    for heading, segs in sections:
        _page_break(doc)
        doc.add_heading(heading, level=2).paragraph_format.keep_with_next = True
        has_ph = any(s.kind=="comment" and classify_comment(s.meta) in ("placeholder","ai_draft") for s in segs)
        if has_ph and not any(s.kind != "comment" for s in segs):
            label = re.sub(r"^\d+\.\s*","", heading)
            _pending_box(doc, label); n_pending += 1
        else:
            for seg in segs: _render_seg(doc, seg)
            n_authored += 1
    _headers_footers(doc, "APPSolve (Pty) Ltd",
                     "Oracle HCM Cloud Proposal — Plennegy Group  |  COMMERCIAL IN CONFIDENCE",
                     tender_id)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path); sz = os.path.getsize(out_path)
    print(f"[PNE CLIENT]      {out_path.split('/')[-1]}  ({sz//1024} KB)  "
          f"Authored={n_authored} Pending={n_pending}")
    return n_authored, n_pending

# ── INTERNAL REVIEW PACK renderer ─────────────────────────────────────────────

def render_internal(sections: List[Tuple], v2_secs, out_path: str, tender_id: str):
    doc = Document(); setup_styles(doc); _page_setup(doc)
    p1=doc.add_paragraph(); p1.alignment=WD_ALIGN_PARAGRAPH.CENTER; p1.paragraph_format.space_before=Pt(24)
    r1=p1.add_run("PNE Internal Review Pack — Plennegy Group")
    r1.font.name="Calibri"; r1.font.size=Pt(22); r1.font.bold=True; r1.font.color.rgb=NAVY
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after=Pt(8)
    r2=p2.add_run(f"{tender_id}  |  {datetime.now().strftime('%Y-%m-%d')}  |  PNE v1.0  |  INTERNAL ONLY")
    r2.font.name="Calibri"; r2.font.size=Pt(11); r2.font.color.rgb=BLUE
    render_callout(doc, ["**INTERNAL USE ONLY — APPSolve Confidential.** "
                         "Contains action register, authored content for review, and complete KB source content. "
                         "Do NOT distribute to clients."])
    # Action register
    _page_break(doc)
    doc.add_heading("Action Register — Pre-Submission Requirements", level=1).paragraph_format.space_before=Pt(0)
    at = doc.add_table(rows=1, cols=3); at.style="Table Grid"
    for ci, lbl in enumerate(["Priority","Action","Owner"]):
        hc=at.rows[0].cells[ci]; _cell_bg(hc,HEX_NAVY); _cell_borders(hc,HEX_NAVY,4)
        hp=hc.paragraphs[0]; hp.paragraph_format.space_before=Pt(3); hp.paragraph_format.space_after=Pt(3)
        add_inline(hp, lbl, sz=10.0)
        hp.runs[0].font.color.rgb=WHITE
    for prio, act, owner in [
        ("P0 URGENT","Renew B-BBEE Level 3 certificate","Finance Director — due 2026-07-31"),
        ("P1","Author Executive Summary (Sec 1) with Plennegy win themes","BU Lead + AM"),
        ("P1","Author Understanding of Requirements (Sec 2) from RFP/qualification notes","Account Manager"),
        ("P1","Confirm scope: modules, entities, headcount, payroll provider, go-live date (OAR-C04)","Account Manager"),
        ("P1","Issue Commercial Proposal once scope confirmed (OAR-C02)","Commercial Director"),
        ("P2","Hollywood Bets reference approval for this tender (OAR-C01)","Account Manager"),
        ("P2","BU Lead review: Sections 3–11 authored content for accuracy","BU Lead"),
        ("P3","Verify compliance document expiry dates","Operations"),
    ]:
        bg = "F8D7DA" if "P0" in prio or "P1" in prio else "FFF3CD" if "P2" in prio else "F2F5FB"
        row=at.add_row()
        for ci2, v in enumerate([prio, act, owner]):
            _cell_bg(row.cells[ci2],bg); _cell_borders(row.cells[ci2],"D0D0D0",4)
            cp=row.cells[ci2].paragraphs[0]; cp.paragraph_format.space_before=Pt(3); cp.paragraph_format.space_after=Pt(3)
            add_inline(cp,v,sz=9.5)
    # Part 1 — Authored sections
    _page_break(doc)
    doc.add_heading("Part 1 — Authored Sections (Published to Client)", level=1).paragraph_format.space_before=Pt(0)
    for heading, segs in sections:
        _page_break(doc)
        doc.add_heading(heading, level=2).paragraph_format.keep_with_next=True
        has_ph=any(s.kind=="comment" and classify_comment(s.meta) in ("placeholder","ai_draft") for s in segs)
        if has_ph and not any(s.kind!="comment" for s in segs):
            render_callout(doc, [f"**PENDING — {heading}:** Action required before submission. "
                                 "See Action Register above."])
        else:
            render_callout(doc, ["**PNE Authored Content** — Review for factual accuracy before submission."])
            for seg in segs: _render_seg(doc, seg)
    # Part 2 — KB source content
    _page_break(doc)
    doc.add_heading("Part 2 — Knowledge Base Source Content (Not Published to Client)", level=1).paragraph_format.space_before=Pt(0)
    p_note=doc.add_paragraph(style="Normal")
    add_inline(p_note, "The following V2 rendered KB sections are source material only — "
               "they are NOT published in the Client Proposal. They are retained for traceability "
               "and authoring reference.")
    for sec in v2_secs:
        if not sec.segs: continue
        _page_break(doc)
        doc.add_heading(sec.heading, level=2).paragraph_format.keep_with_next=True
        badge = {"RENDERED":"RENDERED","PLACEHOLDER":"PLACEHOLDER","AI_DRAFT":"AI-DRAFT","EMPTY":"EMPTY"}.get(sec.status,sec.status)
        p_b=doc.add_paragraph(style="Normal")
        add_inline(p_b, f"KB Source | Status: {badge} | {len(sec.segs)} segs | Not published to client", sz=9.0)
        for seg in sec.segs: _render_seg(doc, seg)
    _headers_footers(doc,"APPSolve INTERNAL REVIEW PACK",f"{tender_id}  |  PNE v1.0  |  INTERNAL",tender_id)
    os.makedirs(os.path.dirname(out_path),exist_ok=True)
    doc.save(out_path); sz=os.path.getsize(out_path)
    print(f"[PNE INTERNAL]    {out_path.split('/')[-1]}  ({sz//1024} KB)  "
          f"V2 secs={len(v2_secs)}")

# ── TRACEABILITY REPORT renderer ──────────────────────────────────────────────

def render_traceability(traces: List[SectionTrace], v2_secs, out_path: str, tender_id: str):
    doc = Document(); setup_styles(doc); _page_setup(doc)
    p1=doc.add_paragraph(); p1.alignment=WD_ALIGN_PARAGRAPH.CENTER; p1.paragraph_format.space_before=Pt(24)
    r1=p1.add_run("PNE Traceability Report — Plennegy Group")
    r1.font.name="Calibri"; r1.font.size=Pt(22); r1.font.bold=True; r1.font.color.rgb=NAVY
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after=Pt(8)
    r2=p2.add_run(f"{tender_id}  |  {datetime.now().strftime('%Y-%m-%d')}  |  PNE v1.0")
    r2.font.name="Calibri"; r2.font.size=Pt(11); r2.font.color.rgb=BLUE
    # Summary table
    doc.add_heading("Section Authoring Summary", level=1).paragraph_format.space_before=Pt(0)
    authored = [t for t in traces if t.status=="AUTHORED"]
    pending  = [t for t in traces if t.status=="PENDING"]
    score = round(len(authored) / len(traces) * 100)
    st = doc.add_table(rows=0, cols=2); st.style="Table Grid"
    for k, v in [
        ("Total sections", str(len(traces))),
        ("Authored by PNE", str(len(authored))),
        ("Pending (human action)", str(len(pending))),
        ("Proposal Readiness Score", f"{score}%"),
        ("Target page range", "25–35 pages"),
        ("KB source sections read", str(len(v2_secs))),
    ]:
        row=st.add_row()
        _cell_bg(row.cells[0],"F2F5FB"); _cell_borders(row.cells[0],"D0D0D0",4)
        _cell_bg(row.cells[1],HEX_WHITE); _cell_borders(row.cells[1],"D0D0D0",4)
        for ci,v2 in enumerate([k,v]):
            cp=row.cells[ci].paragraphs[0]; cp.paragraph_format.space_before=Pt(3); cp.paragraph_format.space_after=Pt(3)
            add_inline(cp,v2,sz=10.0)
    # Section-by-section trace
    _page_break(doc)
    doc.add_heading("Section Traceability", level=1).paragraph_format.space_before=Pt(0)
    tbl=doc.add_table(rows=1, cols=4); tbl.style="Table Grid"
    for ci, lbl in enumerate(["Section","Status","Content Segs","KB Source"]):
        hc=tbl.rows[0].cells[ci]; _cell_bg(hc,HEX_NAVY); _cell_borders(hc,HEX_NAVY,4)
        hp=hc.paragraphs[0]; hp.paragraph_format.space_before=Pt(3); hp.paragraph_format.space_after=Pt(3)
        add_inline(hp,lbl,sz=10.0); hp.runs[0].font.color.rgb=WHITE
    for t in traces:
        bg="D4EDDA" if t.status=="AUTHORED" else "F8D7DA"
        row=tbl.add_row()
        for ci2,v in enumerate([t.heading, t.status, str(t.segs), "\n".join(t.kb_sources) or "PNE authored"]):
            _cell_bg(row.cells[ci2],bg); _cell_borders(row.cells[ci2],"D0D0D0",4)
            cp=row.cells[ci2].paragraphs[0]; cp.paragraph_format.space_before=Pt(3); cp.paragraph_format.space_after=Pt(3)
            add_inline(cp,v,sz=9.0)
    # Compression table
    _page_break(doc)
    doc.add_heading("Content Compression — KB Source vs PNE Output", level=1).paragraph_format.space_before=Pt(0)
    render_callout(doc, ["These figures show the scale of content reduction from KB-verbatim rendering to "
                         "PNE-authored narrative. This is the core purpose of PNE: replacing knowledge asset "
                         "exports with authored proposal narrative."])
    ct=doc.add_table(rows=1, cols=3); ct.style="Table Grid"
    for ci,lbl in enumerate(["KB Section","Source Content Segs (≈ pages)","PNE Treatment"]):
        hc=ct.rows[0].cells[ci]; _cell_bg(hc,HEX_NAVY); _cell_borders(hc,HEX_NAVY,4)
        hp=hc.paragraphs[0]; hp.paragraph_format.space_before=Pt(3); hp.paragraph_format.space_after=Pt(3)
        add_inline(hp,lbl,sz=10.0); hp.runs[0].font.color.rgb=WHITE
    for src_name, segs_count, treatment in [
        ("Oracle Fusion HCM Capability", "555 (≈69 pages)", "Condensed to authored 6-page proposal narrative"),
        ("Oracle Fusion ERP Capability", "165 (≈21 pages)", "Not published — HCM is primary scope"),
        ("Oracle OIC / Integration Capability", "71 (≈9 pages)", "Condensed into Oracle HCM Solution section"),
        ("Implementation Methodology", "140 (≈18 pages)", "Condensed to 4-page Oracle Modern Best Practices narrative"),
        ("Project Governance", "140 (≈18 pages)", "Condensed to 2-page governance table"),
        ("Key Assumptions (Body Section)", "497 (≈62 pages)", "Condensed to 5-row commercial assumption table"),
        ("Complete Assumption Schedule", "494 (≈62 pages)", "Not published — covered by Key Commercial Assumptions"),
        ("Company sections (7 sections)", "76 (≈10 pages)", "Condensed to 3-page Why APPSolve narrative"),
    ]:
        row=ct.add_row(); _cell_bg(row.cells[0],"F8F9FA"); _cell_bg(row.cells[1],"FFF3CD"); _cell_bg(row.cells[2],"D4EDDA")
        for ci2,v in enumerate([src_name, segs_count, treatment]):
            _cell_borders(row.cells[ci2],"D0D0D0",4)
            cp=row.cells[ci2].paragraphs[0]; cp.paragraph_format.space_before=Pt(3); cp.paragraph_format.space_after=Pt(3)
            add_inline(cp,v,sz=9.0)
    _headers_footers(doc,"APPSolve TRACEABILITY REPORT",f"{tender_id}  |  PNE v1.0",tender_id)
    os.makedirs(os.path.dirname(out_path),exist_ok=True)
    doc.save(out_path); sz=os.path.getsize(out_path)
    print(f"[PNE TRACE]       {out_path.split('/')[-1]}  ({sz//1024} KB)")

# ── Orchestrator ──────────────────────────────────────────────────────────────

def run_pne(md_path: str, out_dir: str, tender_id: str) -> dict:
    print(f"[PNE] Tender: {tender_id}")
    with open(md_path,"r",encoding="utf-8") as fh: md_text = fh.read()
    segs    = parse_md(md_text)
    v2_secs = group_v2(segs)
    classify_v2(v2_secs)
    print(f"[PNE] V2: {len(v2_secs)} secs | "
          f"RENDERED={sum(1 for s in v2_secs if s.status=='RENDERED')} "
          f"PH={sum(1 for s in v2_secs if s.status=='PLACEHOLDER')} "
          f"AI={sum(1 for s in v2_secs if s.status=='AI_DRAFT')}")

    sections, traces = _build_sections()
    print(f"[PNE] Sections: {len(sections)} | "
          f"Authored={sum(1 for t in traces if t.status=='AUTHORED')} "
          f"Pending={sum(1 for t in traces if t.status=='PENDING')}")

    prefix = "PLENNEGY" if "PLENNEGY" in tender_id.upper() else tender_id
    os.makedirs(out_dir, exist_ok=True)

    c_path = os.path.join(out_dir, f"CLIENT_PROPOSAL_{prefix}_VNEXT.docx")
    i_path = os.path.join(out_dir, f"INTERNAL_REVIEW_PACK_{prefix}_VNEXT.docx")
    t_path = os.path.join(out_dir, f"TRACEABILITY_REPORT_{prefix}_VNEXT.docx")

    n_auth, n_pend = render_client(sections, c_path, tender_id)
    render_internal(sections, v2_secs, i_path, tender_id)
    render_traceability(traces, v2_secs, t_path, tender_id)

    print(f"\n[PNE] Complete — 3 outputs in {out_dir}")
    return {"client": c_path, "internal": i_path, "traceability": t_path}

# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Proposal Narrative Authoring Engine v1.0 (PNE)")
    parser.add_argument("--tender", metavar="TENDER_ID")
    parser.add_argument("--md",  metavar="MD_PATH")
    parser.add_argument("--out", metavar="OUT_DIR")
    args = parser.parse_args()

    if args.tender:
        tid=args.tender
        td_dir=os.path.join(PROPOSALS_DIR,tid)
        md_path=os.path.join(td_dir,f"PROPOSAL_RENDERED_{tid}.md")
        out_dir=td_dir
    elif args.md and args.out:
        tid=os.path.basename(args.out); md_path=args.md; out_dir=args.out
    else:
        parser.print_help(); sys.exit(1)

    if not os.path.exists(md_path):
        print(f"[PNE] ERROR: {md_path} not found",file=sys.stderr); sys.exit(1)
    run_pne(md_path, out_dir, tid)

if __name__ == "__main__":
    main()
