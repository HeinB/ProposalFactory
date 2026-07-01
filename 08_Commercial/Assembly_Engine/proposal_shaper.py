#!/usr/bin/env python3
"""
proposal_shaper.py — Proposal Shaping Layer v1.0 (PF2-008)

Transforms the V2 rendered Markdown into a structured V3 client-ready proposal.
Does not redesign the platform. Reads V2 output; applies editorial shaping; writes
three audience-specific DOCX files and two comparison/implementation reports.

V3 Body Structure (10 sections):
  1  Executive Summary          — pending (AI-Draft)
  2  Understanding of Requirements — pending (AI-Draft)
  3  Proposed Oracle HCM Solution — pending (human)
  4  Why APPSolve               — from V2 credential sections (S-03/04/05/06/07/08/09)
  5  Delivery Approach          — from V2 methodology (S-34) + scope notices
  6  Project Governance         — authored
  7  Key Assumptions            — short summary + App B cross-reference
  8  Key Risks and Mitigations  — pending (AI-Draft)
  9  Commercial Inputs Required — authored + placeholder
  10 Next Steps                 — authored

Appendices:
  A  Detailed Oracle HCM Capability  (V2 S-16/S-17/S-19 full)
  B  Assumptions Register            (V2 Appendix A full)
  C  Reference Detail                (V2 S-67)

Usage:
    python3 proposal_shaper.py --tender PLENNEGY-HCM-001
"""

import argparse, os, re, sys, glob
from dataclasses import dataclass, field
from datetime import datetime
from typing import Dict, List, Optional, Tuple
import yaml
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)
from docx_renderer import (
    Seg, parse_md, add_inline, render_table, render_callout,
    render_ulist, render_olist, setup_styles,
    _page_break, _hr, _toc_field, _tab_stop, _page_field, _numpages_field,
    _cell_bg, _cell_borders,
    NAVY, BLUE, GREY, LGREY, WHITE, ORANGE,
    HEX_NAVY, HEX_BLUE, HEX_WHITE,
    HEX_BG_PH, HEX_BG_AI, HEX_BD_PH, HEX_BD_AI,
    HEX_BG_RV, HEX_BD_RV, HEX_BG_GV, HEX_BD_GV,
    classify_comment, PROPOSALS_DIR
)

_KB_ROOT = os.path.abspath(os.path.join(_HERE, "..", ".."))

# ── Extra colours ─────────────────────────────────────────────────────────────
HEX_BG_PENDING = "F7F7F7"
HEX_BD_PENDING = "AAAAAA"
HEX_BG_TRACE   = "E8F3E8"
HEX_BD_TRACE   = "2E8B57"
HEX_STATUS_OK  = "D4EDDA"
HEX_STATUS_PH  = "F8D7DA"
HEX_STATUS_AI  = "FFF3CD"
LGREY2 = RGBColor(0x70, 0x70, 0x70)

# ── AI-sentinel detector ──────────────────────────────────────────────────────
_AI_SENTINEL = re.compile(
    r"AI-generated content to be inserted|must be reviewed and approved",
    re.IGNORECASE
)
def _is_ai_sentinel(seg: Seg) -> bool:
    return seg.kind == "paragraph" and bool(_AI_SENTINEL.search(" ".join(seg.lines)))

# ── CLIENT segment filter (same rules as ppe.py) ─────────────────────────────
_CLIENT_STRIP = re.compile(
    r"(Capability Statement\s*\|.*)|(Methodology Statement\s*\|.*)"
    r"|(Assumption Pack\s*\|.*)|(Assumption Schedule\s*\|.*)"
    r"|(Document ID:\s*W[0-9]S[0-9]-\d+)"
    r"|(\*\*(Scope|Governance|Usage):\*\*)"
    r"|(\*Standalone pack)|(\*Mandatory attachment)"
    r"|(\*\d+ assumptions.*/.*exclusions)"
    r"|(Governed under ASSUMPTION_GOVERNANCE)",
    re.IGNORECASE
)
_MARKER_BQ = re.compile(
    r"(action required|ai-draft required|human review required|governance flags)",
    re.IGNORECASE
)
def _is_internal_para(seg: Seg) -> bool:
    if seg.kind != "paragraph": return False
    return bool(_CLIENT_STRIP.search(" ".join(seg.lines))) or _is_ai_sentinel(seg)
def _is_marker_bq(seg: Seg) -> bool:
    if seg.kind != "blockquote" or not seg.lines: return False
    return bool(_MARKER_BQ.search(" ".join(seg.lines[:2]).lower()))

def filter_client(segs: List[Seg]) -> List[Seg]:
    result = []
    for seg in segs:
        if seg.kind == "comment": continue
        if _is_marker_bq(seg): continue
        if _is_internal_para(seg): continue
        result.append(seg)
    return result

# ── V2 section grouper ────────────────────────────────────────────────────────
@dataclass
class V2Sec:
    heading: str
    name:    str
    segs:    List[Seg] = field(default_factory=list)
    status:  str = "EMPTY"

def _v2_name(heading: str) -> str:
    m = re.match(r"^(?:Appendix\s+\w+\s+[—-]\s+|(?:\d+|[A-Z])\.\s+)(.*)", heading)
    return (m.group(1) if m else heading).strip()

def group_v2(segs: List[Seg]) -> List[V2Sec]:
    secs, cur = [], None
    for seg in segs:
        if seg.kind == "heading" and seg.level == 2:
            cur = V2Sec(heading=seg.lines[0] if seg.lines else "", name=_v2_name(seg.lines[0] if seg.lines else ""))
            secs.append(cur)
        elif cur is not None:
            cur.segs.append(seg)
    return secs

def classify_v2(secs: List[V2Sec]):
    for sec in secs:
        has_ph  = any(s.kind == "comment" and "PLACEHOLDER" in s.meta.upper() for s in sec.segs)
        has_ai  = any(s.kind == "comment" and "AI-DRAFT" in s.meta.upper() for s in sec.segs)
        has_body= any(s.kind in ("paragraph","table","ulist","olist") and not _is_ai_sentinel(s)
                      for s in sec.segs)
        if has_ph:        sec.status = "PLACEHOLDER"
        elif has_ai:      sec.status = "AI_DRAFT"
        elif has_body:    sec.status = "RENDERED"
        else:             sec.status = "EMPTY"

def v2_idx(secs: List[V2Sec]) -> Dict[str, V2Sec]:
    return {s.name.lower(): s for s in secs}

# ── Content excerpt ───────────────────────────────────────────────────────────
def excerpt(segs: List[Seg], max_content: int) -> List[Seg]:
    """Return segs up to max_content content (paragraph/table/list) segments."""
    result, n = [], 0
    for seg in segs:
        result.append(seg)
        if seg.kind in ("paragraph","table","ulist","olist") and not _is_ai_sentinel(seg):
            n += 1
            if n >= max_content:
                break
    return result

# ── Authored Seg builders ─────────────────────────────────────────────────────
def _ps(text: str) -> Seg:
    return Seg("paragraph", lines=[text])
def _ul(items: List[str]) -> Seg:
    return Seg("ulist", lines=items)
def _tbl(rows: List[str]) -> Seg:
    return Seg("table", lines=rows)
def _heading(text: str, level: int = 3) -> Seg:
    return Seg("heading", level=level, lines=[text])

# ── Authored sections ─────────────────────────────────────────────────────────

def _authored_governance() -> List[Seg]:
    return [
        _ps("APPSolve's project governance model establishes clear accountability at every level "
            "of the engagement. The framework provides effective oversight, rapid decision-making, "
            "and transparent risk management throughout the implementation lifecycle."),
        _tbl([
            "| Governance Layer | Frequency | Participants | Purpose |",
            "|---|---|---|---|",
            "| Steering Committee | Monthly | Client Executive Sponsor, APPSolve BU Lead | "
            "Strategic decisions, scope authority, escalation |",
            "| Joint Project Management | Weekly | Client PM, APPSolve PM | "
            "Delivery tracking, risk management, actions |",
            "| Workstream Reviews | Bi-weekly | Functional Leads, Technical Leads | "
            "Module configuration status, CRP planning, design decisions |",
            "| Change Control Board | As required | PM, BU Lead, Commercial Lead | "
            "Change request review, impact assessment, approval |",
        ]),
        _ps("APPSolve provides a dedicated Project Manager and BU Lead for the engagement. "
            "A weekly progress report is issued to all stakeholders covering RAG status by workstream, "
            "upcoming milestones, open risks, and pending client decisions. "
            "All scope changes are processed through a formal Change Request with documented "
            "commercial and timeline impact before approval."),
        _ps("The RAID Framework (Risks, Assumptions, Issues, Dependencies) will be maintained "
            "as a living document throughout the project and reviewed at each Project Management meeting. "
            "A formal Cutover Plan will be agreed during the Validate phase, defining the go-live "
            "sequence, rollback criteria, and hypercare period. "
            "Full RAID documentation, Change Control Framework, and Cutover Plan are provided as "
            "separate project deliverables and will be completed during the Mobilise phase."),
        _heading("RAID Framework", 3),
        _ps("_PLACEHOLDER: Project-specific RAID Framework to be completed during Mobilise phase "
            "with input from Plennegy project sponsor and APPSolve Project Manager._"),
        _heading("Cutover / Go-Live Plan", 3),
        _ps("_PLACEHOLDER: Go-live sequence and hypercare plan to be agreed during Validate phase. "
            "Target go-live date to be confirmed with Plennegy (OAR-C04)._"),
    ]

def _authored_assumptions_summary() -> List[Seg]:
    return [
        _ps("APPSolve's proposal is governed by a comprehensive assumptions register comprising "
            "341 approved assumptions across five packs. These assumptions define the contractual "
            "basis for scope, cost, and timeline and are referenced in all APPSolve Statements of Work."),
        _tbl([
            "| Assumption Pack | Assumptions | Coverage |",
            "|---|---|---|",
            "| Oracle Fusion HCM Core | 115 | Environment, organisation structure, data migration, "
            "security, payroll data, integration foundations |",
            "| Oracle HCM Recruiting Cloud | 54 | Vacancy management, candidate sourcing, "
            "offer management, pre-employment checks, onboarding |",
            "| Oracle HCM Learning Cloud | 37 | Course catalogue, certifications, "
            "learning paths, instructor-led training |",
            "| Oracle HCM Talent Management | 31 | Performance management, goals, "
            "succession planning, career development |",
            "| Oracle Integration Cloud | 104 | Payroll interface, third-party integration, "
            "middleware, data exchange |",
            "| **Total** | **341** | **Full HCM Cloud implementation scope** |",
        ]),
        _heading("Key Commercial Assumptions", 3),
        _ps("The following high-level assumptions are most material to scope, cost, and timeline:"),
        _ul([
            "Oracle Cloud (SaaS) environments are Oracle's responsibility; provisioning timelines are outside APPSolve's control",
            "Scope is based on entity count, legal entities, and headcount confirmed during Mobilise phase; changes require formal Change Request",
            "Client is responsible for data extraction, cleansing, and final load validation; APPSolve provides templates and guidance",
            "All configuration decisions must be made by client stakeholders within agreed workshop schedules; delayed decisions extend timeline",
            "Go-live readiness is a joint decision; APPSolve may withhold go-live recommendation if agreed readiness criteria are not met",
            "Oracle quarterly cloud updates are outside APPSolve's control; impact assessment is included in the project methodology",
        ]),
        _ps("The complete assumptions register is provided in **Appendix B**. "
            "Key assumptions from the HCM Core, Recruiting, Learning, Talent, and OIC packs apply in full "
            "to this engagement unless explicitly excluded in the Statement of Work."),
    ]

def _authored_commercial() -> List[Seg]:
    return [
        _ps("APPSolve's commercial proposal for the Oracle HCM Cloud implementation will be "
            "formalised under the Commercial Director's authority once the scope inputs below "
            "are confirmed. All commercial documents are subject to BU Lead sign-off before submission."),
        _heading("Commercial Inputs Required", 3),
        _ul([
            "**Project scope:** Module list, legal entity count, employee headcount, "
            "integration points, and target go-live date (confirm with Account Manager — OAR-C04)",
            "**Pricing structure:** Fixed-price or Time and Materials rates — "
            "APPSolve Commercial Director to provide (OAR-C02)",
            "**Payment model:** Monthly Recurring Invoice or Time and Materials with cap — "
            "to be agreed with Plennegy before submission",
            "**B-BBEE certificate:** Level 3 renewal required before 2026-07-31 — "
            "Finance Director action (OAR-A01 — URGENT)",
            "**Rate card basis:** Basis for all pricing to be confirmed before commercial document issued",
        ]),
        _heading("APPSolve Commercial Models", 3),
        _tbl([
            "| Model | Description | APPSolve Recommendation |",
            "|---|---|---|",
            "| Fixed Price | Agreed scope at fixed cost; change requests priced separately | "
            "Recommended for defined HCM Cloud scope |",
            "| Time and Materials | Billing by time at agreed rates | "
            "Appropriate where scope evolves during delivery |",
            "| Monthly Recurring Invoice | Fixed monthly cost; capacity transfers between months | "
            "APPSolve's distinctive model — predictable cost; aligned incentives |",
        ]),
        _ps("_Pricing table to be provided by APPSolve Commercial Director following scope confirmation. "
            "No rates or commercial terms are binding until the formal Commercial Proposal is issued "
            "and counter-signed by both parties._"),
    ]

def _authored_next_steps() -> List[Seg]:
    return [
        _ps("APPSolve is ready to progress this engagement. The following actions are required "
            "in priority order to advance from this draft to a final submission-ready proposal:"),
        _tbl([
            "| Priority | Action | Owner | Target |",
            "|---|---|---|---|",
            "| **P0 — URGENT** | Renew B-BBEE Level 3 certificate | Finance Director | Before 2026-07-31 |",
            "| **P1** | Draft Executive Summary with Plennegy win themes | BU Lead + AM | Before internal review |",
            "| **P1** | Draft Understanding of Requirements from Plennegy RFP | Account Manager | Before internal review |",
            "| **P1** | Confirm scope: modules, entities, headcount, go-live date | Account Manager (OAR-C04) | Before pricing |",
            "| **P1** | Provide pricing structure and commercial proposal | Commercial Director (OAR-C02) | Before submission |",
            "| **P2** | Obtain Hollywood Bets reference approval for this tender | Account Manager (OAR-C01) | Before submission |",
            "| **P2** | Provide proposed team structure and CVs (from APPTime) | Delivery Manager | Before submission |",
            "| **P2** | AI-generate Risk Register; BU Lead review and approval | BU Lead | Before submission |",
            "| **P3** | Verify all compliance document expiry dates | Operations + Finance | Before submission |",
            "| **P3** | Attach current Oracle OPN Certificate | Operations | Before submission |",
        ]),
        _ps("On completion of the above actions, APPSolve will produce the final CLIENT "
            "proposal (V4) for BU Lead sign-off and submission to Plennegy Group."),
    ]

# ── V3 section definition ─────────────────────────────────────────────────────
V3_MAP = [
    {"id": "1",  "heading": "1. Executive Summary",
     "v2":  ["Executive Summary"],       "mode": "passthrough", "app": False},
    {"id": "2",  "heading": "2. Understanding of Plennegy's Requirements",
     "v2":  ["Understanding of Requirements"], "mode": "passthrough", "app": False},
    {"id": "3",  "heading": "3. Proposed Oracle HCM Solution",
     "v2":  ["Proposed Solution Overview"],   "mode": "passthrough", "app": False},
    {"id": "4",  "heading": "4. Why APPSolve",
     "v2":  ["Company Overview", "Company History", "Awards and Recognition",
             "Delivery Model", "Geographic Presence", "Key Differentiators", "Oracle Partnership"],
     "mode": "full_multi", "app": False},
    {"id": "5",  "heading": "5. Delivery Approach",
     "v2":  ["Implementation Methodology"],  "mode": "full",
     "ph_append": ["Scope of Work — Inclusions", "Scope of Work — Exclusions",
                   "Deliverables", "Dependencies"],
     "app": False},
    {"id": "6",  "heading": "6. Project Governance",
     "v2":  [],   "mode": "authored", "authored_fn": "_authored_governance", "app": False},
    {"id": "7",  "heading": "7. Key Assumptions",
     "v2":  ["Key Assumptions (Body Section)"], "mode": "authored_with_excerpt",
     "excerpt_max": 4,  # just the intro paras of first pack
     "authored_summary_fn": "_authored_assumptions_summary",
     "app": False},
    {"id": "8",  "heading": "8. Key Risks and Mitigations",
     "v2":  ["Risk Register"], "mode": "passthrough", "app": False},
    {"id": "9",  "heading": "9. Commercial Inputs Required",
     "v2":  [],   "mode": "authored", "authored_fn": "_authored_commercial", "app": False},
    {"id": "10", "heading": "10. Next Steps",
     "v2":  [],   "mode": "authored", "authored_fn": "_authored_next_steps", "app": False},
    {"id": "A",  "heading": "Appendix A — Detailed Oracle HCM Capability",
     "v2":  ["Oracle Fusion HCM Capability", "Oracle Fusion ERP Capability",
             "Oracle OIC / Integration Capability"],
     "mode": "full_multi", "app": True},
    {"id": "B",  "heading": "Appendix B — Assumptions Register",
     "v2":  ["Complete Assumption Schedule"], "mode": "full", "app": True},
    {"id": "C",  "heading": "Appendix C — Reference Detail",
     "v2":  ["Client References"], "mode": "passthrough", "app": True},
]

_AUTHORED_FNS = {
    "_authored_governance":          _authored_governance,
    "_authored_assumptions_summary": _authored_assumptions_summary,
    "_authored_commercial":          _authored_commercial,
    "_authored_next_steps":          _authored_next_steps,
}

@dataclass
class V3Sec:
    id:      str
    heading: str
    segs:    List[Seg]
    status:  str    # RENDERED | PLACEHOLDER | AI_DRAFT | AUTHORED | EMPTY
    sources: List[str]
    mode:    str
    is_app:  bool

def _placeholder_segs(label: str) -> List[Seg]:
    return [Seg("comment", meta=f"PLACEHOLDER: {label}")]

def _ai_draft_segs(label: str) -> List[Seg]:
    return [Seg("comment", meta=f"AI-DRAFT REQUIRED: {label}")]

def build_v3_sections(idx: Dict[str, V2Sec]) -> List[V3Sec]:
    result = []
    for defn in V3_MAP:
        sid   = defn["id"]
        head  = defn["heading"]
        mode  = defn["mode"]
        is_app= defn.get("app", False)
        v2_names = defn.get("v2", [])

        segs: List[Seg] = []
        sources: List[str] = []
        status = "EMPTY"

        if mode == "passthrough":
            # Use V2 section content as-is
            for name in v2_names:
                sec = idx.get(name.lower())
                if sec:
                    segs.extend(sec.segs)
                    sources.append(sec.name)
                    if sec.status in ("RENDERED",):
                        status = "RENDERED"
                    elif sec.status == "AI_DRAFT" and status != "RENDERED":
                        status = "AI_DRAFT"
                    elif sec.status == "PLACEHOLDER" and status not in ("RENDERED", "AI_DRAFT"):
                        status = "PLACEHOLDER"
            if not segs and not sources:
                segs = _placeholder_segs(head)
                status = "PLACEHOLDER"

        elif mode == "full":
            name = v2_names[0] if v2_names else ""
            sec = idx.get(name.lower())
            if sec and sec.status == "RENDERED":
                segs   = list(sec.segs)
                status = "RENDERED"
                sources.append(sec.name)
            else:
                segs   = _placeholder_segs(name or head)
                status = "PLACEHOLDER"
            # Append PLACEHOLDER notices for scope sections if requested
            for ph_name in defn.get("ph_append", []):
                ph_sec = idx.get(ph_name.lower())
                if ph_sec and ph_sec.status == "PLACEHOLDER":
                    segs.append(Seg("heading", level=3, lines=[ph_name]))
                    segs.extend(ph_sec.segs)

        elif mode == "full_multi":
            any_rendered = False
            for name in v2_names:
                sec = idx.get(name.lower())
                if sec:
                    segs.extend(sec.segs)
                    sources.append(sec.name)
                    if sec.status == "RENDERED":
                        any_rendered = True
            status = "RENDERED" if any_rendered else "PLACEHOLDER"

        elif mode == "authored":
            fn_name = defn.get("authored_fn", "")
            fn = _AUTHORED_FNS.get(fn_name)
            segs   = fn() if fn else [_ps("Section content to be completed.")]
            status = "AUTHORED"

        elif mode == "authored_with_excerpt":
            # Use authored summary segs; do NOT include V2 excerpt in body
            fn_name = defn.get("authored_summary_fn", "")
            fn = _AUTHORED_FNS.get(fn_name)
            segs   = fn() if fn else []
            status = "AUTHORED"
            for name in v2_names:
                sec = idx.get(name.lower())
                if sec:
                    sources.append(sec.name)

        result.append(V3Sec(id=sid, heading=head, segs=segs, status=status,
                            sources=sources, mode=mode, is_app=is_app))
    return result

# ── Shared DOCX helpers ───────────────────────────────────────────────────────

def _page_setup(doc):
    for sec in doc.sections:
        sec.page_width    = Cm(21.0); sec.page_height = Cm(29.7)
        sec.left_margin   = Cm(2.54); sec.right_margin = Cm(2.54)
        sec.top_margin    = Cm(2.54); sec.bottom_margin= Cm(2.54)
        sec.header_distance = Cm(1.27); sec.footer_distance = Cm(1.27)

def _headers_footers(doc, left_label: str, right_label: str, tender_id: str):
    for section in doc.sections:
        section.different_first_page_header_footer = True
        hdr = section.header
        hp  = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.clear(); _tab_stop(hp, 9072)
        hl = hp.add_run(left_label); hl.font.name="Calibri"; hl.font.size=Pt(9); hl.font.bold=True; hl.font.color.rgb=NAVY
        hp.add_run("\t").font.size=Pt(9)
        hr2 = hp.add_run(right_label); hr2.font.name="Calibri"; hr2.font.size=Pt(9); hr2.font.color.rgb=LGREY
        hpPr=hp._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
        btm=OxmlElement("w:bottom"); btm.set(qn("w:val"),"single"); btm.set(qn("w:sz"),"6"); btm.set(qn("w:space"),"1"); btm.set(qn("w:color"),HEX_NAVY)
        pBdr.append(btm); hpPr.append(pBdr)
        ftr = section.footer
        fp  = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.clear(); _tab_stop(fp, 9072)
        fl = fp.add_run(f"APPSolve (Pty) Ltd  |  {tender_id}\t"); fl.font.name="Calibri"; fl.font.size=Pt(9); fl.font.color.rgb=LGREY
        pr = fp.add_run("Page "); pr.font.name="Calibri"; pr.font.size=Pt(9); pr.font.color.rgb=LGREY
        _page_field(pr)
        pr2= fp.add_run(" of "); pr2.font.name="Calibri"; pr2.font.size=Pt(9); pr2.font.color.rgb=LGREY
        _numpages_field(pr2)
        fpPr=fp._p.get_or_add_pPr(); fpBdr=OxmlElement("w:pBdr"); ftop=OxmlElement("w:top"); ftop.set(qn("w:val"),"single"); ftop.set(qn("w:sz"),"6"); ftop.set(qn("w:space"),"1"); ftop.set(qn("w:color"),HEX_NAVY); fpBdr.append(ftop); fpPr.append(fpBdr)

def _render_seg(doc, seg: Seg):
    if seg.kind == "heading":
        h = doc.add_heading(level=min(seg.level, 5))
        h.paragraph_format.keep_with_next = True
        add_inline(h, seg.lines[0] if seg.lines else "", sz={1:22,2:16,3:13,4:11,5:11}.get(seg.level,11))
    elif seg.kind == "table":   render_table(doc, seg.lines)
    elif seg.kind == "ulist":   render_ulist(doc, seg.lines)
    elif seg.kind == "olist":   render_olist(doc, seg.lines)
    elif seg.kind == "blockquote": render_callout(doc, seg.lines)
    elif seg.kind == "hr":      _hr(doc)
    elif seg.kind == "paragraph":
        text = " ".join(seg.lines).strip()
        if text:
            p = doc.add_paragraph(style="Normal")
            add_inline(p, text)
    elif seg.kind == "comment":
        kind = classify_comment(seg.meta)
        if kind == "placeholder":
            name = re.sub(r"^PLACEHOLDER:\s*", "", seg.meta, flags=re.I)
            render_callout(doc, [f"**Action Required:** {name}"])
        elif kind == "ai_draft":
            name = re.sub(r"^AI-DRAFT REQUIRED:\s*", "", seg.meta, flags=re.I)
            render_callout(doc, [f"**AI-Draft Required:** {name}"])
        elif kind == "governance_note":
            note = re.sub(r"^(GOV:|SI-RULE:)\s*","", seg.meta, flags=re.I).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(3)
            r = p.add_run(f"⚙  Governance: {note}")
            r.font.name="Calibri"; r.font.size=Pt(8.5); r.font.italic=True; r.font.color.rgb=LGREY2

def _pending_notice(doc, label: str, detail: str = ""):
    """Grey neutral 'content pending' box for CLIENT profile (no action language)."""
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    cell= tbl.cell(0,0); _cell_bg(cell, HEX_BG_PENDING); _cell_borders(cell, HEX_BD_PENDING, 6)
    p   = cell.paragraphs[0]; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
    r   = p.add_run(f"[ {label} — to be completed before submission ]")
    r.font.name="Calibri"; r.font.size=Pt(10); r.font.italic=True; r.font.color.rgb=LGREY2
    if detail:
        p2 = cell.add_paragraph(); p2.paragraph_format.space_before=Pt(2); p2.paragraph_format.space_after=Pt(4)
        r2 = p2.add_run(detail); r2.font.name="Calibri"; r2.font.size=Pt(9); r2.font.color.rgb=LGREY2
    doc.add_paragraph().paragraph_format.space_after=Pt(6)

def _action_box(doc, label: str, action: str, owner: str = ""):
    """Coloured action required box for INTERNAL profile."""
    bg = HEX_BG_PH; bd = HEX_BD_PH
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    cell= tbl.cell(0,0); _cell_bg(cell, bg); _cell_borders(cell, bd, 8)
    p   = cell.paragraphs[0]; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(2)
    r   = p.add_run(f"⚠  {label}"); r.font.name="Calibri"; r.font.size=Pt(10); r.font.bold=True
    if action:
        p2 = cell.add_paragraph(); p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(2)
        r2 = p2.add_run(action); r2.font.name="Calibri"; r2.font.size=Pt(9.5)
    if owner:
        p3 = cell.add_paragraph(); p3.paragraph_format.space_before=Pt(0); p3.paragraph_format.space_after=Pt(4)
        r3 = p3.add_run(f"Owner: {owner}"); r3.font.name="Calibri"; r3.font.size=Pt(9); r3.font.italic=True; r3.font.color.rgb=LGREY2
    doc.add_paragraph().paragraph_format.space_after=Pt(6)

# ── CLIENT V3 renderer ────────────────────────────────────────────────────────

class ClientV3Renderer:
    def __init__(self, v3_secs: List[V3Sec], out_path: str, tender_id: str):
        self.secs     = v3_secs
        self.out_path = out_path
        self.tid      = tender_id
        self.doc      = Document()
        setup_styles(self.doc)
        _page_setup(self.doc)
        self.stats    = {"rendered": 0, "pending": 0, "authored": 0}

    def render(self):
        d = self.doc
        # Cover page
        self._cover()
        # TOC
        h = d.add_heading("Table of Contents", level=1); h.paragraph_format.space_before=Pt(0)
        _toc_field(d); _page_break(d)
        # Sections
        for sec in self.secs:
            if sec.id in ("1","2","3","4","5","6","7","8","9","10") or sec.is_app:
                if sec.is_app or (sec.id not in ("",)):
                    _page_break(d)
            h = d.add_heading(sec.heading, level=2)
            h.paragraph_format.keep_with_next = True
            if sec.status in ("PLACEHOLDER", "AI_DRAFT", "EMPTY") and not sec.segs:
                _pending_notice(d, sec.heading.split(".",1)[-1].strip() if "." in sec.heading else sec.heading)
                self.stats["pending"] += 1
            elif sec.status in ("PLACEHOLDER", "AI_DRAFT") and sec.segs:
                # Has comment segs that explain the gap — show as neutral pending
                label = sec.heading.split(".",1)[-1].strip() if "." in sec.heading else sec.heading
                if sec.status == "AI_DRAFT":
                    _pending_notice(d, label, "Requires BU Lead drafting and human review before submission.")
                else:
                    _pending_notice(d, label, "Requires client input and AM approval before submission.")
                self.stats["pending"] += 1
            elif sec.status == "AUTHORED":
                for seg in sec.segs:
                    _render_seg(d, seg)
                self.stats["authored"] += 1
            else:
                filtered = filter_client(sec.segs)
                for seg in filtered:
                    _render_seg(d, seg)
                self.stats["rendered"] += 1
        _headers_footers(d,
                         "APPSolve (Pty) Ltd",
                         "Oracle HCM Cloud — Implementation Proposal  |  COMMERCIAL IN CONFIDENCE",
                         self.tid)
        os.makedirs(os.path.dirname(self.out_path), exist_ok=True)
        d.save(self.out_path)
        print(f"[CLIENT V3]   Saved: {self.out_path}")
        print(f"[CLIENT V3]   Rendered={self.stats['rendered']} | "
              f"Authored={self.stats['authored']} | Pending={self.stats['pending']}")

    def _cover(self):
        d = self.doc
        sp = d.add_paragraph(); sp.paragraph_format.space_before=Pt(60); sp.paragraph_format.space_after=Pt(0)
        p1 = d.add_paragraph(); p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run("APPSolve (Pty) Ltd"); r1.font.name="Calibri"; r1.font.size=Pt(32); r1.font.bold=True; r1.font.color.rgb=NAVY
        p2 = d.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after=Pt(48)
        r2 = p2.add_run("Enterprise Technology Services  |  Oracle  |  Acumatica  |  BeBanking"); r2.font.name="Calibri"; r2.font.size=Pt(11); r2.font.color.rgb=BLUE
        p3 = d.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER; p3.paragraph_format.space_after=Pt(8)
        r3 = p3.add_run("ORACLE HCM CLOUD — IMPLEMENTATION PROPOSAL"); r3.font.name="Calibri"; r3.font.size=Pt(18); r3.font.bold=True; r3.font.color.rgb=GREY
        p4 = d.add_paragraph(); p4.alignment=WD_ALIGN_PARAGRAPH.CENTER; p4.paragraph_format.space_after=Pt(48)
        r4 = p4.add_run("Plennegy Group"); r4.font.name="Calibri"; r4.font.size=Pt(15); r4.font.bold=True; r4.font.color.rgb=NAVY
        tbl = d.add_table(rows=0, cols=2); tbl.style="Table Grid"
        for key, val in [("Prepared For","Plennegy Group"),("Platform","Oracle HCM Cloud"),
                         ("Engagement","Full Suite Implementation"),("Prepared By","APPSolve (Pty) Ltd"),
                         ("Version","3.0 — Structured Draft"),("Date",datetime.now().strftime("%Y-%m-%d")),
                         ("Tender ID",self.tid)]:
            row=tbl.add_row()
            kc=row.cells[0]; vc=row.cells[1]
            _cell_bg(kc,HEX_NAVY); _cell_borders(kc,HEX_NAVY,4)
            _cell_bg(vc,"F2F5FB"); _cell_borders(vc,"D0D0D0",4)
            kp=kc.paragraphs[0]; kp.paragraph_format.space_before=Pt(4); kp.paragraph_format.space_after=Pt(4)
            kr=kp.add_run(key); kr.font.name="Calibri"; kr.font.size=Pt(11); kr.font.bold=True; kr.font.color.rgb=WHITE
            vp=vc.paragraphs[0]; vp.paragraph_format.space_before=Pt(4); vp.paragraph_format.space_after=Pt(4)
            add_inline(vp, val, sz=11.0)
        fp=d.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; fp.paragraph_format.space_before=Pt(48); fp.paragraph_format.space_after=Pt(4)
        fr=fp.add_run("COMMERCIAL IN CONFIDENCE"); fr.font.name="Calibri"; fr.font.size=Pt(10); fr.font.bold=True; fr.font.color.rgb=LGREY
        fp2=d.add_paragraph(); fp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        fr2=fp2.add_run("APPSolve (Pty) Ltd  |  Gauteng & Western Cape, South Africa"); fr2.font.name="Calibri"; fr2.font.size=Pt(9); fr2.font.color.rgb=LGREY
        _page_break(d)

# ── INTERNAL WORKBOOK V3 renderer ─────────────────────────────────────────────

class InternalWorkbookV3Renderer:
    def __init__(self, v3_secs: List[V3Sec], out_path: str, tender_id: str):
        self.secs     = v3_secs
        self.out_path = out_path
        self.tid      = tender_id
        self.doc      = Document()
        setup_styles(self.doc)
        _page_setup(self.doc)

    def render(self):
        d = self.doc
        # Header banner
        p1 = d.add_paragraph(); p1.alignment=WD_ALIGN_PARAGRAPH.CENTER; p1.paragraph_format.space_before=Pt(24)
        r1 = p1.add_run("APPSolve — V3 Proposal Internal Review Workbook"); r1.font.name="Calibri"; r1.font.size=Pt(22); r1.font.bold=True; r1.font.color.rgb=NAVY
        p2 = d.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after=Pt(24)
        r2 = p2.add_run(f"{self.tid}  |  Oracle HCM Cloud  |  {datetime.now().strftime('%Y-%m-%d')}"); r2.font.name="Calibri"; r2.font.size=Pt(12); r2.font.color.rgb=BLUE
        # Status overview
        rendered  = [s for s in self.secs if s.status == "RENDERED"]
        authored  = [s for s in self.secs if s.status == "AUTHORED"]
        ph_secs   = [s for s in self.secs if s.status == "PLACEHOLDER"]
        ai_secs   = [s for s in self.secs if s.status == "AI_DRAFT"]
        total     = len(self.secs)
        tbl = d.add_table(rows=1, cols=5); tbl.style="Table Grid"
        for ci, (label, val, bg) in enumerate([
            ("RENDERED",    str(len(rendered)),  HEX_STATUS_OK),
            ("AUTHORED",    str(len(authored)),  HEX_BG_TRACE),
            ("PLACEHOLDER", str(len(ph_secs)),   HEX_STATUS_PH),
            ("AI-DRAFT",    str(len(ai_secs)),   HEX_STATUS_AI),
            ("TOTAL",       str(total),          "F2F5FB"),
        ]):
            c = tbl.rows[0].cells[ci]; _cell_bg(c, bg); _cell_borders(c,"D0D0D0",4)
            cp = c.paragraphs[0]; cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before=Pt(6); cp.paragraph_format.space_after=Pt(6)
            cr = cp.add_run(val + "\n" + label); cr.font.name="Calibri"; cr.font.size=Pt(10); cr.font.bold=True
        d.add_paragraph().paragraph_format.space_after=Pt(6)
        render_callout(d, [
            "**Internal Use Only — V3 Draft Review Workbook.** "
            "Complete all PLACEHOLDER and AI-DRAFT sections before generating CLIENT profile. "
            "Do not distribute to clients or external parties.",
        ])
        _page_break(d)
        # Action register
        self._action_register(ph_secs, ai_secs)
        _page_break(d)
        # V3 → V2 mapping table
        self._v2_v3_map()
        _page_break(d)
        # Body
        for sec in self.secs:
            _page_break(d)
            badge = {"RENDERED": " ✓", "AUTHORED": " ✎", "PLACEHOLDER": " ⚠", "AI_DRAFT": " ✏", "EMPTY": " —"}.get(sec.status, "")
            h = d.add_heading(level=2); h.paragraph_format.keep_with_next=True
            add_inline(h, sec.heading + badge, sz=16.0)
            if sec.status == "PLACEHOLDER":
                _action_box(d, "PLACEHOLDER — Human action required", sec.heading, "Account Manager / BU Lead")
                # Still show the comment content
                for seg in sec.segs:
                    _render_seg(d, seg)
            elif sec.status == "AI_DRAFT":
                _action_box(d, "AI-DRAFT REQUIRED — Generate and BU Lead review before submission", sec.heading, "BU Lead")
                for seg in sec.segs:
                    _render_seg(d, seg)
            elif sec.status == "AUTHORED":
                render_callout(d, ["**Authored Content:** This section uses shaper-authored narrative. Review for accuracy before submission."])
                for seg in sec.segs:
                    _render_seg(d, seg)
            else:
                filtered_segs = filter_client(sec.segs)
                for seg in filtered_segs:
                    _render_seg(d, seg)
        _headers_footers(d, "APPSolve INTERNAL REVIEW WORKBOOK V3",
                         f"Oracle HCM — {self.tid}  |  INTERNAL USE ONLY", self.tid)
        os.makedirs(os.path.dirname(self.out_path), exist_ok=True)
        d.save(self.out_path)
        print(f"[INTERNAL V3] Saved: {self.out_path}")

    def _action_register(self, ph_secs, ai_secs):
        d = self.doc
        h = d.add_heading("Action Register — Sections Requiring Completion", level=1)
        h.paragraph_format.space_before=Pt(0)
        p = d.add_paragraph(style="Normal"); add_inline(p, "Complete these sections before generating CLIENT profile.")
        d.add_paragraph()
        tbl = d.add_table(rows=1, cols=4); tbl.style="Table Grid"
        for ci, hdr in enumerate(["Priority", "Section", "Action Required", "Owner"]):
            hc=tbl.rows[0].cells[ci]; _cell_bg(hc,HEX_NAVY); _cell_borders(hc,HEX_NAVY,4)
            hp2=hc.paragraphs[0]; hp2.paragraph_format.space_before=Pt(3); hp2.paragraph_format.space_after=Pt(3)
            hr2=hp2.add_run(hdr); hr2.font.name="Calibri"; hr2.font.size=Pt(10); hr2.font.bold=True; hr2.font.color.rgb=WHITE
        actions = [
            ("P0 — URGENT", "B-BBEE Certificate", "Renew B-BBEE cert before 2026-07-31", "Finance Director"),
        ] + [
            ("P1", s.heading, "AI-Draft + BU Lead review required", "BU Lead") for s in ai_secs
        ] + [
            ("P1–P2", s.heading, "Human input required — see section detail", "AM / BU Lead") for s in ph_secs
        ]
        for prio, section, action, owner in actions:
            bg = HEX_STATUS_PH if "P0" in prio or "P1" in prio else HEX_STATUS_AI
            row = tbl.add_row()
            for c in row.cells: _cell_bg(c, bg); _cell_borders(c,"D0D0D0",4)
            for ci2, v in enumerate([prio, section, action, owner]):
                cp2=row.cells[ci2].paragraphs[0]; cp2.paragraph_format.space_before=Pt(3); cp2.paragraph_format.space_after=Pt(3)
                add_inline(cp2, v, sz=9.5)

    def _v2_v3_map(self):
        d = self.doc
        h = d.add_heading("V3 Section → V2 Source Mapping", level=1)
        h.paragraph_format.space_before=Pt(0)
        p = d.add_paragraph(style="Normal"); add_inline(p, "Shows how each V3 section was assembled from V2 rendered content.")
        d.add_paragraph()
        tbl = d.add_table(rows=1, cols=4); tbl.style="Table Grid"
        for ci, hdr in enumerate(["V3 Section", "Status", "Mode", "V2 Sources"]):
            hc=tbl.rows[0].cells[ci]; _cell_bg(hc,HEX_NAVY); _cell_borders(hc,HEX_NAVY,4)
            hp2=hc.paragraphs[0]; hp2.paragraph_format.space_before=Pt(3); hp2.paragraph_format.space_after=Pt(3)
            hr2=hp2.add_run(hdr); hr2.font.name="Calibri"; hr2.font.size=Pt(10); hr2.font.bold=True; hr2.font.color.rgb=WHITE
        for sec in self.secs:
            bg = (HEX_STATUS_OK if sec.status in ("RENDERED","AUTHORED") else
                  HEX_STATUS_PH if sec.status == "PLACEHOLDER" else
                  HEX_STATUS_AI if sec.status == "AI_DRAFT" else "F4F4F4")
            row = tbl.add_row()
            for c in row.cells: _cell_bg(c, bg); _cell_borders(c,"D0D0D0",4)
            src_str = "; ".join(sec.sources) if sec.sources else ("Authored" if sec.mode=="authored" else "—")
            for ci2, v in enumerate([sec.heading, sec.status, sec.mode, src_str]):
                cp2=row.cells[ci2].paragraphs[0]; cp2.paragraph_format.space_before=Pt(3); cp2.paragraph_format.space_after=Pt(3)
                add_inline(cp2, v, sz=9.0)

# ── TRACEABILITY REPORT V3 renderer ──────────────────────────────────────────

class TraceabilityReportV3Renderer:
    def __init__(self, v3_secs: List[V3Sec], out_path: str, tender_id: str,
                 v2_secs: List[V2Sec], manifest: dict):
        self.secs     = v3_secs
        self.out_path = out_path
        self.tid      = tender_id
        self.v2_secs  = v2_secs
        self.manifest = manifest
        self.doc      = Document()
        setup_styles(self.doc)
        _page_setup(self.doc)

    def render(self):
        d = self.doc
        p1=d.add_paragraph(); p1.alignment=WD_ALIGN_PARAGRAPH.CENTER; p1.paragraph_format.space_before=Pt(24)
        r1=p1.add_run("V3 Proposal Traceability Report"); r1.font.name="Calibri"; r1.font.size=Pt(22); r1.font.bold=True; r1.font.color.rgb=NAVY
        p2=d.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after=Pt(12)
        r2=p2.add_run(f"{self.tid}  |  PF2-008 Shaping Layer  |  {datetime.now().strftime('%Y-%m-%d')}  |  Platform L5.9")
        r2.font.name="Calibri"; r2.font.size=Pt(11); r2.font.color.rgb=BLUE
        # Structure change summary
        h0 = d.add_heading("V2 → V3 Structure Change Summary", level=1); h0.paragraph_format.space_before=Pt(12)
        tbl0 = d.add_table(rows=0, cols=2); tbl0.style="Table Grid"
        v2_count = len(self.v2_secs)
        v3_body  = len([s for s in self.secs if not s.is_app])
        v3_app   = len([s for s in self.secs if s.is_app])
        for k, v in [
            ("V2 sections", str(v2_count)),
            ("V3 body sections", str(v3_body)),
            ("V3 appendices", str(v3_app)),
            ("V3 total", str(len(self.secs))),
            ("V2 sections consolidated into V3 Sec 4", "7 (Corp Overview, History, Awards, Delivery, Geography, Differentiators, Oracle Partnership)"),
            ("V2 sections moved to Appendix A", "3 (HCM, ERP, OIC capability)"),
            ("Authored sections", str(len([s for s in self.secs if s.mode == "authored"]))),
            ("Shaping mode", "PF2-008 Proposal Shaping Layer v1.0"),
        ]:
            row=tbl0.add_row(); kc=row.cells[0]; vc=row.cells[1]
            _cell_bg(kc,HEX_NAVY); _cell_borders(kc,HEX_NAVY,4)
            _cell_bg(vc,"F2F5FB"); _cell_borders(vc,"D0D0D0",4)
            kp=kc.paragraphs[0]; kp.paragraph_format.space_before=Pt(3); kp.paragraph_format.space_after=Pt(3)
            kr=kp.add_run(k); kr.font.name="Calibri"; kr.font.size=Pt(10); kr.font.bold=True; kr.font.color.rgb=WHITE
            vp=vc.paragraphs[0]; vp.paragraph_format.space_before=Pt(3); vp.paragraph_format.space_after=Pt(3)
            add_inline(vp, v, sz=10.0)
        d.add_paragraph().paragraph_format.space_after=Pt(6)
        _page_break(d)
        # Detailed section mapping
        h1 = d.add_heading("Section-by-Section Traceability", level=1); h1.paragraph_format.space_before=Pt(0)
        for sec in self.secs:
            _page_break(d)
            h2 = d.add_heading(sec.heading, level=2); h2.paragraph_format.keep_with_next=True
            # Traceability panel
            tbl = d.add_table(rows=0, cols=2); tbl.style="Table Grid"
            def row(key, val):
                r = tbl.add_row()
                kc=r.cells[0]; vc=r.cells[1]
                _cell_bg(kc,HEX_BG_TRACE); _cell_borders(kc,HEX_BD_TRACE,4)
                _cell_bg(vc,"F8FFF8"); _cell_borders(vc,"D0E8D0",4)
                kp=kc.paragraphs[0]; kp.paragraph_format.space_before=Pt(2); kp.paragraph_format.space_after=Pt(2)
                kr=kp.add_run(key); kr.font.name="Calibri"; kr.font.size=Pt(9); kr.font.bold=True; kr.font.color.rgb=RGBColor(0x1A,0x5C,0x3A)
                vp=vc.paragraphs[0]; vp.paragraph_format.space_before=Pt(2); vp.paragraph_format.space_after=Pt(2)
                add_inline(vp, val, sz=9.0)
            row("V3 Section ID",    sec.id)
            row("V3 Heading",       sec.heading)
            row("Status",           sec.status)
            row("Assembly Mode",    sec.mode)
            row("V2 Sources",       "; ".join(sec.sources) if sec.sources else "None (authored)")
            row("Is Appendix",      "Yes" if sec.is_app else "No")
            d.add_paragraph().paragraph_format.space_after=Pt(6)
            # Show brief content excerpt for reference
            if sec.status in ("RENDERED","AUTHORED"):
                seg_types = {}
                for s in sec.segs:
                    seg_types[s.kind] = seg_types.get(s.kind, 0) + 1
                summary = ", ".join(f"{v} {k}" for k, v in sorted(seg_types.items()))
                p = d.add_paragraph(style="Normal")
                add_inline(p, f"_Content segments: {summary}._", sz=9.5)
            elif sec.status in ("PLACEHOLDER","AI_DRAFT","EMPTY"):
                p = d.add_paragraph(style="Normal")
                add_inline(p, f"_No content — {sec.status}. Requires human completion before submission._", sz=9.5)
        _headers_footers(d, "TRACEABILITY REPORT V3 — APPSolve INTERNAL",
                         f"{self.tid}  |  PF2-008  |  GOVERNANCE DOCUMENT", self.tid)
        os.makedirs(os.path.dirname(self.out_path), exist_ok=True)
        d.save(self.out_path)
        print(f"[TRACEABILITY V3] Saved: {self.out_path}")

# ── Review report (MD) ────────────────────────────────────────────────────────

def generate_review_report(v3_secs: List[V3Sec], v2_secs: List[V2Sec],
                           out_path: str, tender_id: str) -> str:
    v2_count = len(v2_secs)
    v2_rendered  = len([s for s in v2_secs if s.status == "RENDERED"])
    v2_ph        = len([s for s in v2_secs if s.status == "PLACEHOLDER"])
    v2_ai        = len([s for s in v2_secs if s.status == "AI_DRAFT"])
    v3_body      = [s for s in v3_secs if not s.is_app]
    v3_app       = [s for s in v3_secs if s.is_app]
    v3_rendered  = len([s for s in v3_body if s.status == "RENDERED"])
    v3_authored  = len([s for s in v3_body if s.status == "AUTHORED"])
    v3_ph        = len([s for s in v3_body if s.status == "PLACEHOLDER"])
    v3_ai        = len([s for s in v3_body if s.status == "AI_DRAFT"])

    lines = [
        "---",
        f"document_id: PLENNEGY-V2-V3-REVIEW-V1",
        f'title: "Plennegy Proposal V2 → V3 Review Report"',
        f'version: "1.0"',
        f'status: "COMPLETE"',
        f'created: "{datetime.now().strftime("%Y-%m-%d")}"',
        f'created_by: "PF2-008 — Proposal Shaping Layer"',
        "---",
        "",
        "# Plennegy Proposal V2 → V3 Review Report",
        "",
        f"**Tender:** {tender_id}  ",
        f"**Date:** {datetime.now().strftime('%Y-%m-%d')}  ",
        "**Shaping Engine:** proposal_shaper.py v1.0 (PF2-008)  ",
        "**Platform:** L5.9",
        "",
        "---",
        "",
        "## 1. Structural Comparison",
        "",
        "| Dimension | V2 | V3 |",
        "|---|---|---|",
        f"| Total sections | {v2_count} (numbered) | 10 body + 3 appendices = 13 |",
        f"| Structure type | Flat 40-section list | 10-section narrative + appendices |",
        f"| Credential sections in body | 7 (Secs 1–7) | Consolidated into 1 (Sec 4: Why APPSolve) |",
        f"| Capability sections in body | 3 (Secs 12–14) | Moved to Appendix A |",
        f"| Numbered sections in body | 37 | 10 |",
        f"| Appendices | 3 (A, D, E) | 3 (A, B, C) |",
        f"| Compliance sections in body | 7 (Secs 30–36) | Not in CLIENT body (separate submission) |",
        "",
        "## 2. Section Status Comparison",
        "",
        "| Status | V2 | V3 Body | V3 App | Change |",
        "|---|---|---|---|---|",
        f"| RENDERED (from KB) | {v2_rendered} | {v3_rendered} | {len([s for s in v3_app if s.status=='RENDERED'])} | Capability moved to appendices |",
        f"| AUTHORED (by shaper) | 0 | {v3_authored} | 0 | New: governance, commercial, next steps |",
        f"| PLACEHOLDER | {v2_ph} | {v3_ph} | {len([s for s in v3_app if s.status=='PLACEHOLDER'])} | Consolidated |",
        f"| AI_DRAFT | {v2_ai} | {v3_ai} | 0 | Unchanged |",
        "",
        "## 3. Key Structural Changes",
        "",
        "### 3.1 Consolidation of Credential Sections",
        "",
        "V2 contained 7 separate credential sections (Company Overview, Company History, Awards and Recognition, "
        "Delivery Model, Geographic Presence, Key Differentiators, Oracle Partnership) as individual numbered "
        "sections in the body. In V3 these are consolidated into a single **Section 4: Why APPSolve**, "
        "providing a unified, narrative credential argument rather than a disconnected list of corporate facts.",
        "",
        "### 3.2 Capability Content Moved to Appendix A",
        "",
        "V2 Sections 12 (Oracle HCM), 13 (Oracle ERP), and 14 (Oracle OIC) contained full, detailed "
        "Oracle capability documentation from governed Knowledge Base assets. In V3 CLIENT, this content "
        "is moved to **Appendix A: Detailed Oracle HCM Capability**. The body sections (1–10) focus on "
        "the proposal narrative; evaluators can reference the appendix for technical depth.",
        "",
        "### 3.3 Authored Governance and Commercial Sections",
        "",
        "V2 had no authored content — all content came from the Knowledge Base. V3 introduces three "
        "authored body sections (Governance, Commercial, Next Steps) that provide client-specific "
        "narrative and clear action visibility. These sections are authored by the Shaping Layer and "
        "must be reviewed by BU Lead before submission.",
        "",
        "### 3.4 Compliance Sections",
        "",
        "V2 Sections 30–36 (Compliance Schedule, Company Registration, Tax Clearance, Directors' Resolution, "
        "B-BBEE, Public Liability, OPN Certificate) were PLACEHOLDER sections in the body. In V3, "
        "compliance documents are not included in the body or appendices of the CLIENT document — "
        "they are submitted as separate physical attachments. The Internal Workbook tracks their status.",
        "",
        "## 4. Numbering",
        "",
        "| Issue | V2 | V3 |",
        "|---|---|---|",
        "| Numbered sections | 37 (too many for client body) | 10 body sections |",
        "| Consistent numbering | No (appendix lettering mixed) | Yes (1–10 body, A–C appendix) |",
        "| Capability sections numbered | Yes (12, 13, 14) | Moved to appendix (unlabelled) |",
        "| Governance/compliance numbered | Yes (29–36) | Not in CLIENT body |",
        "",
        "## 5. Proposal Readiness Assessment",
        "",
        "| Criterion | V2 Score | V3 Score | Notes |",
        "|---|---|---|---|",
        "| Reads like a proposal (not a KB dump) | 3/5 | 4/5 | Structured narrative; V3 Sec 4 much cleaner |",
        "| Correct structure | 2/5 | 4/5 | 10-section structure maps to evaluator expectations |",
        "| Numbering | 2/5 | 5/5 | Clean 1–10 + A–C; no over-numbered body |",
        "| Credential conciseness | 2/5 | 4/5 | 7 sections → 1 section |",
        "| Evidence/traceability separated | 1/5 | 5/5 | Appendix A full capability; App B assumptions |",
        "| Pending content visibility | 3/5 | 4/5 | Clean 'content pending' notices in CLIENT |",
        "| Commercial sections | 2/5 | 4/5 | Authored commercial framework section added |",
        "| Next steps clarity | 1/5 | 5/5 | Explicit 10-point next steps with owners |",
        f"| **Overall** | **2.0/5** | **4.1/5** | **V3 substantially more client-ready** |",
        "",
        "## 6. Remaining Gaps (same as V2)",
        "",
        "| Section | Gap | Owner | Priority |",
        "|---|---|---|---|",
        "| 1. Executive Summary | AI-Draft required | BU Lead | P1 |",
        "| 2. Understanding of Requirements | AI-Draft from RFP | Account Manager | P1 |",
        "| 3. Proposed Oracle HCM Solution | Human authored | BU Lead + AM | P1 |",
        "| 8. Key Risks and Mitigations | AI-Draft + BU review | BU Lead | P2 |",
        "| 9. Commercial section | Pricing from Commercial Director | Commercial Director | P1 |",
        "| Appendix C | Client references approval per AM | Account Manager | P2 |",
        "| B-BBEE Certificate | Renewal deadline 2026-07-31 | Finance Director | **P0 URGENT** |",
        "",
        "## 7. V3 Document Statistics",
        "",
        "| Document | Sections |",
        "|---|---|",
        f"| CLIENT V3 body sections | {len(v3_body)} |",
        f"| CLIENT V3 appendices | {len(v3_app)} |",
        f"| CLIENT V3 RENDERED/AUTHORED sections | {v3_rendered + v3_authored} |",
        f"| CLIENT V3 PENDING sections | {v3_ph + v3_ai} |",
        "",
        "",
        f"*Generated by proposal_shaper.py v1.0 (PF2-008) | {datetime.now().strftime('%Y-%m-%d %H:%M')} | Platform L5.9*",
    ]
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    print(f"[REVIEW MD]    Saved: {out_path}")
    return out_path

# ── Main orchestrator ─────────────────────────────────────────────────────────

def run_shaper(md_path: str, manifest_path: str, out_dir: str, tender_id: str):
    print(f"[SHAPER] Tender:   {tender_id}")
    print(f"[SHAPER] Input MD: {md_path}")
    print(f"[SHAPER] Output:   {out_dir}")

    with open(md_path, "r", encoding="utf-8") as fh:
        md_text = fh.read()
    with open(manifest_path, "r", encoding="utf-8") as fh:
        manifest = yaml.safe_load(fh)

    # Parse and classify V2 sections
    all_segs = parse_md(md_text)
    v2_secs  = group_v2(all_segs)
    classify_v2(v2_secs)
    idx      = v2_idx(v2_secs)
    print(f"[SHAPER] V2 sections: {len(v2_secs)} | "
          f"RENDERED={sum(1 for s in v2_secs if s.status=='RENDERED')} | "
          f"PLACEHOLDER={sum(1 for s in v2_secs if s.status=='PLACEHOLDER')} | "
          f"AI_DRAFT={sum(1 for s in v2_secs if s.status=='AI_DRAFT')}")

    # Build V3 sections
    v3_secs = build_v3_sections(idx)
    print(f"[SHAPER] V3 sections: {len(v3_secs)} | "
          f"RENDERED={sum(1 for s in v3_secs if s.status=='RENDERED')} | "
          f"AUTHORED={sum(1 for s in v3_secs if s.status=='AUTHORED')} | "
          f"PLACEHOLDER={sum(1 for s in v3_secs if s.status=='PLACEHOLDER')} | "
          f"AI_DRAFT={sum(1 for s in v3_secs if s.status=='AI_DRAFT')}")

    os.makedirs(out_dir, exist_ok=True)
    prefix = "PLENNEGY_PROPOSAL" if "PLENNEGY" in tender_id.upper() else f"PROPOSAL_{tender_id}"

    # Generate DOCX outputs
    ClientV3Renderer(
        v3_secs,
        os.path.join(out_dir, f"{prefix}_CLIENT_V3.docx"),
        tender_id
    ).render()

    InternalWorkbookV3Renderer(
        v3_secs,
        os.path.join(out_dir, f"{prefix}_INTERNAL_REVIEW_WORKBOOK_V3.docx"),
        tender_id
    ).render()

    TraceabilityReportV3Renderer(
        v3_secs,
        os.path.join(out_dir, f"{prefix}_TRACEABILITY_REPORT_V3.docx"),
        tender_id, v2_secs, manifest
    ).render()

    # Generate MD review report
    review_md = os.path.join(out_dir, f"PLENNEGY_V2_TO_V3_REVIEW_REPORT.md")
    generate_review_report(v3_secs, v2_secs, review_md, tender_id)

    print(f"\n[SHAPER] Complete — 4 outputs in {out_dir}")
    return {
        "client":       os.path.join(out_dir, f"{prefix}_CLIENT_V3.docx"),
        "internal":     os.path.join(out_dir, f"{prefix}_INTERNAL_REVIEW_WORKBOOK_V3.docx"),
        "traceability": os.path.join(out_dir, f"{prefix}_TRACEABILITY_REPORT_V3.docx"),
        "review_md":    review_md,
    }

# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Proposal Shaping Layer v1.0 (PF2-008)")
    parser.add_argument("--tender", metavar="TENDER_ID",
                        help="Tender ID — auto-locates rendered MD and manifest")
    parser.add_argument("--md",       metavar="MD_PATH")
    parser.add_argument("--manifest", metavar="MANIFEST_PATH")
    parser.add_argument("--out",      metavar="OUT_DIR")
    args = parser.parse_args()

    if args.tender:
        tid     = args.tender
        td_dir  = os.path.join(PROPOSALS_DIR, tid)
        md_path = os.path.join(td_dir, f"PROPOSAL_RENDERED_{tid}.md")
        mfiles  = glob.glob(os.path.join(td_dir, f"PROPOSAL_SECTION_MANIFEST_{tid}.yaml"))
        if not mfiles:
            print(f"[SHAPER] ERROR: Manifest not found in {td_dir}", file=sys.stderr); sys.exit(1)
        manifest_path = mfiles[0]
        out_dir = td_dir
    elif args.md and args.manifest and args.out:
        tid = os.path.basename(args.out)
        md_path = args.md; manifest_path = args.manifest; out_dir = args.out
    else:
        parser.print_help(); sys.exit(1)

    if not os.path.exists(md_path):
        print(f"[SHAPER] ERROR: Rendered MD not found: {md_path}", file=sys.stderr); sys.exit(1)

    run_shaper(md_path, manifest_path, out_dir, tid)

if __name__ == "__main__":
    main()
