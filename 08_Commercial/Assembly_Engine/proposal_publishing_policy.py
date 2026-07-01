#!/usr/bin/env python3
"""
proposal_publishing_policy.py — Proposal Publishing Policy Engine v1.0 (PPPE / PF2-010)

PPPE sits between the Renderer and the DOCX output layer. It owns every decision
about what a customer is allowed to see. It does NOT modify Knowledge Assets,
KRPE, KVE, or PSAE.

10 Publishing Rules (see PROPOSAL_PUBLISHING_STANDARD.md):
  R1  — Remove internal governance information
  R2  — Strip all internal identifiers (HCM-ENV-001, CAP-..., BU-WP..., etc.)
  R3  — Remove implementation annotations (Updated —, Revised —, etc.)
  R4  — Remove BU Decision tables
  R5  — Replace assumption libraries with executive summaries (3–8 bullets per pack)
  R6  — Replace OUM with Oracle's Modern Best Practices (Fusion Cloud proposals)
  R7  — Condense roadmaps to one table
  R8  — Customer-oriented next steps only
  R9  — Remove source evidence references
  R10 — Enforce page budget (target 25–35 pages)

Three publication profiles:
  PROFILE_CUSTOMER — clean customer submission (25–35 pages)
  PROFILE_INTERNAL — everything, no policy applied
  PROFILE_REVIEW   — customer content + policy annotations showing removals

Usage:
    python3 proposal_publishing_policy.py --tender PLENNEGY-HCM-001
"""

import argparse, glob, os, re, sys
from dataclasses import dataclass, field
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)
from docx_renderer import (
    Seg, parse_md, add_inline, render_table, render_callout,
    render_ulist, render_olist, setup_styles,
    _page_break, _hr, _toc_field, _cell_bg, _cell_borders,
    NAVY, BLUE, GREY, LGREY, WHITE, ORANGE,
    HEX_NAVY, HEX_BLUE, HEX_WHITE, HEX_BG_PH, HEX_BG_AI,
    classify_comment, PROPOSALS_DIR
)
from proposal_shaper import (
    group_v2, classify_v2, _page_setup, _headers_footers,
    HEX_BG_PENDING, HEX_BD_PENDING, LGREY2,
)
from pee import apply_budget

# ── Profile constants ─────────────────────────────────────────────────────────
PROFILE_CUSTOMER = "CUSTOMER"
PROFILE_INTERNAL = "INTERNAL"
PROFILE_REVIEW   = "REVIEW"

# ── Rule 1 — governance subsection headings (H3/H4) to exclude ───────────────
_GOVERNANCE_SUB_RE = re.compile(
    r'\b(?:extraction\s+notes?|fact\s+verif|governance\s+review|'
    r'review\s+notes?|approval\s+history|approval\s+record|'
    r'approval\s+status|document\s+register|author\s+note|'
    r'bu\s+lead\s+decision|bu\s+decision|update\s+history|'
    r'source\s+evidence|primary\s+source|traceability\s+note)\b',
    re.IGNORECASE
)

# Top-level V2 sections that are ALWAYS internal-only (match group_v2-stripped names, lowercased)
_INTERNAL_SECTION_NAMES = {
    "table of contents",
    "extraction notes",
    "fact verification summary",
    "pre-tender checks",
    "pre tender checks",
    "approval record",
    "governance review",
    "internal review",
    # Duplicate assumption sections — only the appendix copy is published to customer
    "commercial assumptions",
}

# ── Rule 2 — internal identifier pattern ─────────────────────────────────────
_ID_RE = re.compile(
    r'\b(?:HCM|ERP|OIC|RCR|LRN|TLN|CAP|ASM|ASP|SEC|PAT|REF|RSK|BU|KB|RC|SI|GOV|ADR|IG|AV|RI|TD|WP\d+)'
    r'[-–][A-Z0-9]+(?:[-–][A-Z0-9]+)*\b'
)

# ── Rule 3 — implementation annotation lines ──────────────────────────────────
_ANNOTATION_LINE_RE = re.compile(
    r'^\*?\(?(?:Updated?|Revised?|Added?|Removed?|Applied?)\s*[—–-]\s*(?:BU-|WP|KB|RC)[^\)]*\)?\*?$',
    re.IGNORECASE
)
_ANNOTATION_INLINE_RE = re.compile(
    r'\s*\*?\(?(?:Updated?|Revised?)\s*[—–-]\s*(?:BU-\w+-\d+|WP\d+\w*[-–]\d+)[^\)]*\)?\*?',
    re.IGNORECASE
)

# ── Rule 4 — BU decision table markers ───────────────────────────────────────
_BU_DECISION_RE = re.compile(r'\bBU\s+(?:Lead\s+)?Decision|BU-WP\d+', re.IGNORECASE)

# ── Rule 6 — OUM replacement (Fusion Cloud only) ─────────────────────────────
_OUM_RE = re.compile(
    r'Oracle\s+Unified\s+Method(?:ology)?(?:\s*\(OUM\))?|'
    r'\bOUM\s+(?:methodology|framework|approach|phase|structured|standard)|'
    r'the\s+OUM\b',
    re.IGNORECASE
)
_OUM_BARE_RE = re.compile(r'\bOUM\b', re.IGNORECASE)
OUM_REPLACE = "Oracle's Modern Best Practices"

# ── Rule 9 — evidence column headers ─────────────────────────────────────────
_EVIDENCE_COLS = {"primary source", "evidence", "source document", "knowledge source",
                  "capability source", "traceability", "internal reference",
                  "fact reference", "source reference"}

# ── Section ordering for customer proposal (use group_v2-stripped names, lowercased) ───
_CUSTOMER_ORDER = [
    "executive summary",
    "understanding of requirements",
    "proposed solution overview",
    "company overview",
    "company history",
    "awards and recognition",
    "delivery model",
    "geographic presence",
    "key differentiators",
    "oracle partnership",
    "b-bbee compliance statement",
    "implementation methodology",
    "oracle fusion hcm capability",
    "oracle fusion erp capability",
    "oracle oic / integration capability",
    "project governance",
    "project plan / timeline",
    "raid framework",
    "change control framework",
    "team structure",
    "key assumptions (body section)",   # summarized — body overview
    "risk register",
    "commercials / pricing",
    "client references",
    "oracle opn certificate",
    "complete assumption schedule",      # summarized — appendix detail
    "certifications and compliance",
    "company registration",
    "tax clearance",
    "directors' resolution",
    "b-bbee certificate",
    "public liability insurance",
    "oracle opn certificate",
]

# Sections that contain assumption libraries (replace with executive summaries — Rule 5)
_ASSUMPTION_SECTION_NAMES = {
    "key assumptions (body section)",
    "complete assumption schedule",
}

# Section budgets (content segment caps — Rule 10)
_SECTION_BUDGETS: Dict[str, int] = {
    "executive summary": 24,
    "understanding of requirements": 20,
    "proposed solution overview": 16,
    "company overview": 16,
    "company history": 12,
    "awards and recognition": 12,
    "delivery model": 16,
    "geographic presence": 12,
    "key differentiators": 16,
    "oracle partnership": 12,
    "implementation methodology": 40,
    "oracle fusion hcm capability": 56,
    "oracle fusion erp capability": 32,
    "oracle oic / integration capability": 32,
    "project governance": 24,
    "project plan / timeline": 16,
    "team structure": 16,
    "risk register": 20,
    "complete assumption schedule": 48,  # executive summaries only
    "key assumptions (body section)": 24,  # executive summaries only
}
_DEFAULT_BUDGET = 24

# ── Assumption pack executive summaries (Rule 5) ──────────────────────────────

def _asm_segs(heading: str, intro: str, bullets: List[str]) -> List[Seg]:
    out = [Seg("heading", level=3, lines=[heading])]
    out.append(Seg("paragraph", lines=[intro]))
    out.append(Seg("ulist", lines=bullets))
    out.append(Seg("paragraph", lines=[
        "_The detailed contractual assumptions will form part of the Statement of Work "
        "agreed between APPSolve and Plennegy._"
    ]))
    return out

_PACK_SUMMARIES: Dict[str, callable] = {}

def _pack_hcm_base() -> List[Seg]:
    return _asm_segs(
        "Oracle Fusion HCM Cloud — Core Platform",
        "This proposal is based on a standard Oracle Fusion HCM Cloud SaaS implementation "
        "scoped to Plennegy's confirmed legal entities, employee headcount, and module list. "
        "The following core assumptions govern all phases of the engagement.",
        [
            "Oracle Corporation provisions and manages all Oracle Fusion HCM Cloud environments; "
            "APPSolve's implementation timeline commences from environment access date",
            "Oracle applies quarterly cloud updates automatically; APPSolve manages update "
            "assessment and adoption as part of implementation and subsequent AMS",
            "All Oracle Fusion HCM configuration uses Oracle's standard tooling and "
            "Low-Code/No-Code extension framework — no custom code development",
            "Employee data migration (extraction, cleansing, validation) is Plennegy's "
            "responsibility; APPSolve provides migration templates and technical guidance",
            "Implementation scope, timeline, and pricing are based on entities and headcount "
            "confirmed at Mobilise; any additions require a formal Change Request",
            "Plennegy provides a nominated project team with availability for design workshops, "
            "testing, and decision-making throughout the project",
        ]
    )

def _pack_recruiting() -> List[Seg]:
    return _asm_segs(
        "Oracle Recruiting Cloud",
        "The Oracle Recruiting Cloud implementation is scoped within the Phase 1 HCM delivery "
        "and assumes the following.",
        [
            "Oracle Recruiting Cloud is licensed as part of Plennegy's Oracle HCM Cloud subscription",
            "Vacancy workflows, approval routes, and offer letter templates are configured "
            "during Scope & Design using Oracle's standard recruiting configuration",
            "External job board posting requires direct configuration between Plennegy and "
            "each job board provider; interface effort is separately scoped",
            "Pre-employment screening integration is not in scope unless separately quoted "
            "and agreed in the Statement of Work",
            "Plennegy nominates key recruiting users to participate in all design workshops "
            "and User Acceptance Testing",
        ]
    )

def _pack_learning() -> List[Seg]:
    return _asm_segs(
        "Oracle Learning Cloud",
        "Oracle Learning Cloud is delivered in Phase 2 of the engagement and assumes the following.",
        [
            "Oracle Learning Cloud is included in the Oracle HCM Cloud subscription tier confirmed by Plennegy",
            "Plennegy provides an approved course catalogue and learning content prior to configuration",
            "SCORM/xAPI content hosting within Oracle Learning is in-scope; "
            "migration from a third-party LMS is separately scoped",
            "Instructor-led training scheduling uses Oracle Learning standard functionality",
            "Learning completion records are maintained in Oracle Learning — not in a parallel system",
            "Plennegy nominates a Learning Administrator to own content and catalogue management post go-live",
        ]
    )

def _pack_talent() -> List[Seg]:
    return _asm_segs(
        "Oracle Talent Management",
        "Oracle Talent Management (Performance, Goals, Succession) is delivered in Phase 2 "
        "and assumes the following.",
        [
            "Plennegy confirms performance review cycle design, rating scales, and goal taxonomy "
            "before configuration begins in Phase 2",
            "Succession planning scope and nomination criteria are confirmed in Scope & Design",
            "Integration to third-party performance or engagement tools is not in scope "
            "unless separately quoted",
            "Plennegy managers and HR business partners participate in UAT for all "
            "talent management workflows",
        ]
    )

def _pack_oic() -> List[Seg]:
    return _asm_segs(
        "Oracle Integration Cloud (OIC) — Integration Services",
        "All integration services use Oracle Integration Cloud (OIC) as the standard "
        "integration layer. The following assumptions apply.",
        [
            "Oracle Integration Cloud is separately licensed; Plennegy confirms OIC licensing "
            "before integration development commences",
            "Each third-party integration point (payroll provider, time and attendance, etc.) "
            "is separately scoped, designed, and priced",
            "Plennegy's payroll provider participates in integration design workshops and "
            "provides API documentation and test environment access",
            "All APIs used for integration are published, versioned, and vendor-supported; "
            "undocumented or legacy API access is out of scope",
            "Integration monitoring, alerting, and error handling are included within each "
            "individually scoped integration item",
            "APPSolve is not responsible for third-party system downtime affecting integration testing",
        ]
    )

def _pack_generic(heading: str) -> List[Seg]:
    return _asm_segs(
        heading.replace(" V1.0", "").replace(" Assumptions", "").strip(),
        "This assumption pack governs the scope and delivery expectations for this component.",
        [
            "Scope is confirmed during the Mobilise phase and documented in the Statement of Work",
            "Any additions to agreed scope are subject to a formal Change Request",
            "Plennegy provides required access, data, and resources as agreed in the project plan",
        ]
    )

def _get_pack_summary(heading: str) -> List[Seg]:
    h = heading.lower()
    if "hcm" in h and ("base" in h or "core" in h or "hcm" == h.strip()):
        return _pack_hcm_base()
    if "recruit" in h:
        return _pack_recruiting()
    if "learn" in h:
        return _pack_learning()
    if "talent" in h:
        return _pack_talent()
    if "integrat" in h or "oic" in h:
        return _pack_oic()
    return _pack_generic(heading)

# ── Policy record ─────────────────────────────────────────────────────────────

@dataclass
class PolicyRecord:
    tender_id: str
    date: str = field(default_factory=lambda: datetime.now().strftime("%Y-%m-%d"))
    sections_excluded: List[str] = field(default_factory=list)
    sections_summarized: List[str] = field(default_factory=list)
    oum_replacements: int = 0
    id_removals: int = 0
    annotation_removals: int = 0
    bu_tables_removed: int = 0
    gov_subs_removed: int = 0
    segs_in: int = 0
    segs_out: int = 0

# ── Text-level transforms ─────────────────────────────────────────────────────

def _transform_text(text: str, rec: PolicyRecord) -> str:
    # Rule 6 — OUM replacement
    def _oum_sub(m):
        rec.oum_replacements += 1
        return OUM_REPLACE
    text = _OUM_RE.sub(_oum_sub, text)
    text = _OUM_BARE_RE.sub(_oum_sub, text)

    # Rule 3 — inline annotations  *(Updated — BU-WP11-002)*
    def _ann_sub(m):
        rec.annotation_removals += 1
        return ""
    text = _ANNOTATION_INLINE_RE.sub(_ann_sub, text)

    # Rule 2 — strip internal IDs inline (replace with nothing, preserve surrounding text)
    def _id_sub(m):
        rec.id_removals += 1
        return ""
    text = _ID_RE.sub(_id_sub, text)
    # Clean up orphaned punctuation from ID removal (e.g. "( )" "- -" "**  **")
    text = re.sub(r'\*\*\s*\*\*', '', text)
    text = re.sub(r'\(\s*\)', '', text)
    text = re.sub(r'\s{2,}', ' ', text).strip()

    return text

def _transform_seg(seg: Seg, rec: PolicyRecord) -> Seg:
    if seg.lines:
        new_lines = [_transform_text(ln, rec) for ln in seg.lines]
        seg = Seg(seg.kind, seg.level, new_lines, seg.meta)
    return seg

# ── Table filtering (Rule 4 — BU decisions; Rule 9 — evidence columns) ───────

def _is_bu_decision_table(lines: List[str]) -> bool:
    if not lines:
        return False
    header = lines[0] if lines else ""
    return bool(_BU_DECISION_RE.search(header))

def _strip_evidence_columns(lines: List[str], rec: PolicyRecord) -> List[str]:
    if len(lines) < 2:
        return lines
    header_row = lines[0]
    cols = [c.strip().strip("*").lower() for c in header_row.split("|") if c.strip()]
    keep_indices = [i for i, c in enumerate(cols) if c not in _EVIDENCE_COLS]
    if len(keep_indices) == len(cols):
        return lines  # nothing to strip
    result = []
    for row in lines:
        cells = row.split("|")
        if len(cells) <= 2:
            result.append(row)
            continue
        new_cells = [""] + [cells[i+1] for i in keep_indices if i+1 < len(cells)] + [""]
        result.append("|".join(new_cells))
    return result

# ── Segment-level filtering pipeline ─────────────────────────────────────────

def _apply_customer_rules(segs: List[Seg], rec: PolicyRecord) -> List[Seg]:
    result: List[Seg] = []
    in_excluded_sub = False   # tracking governance subsection exclusion
    excluded_sub_level = 0

    for seg in segs:
        # Rule 1 — governance comment segs
        if seg.kind == "comment":
            kind = classify_comment(seg.meta)
            if kind == "governance_note":
                rec.gov_subs_removed += 1
                continue
            # Keep placeholder/ai_draft as pending notice

        # Handle subsection exclusion state
        if seg.kind == "heading":
            heading_text = (seg.lines[0] if seg.lines else "")
            if _GOVERNANCE_SUB_RE.search(heading_text):
                # Rule 1 — enter excluded subsection
                in_excluded_sub = True
                excluded_sub_level = seg.level
                rec.gov_subs_removed += 1
                continue
            elif in_excluded_sub and seg.level <= excluded_sub_level:
                # Exit excluded subsection when we reach same or higher level heading
                in_excluded_sub = False

        if in_excluded_sub:
            rec.gov_subs_removed += 1
            continue

        # Rule 3 — standalone annotation lines (paragraph that is only annotation)
        if seg.kind == "paragraph" and seg.lines:
            text = " ".join(seg.lines).strip()
            if _ANNOTATION_LINE_RE.match(text):
                rec.annotation_removals += 1
                continue

        # Rule 4 + text transforms on tables
        if seg.kind == "table":
            if _is_bu_decision_table(seg.lines):
                rec.bu_tables_removed += 1
                continue
            # Rule 9 — strip evidence columns
            new_lines = _strip_evidence_columns(seg.lines, rec)
            seg = Seg(seg.kind, seg.level, new_lines, seg.meta)

        # Rule 2 + 3 + 6 — text transforms
        seg = _transform_seg(seg, rec)

        # Rule 2 — drop paragraphs that become empty after ID stripping
        if seg.kind == "paragraph" and seg.lines:
            cleaned = _ID_RE.sub("", " ".join(seg.lines)).strip(" |-:")
            if not cleaned:
                rec.id_removals += 1
                continue

        result.append(seg)

    return result

# ── Assumption section replacement (Rule 5) ───────────────────────────────────

def _summarize_assumption_section(segs: List[Seg], rec: PolicyRecord) -> List[Seg]:
    # Identify H3 pack headings within this section
    pack_headings = [seg.lines[0] for seg in segs
                     if seg.kind == "heading" and seg.level == 3 and seg.lines]
    rec.sections_summarized.extend(pack_headings)
    result = []
    for ph in pack_headings:
        result.extend(_get_pack_summary(ph))
    if not result:
        result.extend(_get_pack_summary("Assumption Schedule"))
    return result

# ── Core processing ───────────────────────────────────────────────────────────

def _is_assumption_section(name: str) -> bool:
    return name.lower() in _ASSUMPTION_SECTION_NAMES

def _content_count(segs: List[Seg]) -> int:
    return sum(1 for s in segs if s.kind in ("paragraph", "table", "ulist", "olist"))

def process_for_customer(v2_secs, rec: PolicyRecord) -> List:
    """Apply all 10 rules and return ordered list of (heading, segs, status) for customer."""
    # Index V2 sections by normalised name
    idx = {s.name.lower(): s for s in v2_secs}

    ordered = []
    for name in _CUSTOMER_ORDER:
        sec = idx.get(name)
        if sec is None:
            continue
        if sec.name.lower() in _INTERNAL_SECTION_NAMES:
            rec.sections_excluded.append(sec.name)
            continue

        rec.segs_in += len(sec.segs)

        if _is_assumption_section(sec.name):
            # Rule 5 — replace with executive summaries
            segs = _summarize_assumption_section(sec.segs, rec)
        else:
            # Apply customer rules pipeline
            segs = _apply_customer_rules(list(sec.segs), rec)
            # Rule 10 — per-section page budget
            budget = _SECTION_BUDGETS.get(sec.name.lower(), _DEFAULT_BUDGET)
            if _content_count(segs) > budget:
                segs, _, _ = apply_budget(segs, budget, "Internal Review Pack")

        rec.segs_out += len(segs)
        ordered.append((sec.heading, segs, sec.status))

    return ordered

# ── Shared DOCX rendering helpers ─────────────────────────────────────────────

def _render_seg(doc, seg: Seg):
    if seg.kind == "heading":
        h = doc.add_heading(level=min(seg.level, 5))
        h.paragraph_format.keep_with_next = True
        add_inline(h, seg.lines[0] if seg.lines else "",
                   sz={1:22, 2:16, 3:13, 4:11, 5:11}.get(seg.level, 11))
    elif seg.kind == "table":
        render_table(doc, seg.lines)
    elif seg.kind == "ulist":
        render_ulist(doc, seg.lines)
    elif seg.kind == "olist":
        render_olist(doc, seg.lines)
    elif seg.kind == "blockquote":
        render_callout(doc, seg.lines)
    elif seg.kind == "hr":
        from docx_renderer import _hr; _hr(doc)
    elif seg.kind == "paragraph":
        text = " ".join(seg.lines).strip()
        if text:
            p = doc.add_paragraph(style="Normal"); add_inline(p, text)
    elif seg.kind == "comment":
        kind = classify_comment(seg.meta)
        if kind in ("placeholder", "ai_draft"):
            label = "Action Required" if kind == "placeholder" else "AI-Draft Required"
            render_callout(doc, [f"**{label}:** {seg.meta}"])

def _pending_box(doc, label: str):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    c = tbl.cell(0, 0); _cell_bg(c, "F7F7F7"); _cell_borders(c, "AAAAAA", 6)
    p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"[ {label} — content to be authored before submission ]")
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = LGREY2
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def _policy_strip_banner(doc, n_removed: int, rule: str):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    c = tbl.cell(0, 0); _cell_bg(c, "E8F3E8"); _cell_borders(c, "2E8B57", 6)
    p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"[PPPE {rule}] {n_removed} item(s) removed — policy-governed content excluded from customer proposal")
    r.font.name = "Calibri"; r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = RGBColor(0x2E, 0x8B, 0x57)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

# ── CUSTOMER PROPOSAL renderer ────────────────────────────────────────────────

class CustomerProposalRenderer:
    def __init__(self, sections, out_path: str, tender_id: str):
        self.sections = sections
        self.out_path = out_path; self.tid = tender_id
        self.doc = Document(); setup_styles(self.doc); _page_setup(self.doc)

    def render(self):
        d = self.doc
        self._cover()
        h = d.add_heading("Table of Contents", level=1); h.paragraph_format.space_before = Pt(0)
        from docx_renderer import _toc_field; _toc_field(d)
        _page_break(d)
        for heading, segs, status in self.sections:
            _page_break(d)
            d.add_heading(heading, level=2).paragraph_format.keep_with_next = True
            has_ph = any(s.kind == "comment" and classify_comment(s.meta)
                         in ("placeholder", "ai_draft") for s in segs)
            if status == "PLACEHOLDER" and not any(s.kind != "comment" for s in segs):
                label = re.sub(r"^\d+\.\s*", "", heading)
                _pending_box(d, label)
            elif has_ph:
                for seg in segs:
                    _render_seg(d, seg)
            else:
                for seg in segs:
                    _render_seg(d, seg)
        _headers_footers(d,
                         "APPSolve (Pty) Ltd",
                         "Oracle HCM Cloud Proposal  |  COMMERCIAL IN CONFIDENCE",
                         self.tid)
        os.makedirs(os.path.dirname(self.out_path), exist_ok=True)
        d.save(self.out_path)
        sz = os.path.getsize(self.out_path)
        print(f"[PPPE CUSTOMER] {self.out_path.split('/')[-1]}  ({sz//1024} KB)  {len(self.sections)} sections")

    def _cover(self):
        d = self.doc
        d.add_paragraph().paragraph_format.space_before = Pt(60)
        for txt, sz, bold, colour in [
            ("APPSolve (Pty) Ltd", 32, True, NAVY),
            ("Enterprise Technology Services  |  Oracle  |  Acumatica  |  BeBanking", 11, False, BLUE),
        ]:
            p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if "Enterprise" in txt: p.paragraph_format.space_after = Pt(48)
            r = p.add_run(txt); r.font.name = "Calibri"; r.font.size = Pt(sz); r.font.bold = bold; r.font.color.rgb = colour
        for txt, sz in [("ORACLE FUSION HCM CLOUD", 18), ("IMPLEMENTATION PROPOSAL", 14)]:
            p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)
            r = p.add_run(txt); r.font.name = "Calibri"; r.font.size = Pt(sz); r.font.bold = True; r.font.color.rgb = GREY
        pc = d.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER; pc.paragraph_format.space_after = Pt(40)
        rc = pc.add_run("Plennegy Group"); rc.font.name = "Calibri"; rc.font.size = Pt(16); rc.font.bold = True; rc.font.color.rgb = NAVY
        tbl = d.add_table(rows=0, cols=2); tbl.style = "Table Grid"
        for k, v in [("Prepared For", "Plennegy Group"), ("Platform", "Oracle Fusion HCM Cloud"),
                     ("Engagement", "Full Suite HCM Implementation"),
                     ("Prepared By", "APPSolve (Pty) Ltd"),
                     ("Policy Profile", "PROFILE_CUSTOMER (PPPE v1.0)"),
                     ("Date", datetime.now().strftime("%Y-%m-%d")), ("Tender ID", self.tid)]:
            row = tbl.add_row(); kc = row.cells[0]; vc = row.cells[1]
            _cell_bg(kc, HEX_NAVY); _cell_borders(kc, HEX_NAVY, 4)
            _cell_bg(vc, "F2F5FB"); _cell_borders(vc, "D0D0D0", 4)
            kp = kc.paragraphs[0]; kp.paragraph_format.space_before = Pt(4); kp.paragraph_format.space_after = Pt(4)
            kr = kp.add_run(k); kr.font.name = "Calibri"; kr.font.size = Pt(11); kr.font.bold = True; kr.font.color.rgb = WHITE
            vp = vc.paragraphs[0]; vp.paragraph_format.space_before = Pt(4); vp.paragraph_format.space_after = Pt(4)
            add_inline(vp, v, sz=11.0)
        fp = d.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER; fp.paragraph_format.space_before = Pt(48)
        fr = fp.add_run("COMMERCIAL IN CONFIDENCE  |  PPPE v1.0 — Policy-Governed Publication")
        fr.font.name = "Calibri"; fr.font.size = Pt(10); fr.font.bold = True; fr.font.color.rgb = LGREY2
        _page_break(d)

# ── INTERNAL REVIEW renderer ──────────────────────────────────────────────────

class InternalReviewRenderer:
    def __init__(self, v2_secs, rec: PolicyRecord, out_path: str, tender_id: str):
        self.v2_secs = v2_secs; self.rec = rec
        self.out_path = out_path; self.tid = tender_id
        self.doc = Document(); setup_styles(self.doc); _page_setup(self.doc)

    def render(self):
        d = self.doc
        p1 = d.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; p1.paragraph_format.space_before = Pt(24)
        r1 = p1.add_run("PPPE Internal Review Pack")
        r1.font.name = "Calibri"; r1.font.size = Pt(22); r1.font.bold = True; r1.font.color.rgb = NAVY
        p2 = d.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(f"{self.tid}  |  {self.rec.date}  |  PPPE v1.0  |  PROFILE_INTERNAL")
        r2.font.name = "Calibri"; r2.font.size = Pt(11); r2.font.color.rgb = BLUE
        render_callout(d, ["**INTERNAL USE ONLY.** Complete unfiltered rendered content. "
                           "All sections present. Policy rules NOT applied. "
                           "Do NOT distribute to clients."])
        self._policy_summary()
        _page_break(d)
        for sec in self.v2_secs:
            _page_break(d)
            d.add_heading(sec.heading, level=2).paragraph_format.keep_with_next = True
            badge = {"RENDERED": "RENDERED", "PLACEHOLDER": "PLACEHOLDER",
                     "AI_DRAFT": "AI-DRAFT", "EMPTY": "EMPTY"}.get(sec.status, sec.status)
            p_b = d.add_paragraph(style="Normal")
            add_inline(p_b, f"[INTERNAL] Status: {badge} — unfiltered content", sz=9.0)
            for seg in sec.segs:
                _render_seg(d, seg)
        _headers_footers(d, "APPSolve INTERNAL REVIEW",
                         f"{self.tid}  |  PPPE v1.0  |  INTERNAL ONLY", self.tid)
        os.makedirs(os.path.dirname(self.out_path), exist_ok=True)
        d.save(self.out_path)
        sz = os.path.getsize(self.out_path)
        print(f"[PPPE INTERNAL] {self.out_path.split('/')[-1]}  ({sz//1024} KB)  V2 secs={len(self.v2_secs)}")

    def _policy_summary(self):
        d = self.doc
        h = d.add_heading("PPPE Policy Application Summary", level=1); h.paragraph_format.space_before = Pt(0)
        tbl = d.add_table(rows=0, cols=2); tbl.style = "Table Grid"
        for k, v in [
            ("Sections excluded (customer)", str(len(self.rec.sections_excluded))),
            ("Assumption packs summarized", str(len(self.rec.sections_summarized))),
            ("OUM replacements (Rule 6)", str(self.rec.oum_replacements)),
            ("ID removals (Rule 2)", str(self.rec.id_removals)),
            ("Annotation removals (Rule 3)", str(self.rec.annotation_removals)),
            ("BU decision tables removed (Rule 4)", str(self.rec.bu_tables_removed)),
            ("Governance subsections removed (Rule 1)", str(self.rec.gov_subs_removed)),
            ("Segments in / out", f"{self.rec.segs_in} → {self.rec.segs_out}"),
        ]:
            row = tbl.add_row(); kc = row.cells[0]; vc = row.cells[1]
            _cell_bg(kc, "F2F5FB"); _cell_borders(kc, "D0D0D0", 4)
            _cell_bg(vc, HEX_WHITE); _cell_borders(vc, "D0D0D0", 4)
            kp = kc.paragraphs[0]; kp.paragraph_format.space_before = Pt(3); kp.paragraph_format.space_after = Pt(3)
            add_inline(kp, k, sz=10.0)
            vp = vc.paragraphs[0]; vp.paragraph_format.space_before = Pt(3); vp.paragraph_format.space_after = Pt(3)
            add_inline(vp, v, sz=10.0)

# ── REVIEW renderer (customer + policy annotations) ───────────────────────────

class ReviewProposalRenderer:
    def __init__(self, sections, rec: PolicyRecord, out_path: str, tender_id: str):
        self.sections = sections; self.rec = rec
        self.out_path = out_path; self.tid = tender_id
        self.doc = Document(); setup_styles(self.doc); _page_setup(self.doc)

    def render(self):
        d = self.doc
        p1 = d.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; p1.paragraph_format.space_before = Pt(24)
        r1 = p1.add_run("PPPE Review Proposal"); r1.font.name = "Calibri"; r1.font.size = Pt(22); r1.font.bold = True; r1.font.color.rgb = NAVY
        p2 = d.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(f"{self.tid}  |  {self.rec.date}  |  PPPE v1.0  |  PROFILE_REVIEW"); r2.font.name = "Calibri"; r2.font.size = Pt(11); r2.font.color.rgb = BLUE
        render_callout(d, ["**PROFILE_REVIEW** — Customer content with policy annotations. "
                           "Shows what was published AND what was removed by each rule. "
                           "For editorial review only — not for client distribution."])
        render_callout(d, [
            f"**Policy Application:** OUM replacements={self.rec.oum_replacements} | "
            f"ID removals={self.rec.id_removals} | Annotations removed={self.rec.annotation_removals} | "
            f"BU tables removed={self.rec.bu_tables_removed} | Gov subsections removed={self.rec.gov_subs_removed} | "
            f"Sections excluded={len(self.rec.sections_excluded)} | "
            f"Packs summarized={len(self.rec.sections_summarized)}"
        ])
        _page_break(d)
        for heading, segs, status in self.sections:
            _page_break(d)
            d.add_heading(heading, level=2).paragraph_format.keep_with_next = True
            for seg in segs:
                _render_seg(d, seg)
            if heading.lower() in [h.lower() for h in self.rec.sections_summarized]:
                _policy_strip_banner(d, 1, "R5")
        if self.rec.sections_excluded:
            _page_break(d)
            d.add_heading("Sections Excluded from Customer Proposal", level=1)
            for name in self.rec.sections_excluded:
                p = d.add_paragraph(style="Normal"); add_inline(p, f"• {name} — excluded by PPPE (internal-only section)")
        _headers_footers(d, "APPSolve REVIEW PROPOSAL",
                         f"{self.tid}  |  PPPE v1.0  |  REVIEW ONLY", self.tid)
        os.makedirs(os.path.dirname(self.out_path), exist_ok=True)
        d.save(self.out_path)
        sz = os.path.getsize(self.out_path)
        print(f"[PPPE REVIEW]   {self.out_path.split('/')[-1]}  ({sz//1024} KB)")

# ── Policy report ─────────────────────────────────────────────────────────────

def _write_policy_report(rec: PolicyRecord, out_path: str):
    lines = [
        f"# PPPE Policy Report — {rec.tender_id}",
        f"**Date:** {rec.date}  |  **Engine:** PPPE v1.0  |  **Profile:** PROFILE_CUSTOMER",
        "",
        "## Rule Application Summary",
        "",
        "| Rule | Description | Count |",
        "|---|---|---|",
        f"| R1 | Governance subsections removed | {rec.gov_subs_removed} |",
        f"| R2 | Internal IDs removed | {rec.id_removals} |",
        f"| R3 | Implementation annotations removed | {rec.annotation_removals} |",
        f"| R4 | BU decision tables removed | {rec.bu_tables_removed} |",
        f"| R5 | Assumption packs summarized | {len(rec.sections_summarized)} |",
        f"| R6 | OUM → Oracle Modern Best Practices replacements | {rec.oum_replacements} |",
        f"| R10 | Segments: {rec.segs_in} in → {rec.segs_out} out (reduction: {rec.segs_in - rec.segs_out}) | — |",
        "",
        "## Sections Excluded from Customer Proposal",
        "",
    ]
    for s in rec.sections_excluded:
        lines.append(f"- {s}")
    lines += [
        "",
        "## Assumption Packs Summarized",
        "",
    ]
    for s in rec.sections_summarized:
        lines.append(f"- {s}")
    lines += ["", f"_Generated: {datetime.now().isoformat()}_"]
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    print(f"[PPPE REPORT]   {out_path.split('/')[-1]}")

# ── Orchestrator ──────────────────────────────────────────────────────────────

def run_pppe(md_path: str, out_dir: str, tender_id: str) -> dict:
    print(f"[PPPE] Tender: {tender_id}")
    print(f"[PPPE] Input:  {md_path}")
    with open(md_path, "r", encoding="utf-8") as fh:
        md_text = fh.read()
    segs    = parse_md(md_text)
    v2_secs = group_v2(segs)
    classify_v2(v2_secs)
    print(f"[PPPE] V2: {len(v2_secs)} secs | "
          f"RENDERED={sum(1 for s in v2_secs if s.status=='RENDERED')} "
          f"PH={sum(1 for s in v2_secs if s.status=='PLACEHOLDER')} "
          f"AI={sum(1 for s in v2_secs if s.status=='AI_DRAFT')}")

    rec = PolicyRecord(tender_id=tender_id)
    customer_sections = process_for_customer(v2_secs, rec)
    print(f"[PPPE] Rules applied: OUM={rec.oum_replacements} IDs={rec.id_removals} "
          f"Annot={rec.annotation_removals} BU-tables={rec.bu_tables_removed} "
          f"Gov-subs={rec.gov_subs_removed} Packs-summarized={len(rec.sections_summarized)}")

    os.makedirs(out_dir, exist_ok=True)
    prefix = "PLENNEGY" if "PLENNEGY" in tender_id.upper() else tender_id

    c_path  = os.path.join(out_dir, f"PPPE_CUSTOMER_{prefix}.docx")
    i_path  = os.path.join(out_dir, f"PPPE_INTERNAL_{prefix}.docx")
    r_path  = os.path.join(out_dir, f"PPPE_REVIEW_{prefix}.docx")
    rp_path = os.path.join(out_dir, f"PPPE_POLICY_REPORT_{prefix}.md")

    CustomerProposalRenderer(customer_sections, c_path, tender_id).render()
    InternalReviewRenderer(v2_secs, rec, i_path, tender_id).render()
    ReviewProposalRenderer(customer_sections, rec, r_path, tender_id).render()
    _write_policy_report(rec, rp_path)

    print(f"\n[PPPE] Complete — 4 outputs in {out_dir}")
    return {"customer": c_path, "internal": i_path, "review": r_path, "report": rp_path}

# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Proposal Publishing Policy Engine v1.0 (PPPE)")
    parser.add_argument("--tender", metavar="TENDER_ID")
    parser.add_argument("--md",  metavar="MD_PATH")
    parser.add_argument("--out", metavar="OUT_DIR")
    args = parser.parse_args()

    if args.tender:
        tid    = args.tender
        td_dir = os.path.join(PROPOSALS_DIR, tid)
        md_path = os.path.join(td_dir, f"PROPOSAL_RENDERED_{tid}.md")
        out_dir = td_dir
    elif args.md and args.out:
        tid     = os.path.basename(args.out)
        md_path = args.md
        out_dir = args.out
    else:
        parser.print_help(); sys.exit(1)

    if not os.path.exists(md_path):
        print(f"[PPPE] ERROR: {md_path} not found", file=sys.stderr); sys.exit(1)
    run_pppe(md_path, out_dir, tid)

if __name__ == "__main__":
    main()
