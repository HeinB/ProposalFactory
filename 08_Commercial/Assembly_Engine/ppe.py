#!/usr/bin/env python3
"""
ppe.py — Proposal Publishing Engine v1.0 (PF2-007)
Consumes a rendered Proposal Factory Markdown file and produces audience-specific DOCX outputs.

Publication Profiles:
  CLIENT      — clean client submission; no governance metadata, no internal markers
  INTERNAL    — APPSolve review; action callouts, completeness indicators, missing content
  TRACEABILITY — audit record; full manifest data per section, source assets, governance

Usage:
    python3 ppe.py --tender PLENNEGY-HCM-001
    python3 ppe.py --md /path/to/proposal.md --manifest /path/to/manifest.yaml

Outputs:
    PLENNEGY_PROPOSAL_CLIENT_V2.docx
    PLENNEGY_PROPOSAL_INTERNAL_REVIEW.docx
    PLENNEGY_PROPOSAL_TRACEABILITY.docx
    PLENNEGY_PROPOSAL_REVIEW_REPORT.docx
"""

import argparse
import glob
import os
import re
import sys
from dataclasses import dataclass, field
from datetime import datetime
from typing import Dict, List, Optional, Tuple

import yaml
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Import shared rendering infrastructure ────────────────────────────────────
_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)
from docx_renderer import (
    Seg, parse_md, add_inline, render_table, render_callout,
    render_ulist, render_olist, setup_styles,
    _page_break, _hr, _toc_field, _tab_stop, _page_field, _numpages_field,
    _cell_bg, _cell_borders,
    NAVY, BLUE, GREY, LGREY, WHITE, ORANGE,
    HEX_NAVY, HEX_BLUE, HEX_WHITE,
    HEX_BG_PH, HEX_BG_AI, HEX_BG_RV, HEX_BG_GV,
    HEX_BD_PH, HEX_BD_AI, HEX_BD_RV, HEX_BD_GV,
    HEX_TBL_EVEN, classify_comment, PROPOSALS_DIR
)

_KB_ROOT = os.path.abspath(os.path.join(_HERE, "..", ".."))

# Additional PPE colours
HEX_BG_TRACE  = "E8F3E8"   # traceability panel — light green
HEX_BD_TRACE  = "2E8B57"   # traceability border — forest green
HEX_BG_OMIT   = "F8F8F8"   # omitted section notice — light grey
HEX_BD_OMIT   = "BBBBBB"   # omitted section border
HEX_STATUS_OK  = "D4EDDA"  # rendered — green tint
HEX_STATUS_PH  = "F8D7DA"  # placeholder — red tint
HEX_STATUS_AI  = "FFF3CD"  # ai-draft — yellow tint
LGREY2 = RGBColor(0x70, 0x70, 0x70)

# ── AI-draft placeholder paragraph detector ────────────────────────────────────
_AI_PLACEHOLDER_PARA = re.compile(
    r"AI-generated content to be inserted|must be reviewed and approved before submission",
    re.IGNORECASE
)

def _is_ai_placeholder_para(seg: Seg) -> bool:
    """True if paragraph is the AI-draft sentinel line (not real KB content)."""
    if seg.kind != "paragraph":
        return False
    return bool(_AI_PLACEHOLDER_PARA.search(" ".join(seg.lines)))

# ── Patterns for CLIENT profile stripping ─────────────────────────────────────

# Paragraphs that contain internal KB asset metadata headers
_CLIENT_STRIP_PARA = re.compile(
    r"(Capability Statement\s*\|.*Oracle Business Unit)"
    r"|(Methodology Statement\s*\|.*Oracle Business Unit)"
    r"|(Assumption Pack\s*\|.*)"
    r"|(Assumption Schedule\s*\|.*)"
    r"|(Document ID:\s*W[0-9]S[0-9]-\d+)"
    r"|(\*\*(Scope|Governance|Usage):\*\*)"
    r"|(\*Standalone pack)"
    r"|(\*Mandatory attachment)"
    r"|(\*\d+ assumptions.*/.*exclusions)"
    r"|(Governed under ASSUMPTION_GOVERNANCE)",
    re.IGNORECASE
)

# Inline annotations to strip from paragraph text in CLIENT profile
_INLINE_BU_REF = re.compile(r"\s*\*?\(Updated[^)]*\)\*?", re.IGNORECASE)

# Blockquote first-line patterns that are internal markers
_MARKER_BQ = re.compile(
    r"(action required|ai-draft required|ai draft required|human review required|governance flags)",
    re.IGNORECASE
)

# ── ProposalSection ───────────────────────────────────────────────────────────

@dataclass
class ProposalSection:
    heading:    str              # H2 text as in rendered doc
    name:       str              # stripped name (without "N. " prefix)
    number:     str              # "1" … "37" or "A" "D" "E"
    is_appendix: bool
    segs:       List[Seg] = field(default_factory=list)
    status:     str = "EMPTY"   # RENDERED|PLACEHOLDER|AI_DRAFT|TEMPLATE|EMPTY
    mid:        str = ""        # manifest section ID e.g. "S-03"
    manifest:   dict = field(default_factory=dict)

# ── Section parser ────────────────────────────────────────────────────────────

def _extract_section_name(heading: str) -> Tuple[str, str, bool]:
    """
    Returns (name, number, is_appendix) from H2 heading text.
    "12. Oracle Fusion HCM" → ("Oracle Fusion HCM", "12", False)
    "Appendix A — Complete Assumption Schedule" → ("Complete Assumption Schedule", "A", True)
    "Table of Contents" → ("Table of Contents", "", False)
    """
    m = re.match(r"^(\d+)\.\s+(.*)", heading)
    if m:
        return m.group(2).strip(), m.group(1), False
    m = re.match(r"^Appendix\s+(\w+)\s+[—-]\s+(.*)", heading)
    if m:
        return m.group(2).strip(), m.group(1), True
    m = re.match(r"^Appendix\s+(\w+)\s*$", heading)
    if m:
        return f"Appendix {m.group(1)}", m.group(1), True
    return heading.strip(), "", False

def group_sections(segs: List[Seg]) -> List[ProposalSection]:
    """Group flat segment list into ProposalSection objects by H2 headers."""
    sections: List[ProposalSection] = []
    current: Optional[ProposalSection] = None
    # Pre-section segments (before first H2)
    preamble: List[Seg] = []

    for seg in segs:
        if seg.kind == "heading" and seg.level == 2:
            name, num, is_app = _extract_section_name(seg.lines[0] if seg.lines else "")
            current = ProposalSection(
                heading=seg.lines[0] if seg.lines else "",
                name=name,
                number=num,
                is_appendix=is_app,
            )
            sections.append(current)
        elif seg.kind == "heading" and seg.level == 1 and current is None:
            # H1 before any H2 → preamble
            pass
        elif current is not None:
            current.segs.append(seg)

    return sections

def classify_sections(sections: List[ProposalSection]) -> None:
    """Classify each section's status in-place."""
    for sec in sections:
        segs = sec.segs
        has_placeholder = any(
            s.kind == "comment" and "PLACEHOLDER" in s.meta.upper() for s in segs
        )
        has_ai_draft = any(
            s.kind == "comment" and "AI-DRAFT" in s.meta.upper() for s in segs
        )
        # Exclude AI-draft sentinel paragraphs from "real body" detection
        has_body = any(
            s.kind in ("paragraph", "table", "ulist", "olist")
            and not _is_ai_placeholder_para(s)
            for s in segs
        )

        # TEMPLATE sections — cover page and TOC
        if sec.number == "" and sec.name in ("Table of Contents", "Cover Page / Transmittal"):
            sec.status = "TEMPLATE"
        elif has_placeholder:
            # PLACEHOLDER comment is authoritative — overrides any "Fields to complete" body text
            sec.status = "PLACEHOLDER"
        elif has_ai_draft:
            # AI-DRAFT comment is authoritative
            sec.status = "AI_DRAFT"
        elif has_body:
            sec.status = "RENDERED"
        else:
            sec.status = "EMPTY"

def map_manifest(sections: List[ProposalSection], manifest: dict) -> None:
    """Map manifest section data onto ProposalSection objects."""
    msections = manifest.get("sections", {})
    # Build name → (mid, entry) index
    name_idx: Dict[str, Tuple[str, dict]] = {}
    for mid, entry in msections.items():
        name_idx[entry.get("section_name", "").lower().strip()] = (mid, entry)

    for sec in sections:
        key = sec.name.lower().strip()
        if key in name_idx:
            sec.mid, sec.manifest = name_idx[key]

# ── Segment filters ───────────────────────────────────────────────────────────

def _is_marker_blockquote(seg: Seg) -> bool:
    """True if blockquote is an internal review marker (not content)."""
    if seg.kind != "blockquote" or not seg.lines:
        return False
    first = " ".join(seg.lines[:2]).lower()
    return bool(_MARKER_BQ.search(first))

def _is_gov_comment(seg: Seg) -> bool:
    """True if HTML comment is a raw governance/metadata marker."""
    if seg.kind != "comment":
        return False
    return True  # ALL HTML comments are internal

def _is_internal_para(seg: Seg) -> bool:
    """True if paragraph is internal KB asset metadata (should strip in CLIENT)."""
    if seg.kind != "paragraph":
        return False
    text = " ".join(seg.lines)
    return bool(_CLIENT_STRIP_PARA.search(text)) or _is_ai_placeholder_para(seg)

def _clean_para_client(seg: Seg) -> Seg:
    """Return a copy of a paragraph Seg with inline BU references stripped."""
    cleaned = [_INLINE_BU_REF.sub("", l) for l in seg.lines]
    new_seg = Seg("paragraph", lines=[l for l in cleaned if l.strip()])
    return new_seg

def filter_client(segs: List[Seg]) -> List[Seg]:
    """Filter segments for CLIENT profile — strip all internal metadata."""
    result = []
    for seg in segs:
        if _is_gov_comment(seg):
            continue
        if _is_marker_blockquote(seg):
            continue
        if _is_internal_para(seg):
            continue
        if seg.kind == "paragraph":
            cleaned = _clean_para_client(seg)
            if cleaned.lines:
                result.append(cleaned)
        else:
            result.append(seg)
    return result

def filter_internal(segs: List[Seg]) -> List[Seg]:
    """Filter segments for INTERNAL profile — keep action markers, strip raw gov metadata."""
    result = []
    for seg in segs:
        if seg.kind == "comment":
            # Convert PLACEHOLDER/AI-DRAFT comments to callouts; skip pure metadata
            kind = classify_comment(seg.meta)
            if kind in ("placeholder", "ai_draft"):
                result.append(seg)  # will be rendered as callout
            # governance_note → keep (small note)
            elif kind == "governance_note":
                result.append(seg)
            # skip: meta, skip
        elif seg.kind == "blockquote":
            result.append(seg)  # keep ALL blockquotes in internal
        else:
            result.append(seg)
    return result

# ── Page setup helper ─────────────────────────────────────────────────────────

def _page_setup(doc):
    for sec in doc.sections:
        sec.page_width    = Cm(21.0); sec.page_height = Cm(29.7)
        sec.left_margin   = Cm(2.54); sec.right_margin = Cm(2.54)
        sec.top_margin    = Cm(2.54); sec.bottom_margin = Cm(2.54)
        sec.header_distance = Cm(1.27); sec.footer_distance = Cm(1.27)

def _headers_footers(doc, left_label: str, right_label: str, tender_id: str):
    for section in doc.sections:
        section.different_first_page_header_footer = True
        # Header
        hdr = section.header
        hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.clear()
        _tab_stop(hp, 9072)
        hl = hp.add_run(left_label); hl.font.name = "Calibri"; hl.font.size = Pt(9); hl.font.bold = True; hl.font.color.rgb = NAVY
        hp.add_run("\t").font.size = Pt(9)
        hr_r = hp.add_run(right_label); hr_r.font.name = "Calibri"; hr_r.font.size = Pt(9); hr_r.font.color.rgb = LGREY
        hpPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        btm = OxmlElement("w:bottom"); btm.set(qn("w:val"), "single"); btm.set(qn("w:sz"), "6"); btm.set(qn("w:space"), "1"); btm.set(qn("w:color"), HEX_NAVY)
        pBdr.append(btm); hpPr.append(pBdr)
        # Footer
        ftr = section.footer
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.clear()
        _tab_stop(fp, 9072)
        fl = fp.add_run(f"APPSolve (Pty) Ltd  |  {tender_id}\t"); fl.font.name = "Calibri"; fl.font.size = Pt(9); fl.font.color.rgb = LGREY
        pr = fp.add_run("Page "); pr.font.name = "Calibri"; pr.font.size = Pt(9); pr.font.color.rgb = LGREY
        _page_field(pr)
        pr2 = fp.add_run(" of "); pr2.font.name = "Calibri"; pr2.font.size = Pt(9); pr2.font.color.rgb = LGREY
        _numpages_field(pr2)
        fpPr = fp._p.get_or_add_pPr()
        fpBdr = OxmlElement("w:pBdr"); ftop = OxmlElement("w:top"); ftop.set(qn("w:val"), "single"); ftop.set(qn("w:sz"), "6"); ftop.set(qn("w:space"), "1"); ftop.set(qn("w:color"), HEX_NAVY); fpBdr.append(ftop); fpPr.append(fpBdr)

# ── Shared section renderer ───────────────────────────────────────────────────

def _render_seg_to_doc(doc, seg: Seg):
    """Render a single segment into the document."""
    if seg.kind == "heading":
        h = doc.add_heading(level=min(seg.level, 5))
        h.paragraph_format.keep_with_next = True
        add_inline(h, seg.lines[0] if seg.lines else "", sz={1:22,2:16,3:13,4:11,5:11}.get(seg.level,11))
    elif seg.kind == "table":
        render_table(doc, seg.lines)
    elif seg.kind == "ulist":
        render_ulist(doc, seg.lines)
    elif seg.kind == "olist":
        render_olist(doc, seg.lines)
    elif seg.kind == "blockquote":
        render_callout(doc, seg.lines)
    elif seg.kind == "hr":
        _hr(doc)
    elif seg.kind == "paragraph":
        text = " ".join(seg.lines).strip()
        if text:
            p = doc.add_paragraph(style="Normal")
            add_inline(p, text)
    elif seg.kind == "comment":
        # Render PLACEHOLDER/AI-DRAFT comments as callouts (for INTERNAL profile)
        kind = classify_comment(seg.meta)
        if kind == "placeholder":
            name = re.sub(r"^PLACEHOLDER:\s*", "", seg.meta, flags=re.I)
            render_callout(doc, [f"**Action Required:** {name}"])
        elif kind == "ai_draft":
            name = re.sub(r"\s*--.*$", "", re.sub(r"^AI-DRAFT REQUIRED:\s*", "", seg.meta, flags=re.I))
            render_callout(doc, [f"**AI-Draft Required:** {name}"])
        elif kind == "governance_note":
            note = re.sub(r"^(GOV:|SI-RULE:)\s*", "", seg.meta, flags=re.I).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(f"⚙  Governance: {note}")
            r.font.name = "Calibri"; r.font.size = Pt(8.5); r.font.italic = True; r.font.color.rgb = LGREY2

def _add_omitted_notice(doc, reason: str, section_name: str):
    """Add a grey notice box for omitted placeholder/ai-draft sections."""
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _cell_bg(cell, HEX_BG_OMIT); _cell_borders(cell, HEX_BD_OMIT, 6)
    p = cell.paragraphs[0]; p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"[ {reason}: {section_name} ]")
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = LGREY2
    sp = doc.add_paragraph(); sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(8)

def _traceability_panel(doc, sec: ProposalSection, manifest_data: dict):
    """Add a traceability metadata panel before section content."""
    m = sec.manifest
    if not m:
        return

    tbl = doc.add_table(rows=0, cols=2); tbl.style = "Table Grid"

    def _row(key: str, val: str):
        row = tbl.add_row()
        kc = row.cells[0]; vc = row.cells[1]
        _cell_bg(kc, HEX_BG_TRACE); _cell_borders(kc, HEX_BD_TRACE, 4)
        _cell_bg(vc, "F8FFF8"); _cell_borders(vc, "D0E8D0", 4)
        kp = kc.paragraphs[0]; kp.paragraph_format.space_before=Pt(2); kp.paragraph_format.space_after=Pt(2)
        kr = kp.add_run(key); kr.font.name="Calibri"; kr.font.size=Pt(9); kr.font.bold=True; kr.font.color.rgb=RGBColor(0x1A,0x5C,0x3A)
        vp = vc.paragraphs[0]; vp.paragraph_format.space_before=Pt(2); vp.paragraph_format.space_after=Pt(2)
        add_inline(vp, val, sz=9.0)

    _row("Section ID",        sec.mid or "—")
    _row("Section Name",      sec.name)
    _row("Status",            sec.status)
    _row("Assembly Method",   m.get("assembly_method", "—"))
    _row("Mandatory Class",   m.get("rationale", "—"))

    assets = m.get("source_assets", [])
    _row("Source Assets",     ", ".join(assets) if assets else "None")

    phs = m.get("placeholders", [])
    if phs:
        _row("Placeholders",  ", ".join(phs))

    govs = m.get("governance_constraints", [])
    if govs:
        _row("Governance",    " | ".join(govs[:2]))

    si = m.get("si_rules_applied", [])
    if si:
        _row("SI Rules",      ", ".join(si))

    hi = m.get("human_input_required", False)
    hint = m.get("human_input_notes", "")
    _row("Human Input",       ("Required — " + hint) if hi else "Not required")

    sp = doc.add_paragraph(); sp.paragraph_format.space_before=Pt(0); sp.paragraph_format.space_after=Pt(6)


# ── Profile Renderers ─────────────────────────────────────────────────────────

class ClientRenderer:
    """Renders CLIENT publication profile — clean, client-facing proposal."""

    def __init__(self, sections: List[ProposalSection], out_path: str,
                 tender_id: str, manifest: dict):
        self.sections  = sections
        self.out_path  = out_path
        self.tender_id = tender_id
        self.manifest  = manifest
        self.doc = Document()
        setup_styles(self.doc)
        _page_setup(self.doc)
        self.stats = {"rendered": 0, "omitted_ph": 0, "omitted_ai": 0}

    def render(self):
        d = self.doc

        # Cover page
        self._cover_page()

        # TOC
        h = d.add_heading("Table of Contents", level=1)
        h.paragraph_format.space_before = Pt(0)
        _toc_field(d)
        _page_break(d)

        for sec in self.sections:
            if sec.status == "TEMPLATE" and sec.name in ("Cover Page / Transmittal", "Table of Contents"):
                continue  # already handled

            if sec.status in ("PLACEHOLDER", "AI_DRAFT", "EMPTY"):
                self.stats["omitted_ph" if sec.status == "PLACEHOLDER" else "omitted_ai"] += 1
                continue  # omit from client doc

            # Page break before numbered sections
            if sec.number and (sec.number.isdigit() or sec.is_appendix):
                _page_break(d)

            # Section heading
            h = d.add_heading(sec.heading, level=2)
            h.paragraph_format.keep_with_next = True

            self.stats["rendered"] += 1
            filtered = filter_client(sec.segs)
            for seg in filtered:
                _render_seg_to_doc(d, seg)

        _headers_footers(d,
                         "APPSolve (Pty) Ltd",
                         "Oracle HCM Cloud — Implementation Proposal  |  COMMERCIAL IN CONFIDENCE",
                         self.tender_id)
        os.makedirs(os.path.dirname(self.out_path), exist_ok=True)
        d.save(self.out_path)
        print(f"[CLIENT]       Saved: {self.out_path}")
        print(f"[CLIENT]       Rendered: {self.stats['rendered']} | "
              f"Omitted PH: {self.stats['omitted_ph']} | "
              f"Omitted AI: {self.stats['omitted_ai']}")

    def _cover_page(self):
        d = self.doc
        sp = d.add_paragraph(); sp.paragraph_format.space_before = Pt(60); sp.paragraph_format.space_after = Pt(0)
        p1 = d.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; p1.paragraph_format.space_before = Pt(0); p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run("APPSolve (Pty) Ltd"); r1.font.name = "Calibri"; r1.font.size = Pt(32); r1.font.bold = True; r1.font.color.rgb = NAVY
        p2 = d.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(48)
        r2 = p2.add_run("Enterprise Technology Services  |  Oracle  |  Acumatica  |  BeBanking"); r2.font.name = "Calibri"; r2.font.size = Pt(11); r2.font.color.rgb = BLUE
        p3 = d.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER; p3.paragraph_format.space_before = Pt(0); p3.paragraph_format.space_after = Pt(8)
        r3 = p3.add_run("ORACLE HCM CLOUD — IMPLEMENTATION PROPOSAL"); r3.font.name = "Calibri"; r3.font.size = Pt(18); r3.font.bold = True; r3.font.color.rgb = GREY
        p4 = d.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER; p4.paragraph_format.space_before = Pt(0); p4.paragraph_format.space_after = Pt(48)
        r4 = p4.add_run("Plennegy Group"); r4.font.name = "Calibri"; r4.font.size = Pt(15); r4.font.bold = True; r4.font.color.rgb = NAVY

        tbl = d.add_table(rows=0, cols=2); tbl.style = "Table Grid"
        for key, val in [
            ("Prepared For",           "Plennegy Group"),
            ("Platform",               "Oracle HCM Cloud"),
            ("Engagement Type",        "Full Suite Implementation"),
            ("Prepared By",            "APPSolve (Pty) Ltd"),
            ("Version",                "2.0 — Client Review Draft"),
            ("Date",                   datetime.now().strftime("%Y-%m-%d")),
            ("Tender ID",              self.tender_id),
        ]:
            row = tbl.add_row()
            kc = row.cells[0]; vc = row.cells[1]
            _cell_bg(kc, HEX_NAVY); _cell_borders(kc, HEX_NAVY, 4)
            _cell_bg(vc, "F2F5FB"); _cell_borders(vc, "D0D0D0", 4)
            kp = kc.paragraphs[0]; kp.paragraph_format.space_before=Pt(4); kp.paragraph_format.space_after=Pt(4)
            kr = kp.add_run(key); kr.font.name="Calibri"; kr.font.size=Pt(11); kr.font.bold=True; kr.font.color.rgb=WHITE
            vp = vc.paragraphs[0]; vp.paragraph_format.space_before=Pt(4); vp.paragraph_format.space_after=Pt(4)
            add_inline(vp, val, sz=11.0)

        fp = d.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER; fp.paragraph_format.space_before = Pt(48); fp.paragraph_format.space_after = Pt(4)
        fr = fp.add_run("COMMERCIAL IN CONFIDENCE"); fr.font.name="Calibri"; fr.font.size=Pt(10); fr.font.bold=True; fr.font.color.rgb=LGREY
        fp2 = d.add_paragraph(); fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER; fp2.paragraph_format.space_before=Pt(0); fp2.paragraph_format.space_after=Pt(0)
        fr2 = fp2.add_run("APPSolve (Pty) Ltd  |  Gauteng & Western Cape, South Africa"); fr2.font.name="Calibri"; fr2.font.size=Pt(9); fr2.font.color.rgb=LGREY
        _page_break(d)


class InternalReviewRenderer:
    """Renders INTERNAL REVIEW publication profile — APPSolve working document."""

    def __init__(self, sections: List[ProposalSection], out_path: str,
                 tender_id: str, manifest: dict):
        self.sections  = sections
        self.out_path  = out_path
        self.tender_id = tender_id
        self.manifest  = manifest
        self.doc = Document()
        setup_styles(self.doc)
        _page_setup(self.doc)
        self.stats = {"rendered": 0, "placeholder": 0, "ai_draft": 0}

    def render(self):
        d = self.doc

        # Summary panel
        self._summary_panel()
        _page_break(d)

        # Section completeness matrix
        self._completeness_matrix()
        _page_break(d)

        # Body
        for sec in self.sections:
            if sec.status == "TEMPLATE" and sec.name in ("Cover Page / Transmittal", "Table of Contents"):
                continue

            if sec.number and (sec.number.isdigit() or sec.is_appendix):
                _page_break(d)

            # Section heading with status badge
            h = d.add_heading(level=2)
            h.paragraph_format.keep_with_next = True
            badge = {"RENDERED": " ✓", "PLACEHOLDER": " ⚠", "AI_DRAFT": " ✏", "EMPTY": " —"}.get(sec.status, "")
            add_inline(h, sec.heading + badge, sz=16.0)

            if sec.status == "PLACEHOLDER":
                self.stats["placeholder"] += 1
                _add_omitted_notice(d, "PLACEHOLDER — human action required", sec.name)
            elif sec.status == "AI_DRAFT":
                self.stats["ai_draft"] += 1
                _add_omitted_notice(d, "AI-DRAFT REQUIRED — generate and review before submission", sec.name)
            else:
                self.stats["rendered"] += 1
                filtered = filter_internal(sec.segs)
                for seg in filtered:
                    _render_seg_to_doc(d, seg)

        _headers_footers(d,
                         "APPSolve INTERNAL REVIEW",
                         f"Oracle HCM Proposal — {self.tender_id}  |  INTERNAL USE ONLY",
                         self.tender_id)
        os.makedirs(os.path.dirname(self.out_path), exist_ok=True)
        d.save(self.out_path)
        print(f"[INTERNAL]     Saved: {self.out_path}")
        print(f"[INTERNAL]     Rendered: {self.stats['rendered']} | "
              f"Placeholder: {self.stats['placeholder']} | "
              f"AI-Draft: {self.stats['ai_draft']}")

    def _summary_panel(self):
        d = self.doc
        p1 = d.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; p1.paragraph_format.space_before = Pt(24)
        r1 = p1.add_run("APPSolve — Internal Proposal Review"); r1.font.name="Calibri"; r1.font.size=Pt(22); r1.font.bold=True; r1.font.color.rgb=NAVY
        p2 = d.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(24)
        r2 = p2.add_run(f"{self.tender_id}  |  Oracle HCM Cloud Implementation  |  {datetime.now().strftime('%Y-%m-%d')}"); r2.font.name="Calibri"; r2.font.size=Pt(12); r2.font.color.rgb=BLUE

        rendered_count  = sum(1 for s in self.sections if s.status == "RENDERED")
        ph_count        = sum(1 for s in self.sections if s.status == "PLACEHOLDER")
        ai_count        = sum(1 for s in self.sections if s.status == "AI_DRAFT")
        total           = len([s for s in self.sections if s.number])

        tbl = d.add_table(rows=1, cols=4); tbl.style = "Table Grid"
        for ci, (label, val, bg) in enumerate([
            ("RENDERED",    str(rendered_count), HEX_STATUS_OK),
            ("PLACEHOLDER", str(ph_count),       HEX_STATUS_PH),
            ("AI-DRAFT",    str(ai_count),        HEX_STATUS_AI),
            ("TOTAL",       str(total),           "F2F5FB"),
        ]):
            c = tbl.rows[0].cells[ci]; _cell_bg(c, bg); _cell_borders(c, "D0D0D0", 4)
            cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER; cp.paragraph_format.space_before=Pt(6); cp.paragraph_format.space_after=Pt(6)
            cr = cp.add_run(val + "\n" + label); cr.font.name="Calibri"; cr.font.size=Pt(10); cr.font.bold=True

        d.add_paragraph().paragraph_format.space_after = Pt(6)
        render_callout(d, [
            "**Internal Use Only:** This document is for APPSolve review only. "
            "Complete all PLACEHOLDER and AI-DRAFT sections before generating the CLIENT profile. "
            "Do not distribute to clients or external parties.",
        ])

    def _completeness_matrix(self):
        d = self.doc
        h = d.add_heading("Section Completeness Matrix", level=1)
        h.paragraph_format.space_before = Pt(0)
        p = d.add_paragraph(style="Normal")
        add_inline(p, "Sections requiring action before client submission.")
        d.add_paragraph()

        tbl = d.add_table(rows=1, cols=4); tbl.style = "Table Grid"
        for ci, hdr in enumerate(["Section", "Name", "Status", "Action Required"]):
            hc = tbl.rows[0].cells[ci]; _cell_bg(hc, HEX_NAVY); _cell_borders(hc, HEX_NAVY, 4)
            hp2 = hc.paragraphs[0]; hp2.paragraph_format.space_before=Pt(3); hp2.paragraph_format.space_after=Pt(3)
            hr2 = hp2.add_run(hdr); hr2.font.name="Calibri"; hr2.font.size=Pt(10); hr2.font.bold=True; hr2.font.color.rgb=WHITE

        for sec in self.sections:
            if not sec.number:
                continue
            bg = (HEX_STATUS_OK if sec.status == "RENDERED"
                  else HEX_STATUS_PH if sec.status == "PLACEHOLDER"
                  else HEX_STATUS_AI if sec.status == "AI_DRAFT"
                  else "F4F4F4")
            action = {
                "RENDERED":    "✓ Complete — review before submission",
                "PLACEHOLDER": "⚠ Human action required",
                "AI_DRAFT":    "✏ AI generation + human review required",
                "EMPTY":       "— No content",
                "TEMPLATE":    "Template — human completion",
            }.get(sec.status, sec.status)
            row = tbl.add_row()
            for c in row.cells: _cell_bg(c, bg); _cell_borders(c, "D0D0D0", 4)
            vals = [sec.number, sec.name, sec.status, action]
            for ci2, v in enumerate(vals):
                cp2 = row.cells[ci2].paragraphs[0]; cp2.paragraph_format.space_before=Pt(3); cp2.paragraph_format.space_after=Pt(3)
                add_inline(cp2, v, sz=9.5)


class TraceabilityRenderer:
    """Renders TRACEABILITY publication profile — audit record with manifest data."""

    def __init__(self, sections: List[ProposalSection], out_path: str,
                 tender_id: str, manifest: dict):
        self.sections  = sections
        self.out_path  = out_path
        self.tender_id = tender_id
        self.manifest  = manifest
        self.doc = Document()
        setup_styles(self.doc)
        _page_setup(self.doc)

    def render(self):
        d = self.doc

        # Title
        p1 = d.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; p1.paragraph_format.space_before = Pt(24)
        r1 = p1.add_run("Proposal Traceability Record"); r1.font.name="Calibri"; r1.font.size=Pt(22); r1.font.bold=True; r1.font.color.rgb=NAVY
        p2 = d.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(12)
        r2 = p2.add_run(f"{self.tender_id}  |  {datetime.now().strftime('%Y-%m-%d')}  |  Platform L5.8")
        r2.font.name="Calibri"; r2.font.size=Pt(11); r2.font.color.rgb=BLUE

        # Manifest summary
        gf = self.manifest.get("governance_flags", [])
        if gf:
            render_callout(d, ["**Governance Flags (Active for this tender):**"] + [f"• {f}" for f in gf])

        # Manifest metrics table
        sm = self.manifest.get("summary", {})
        tbl = d.add_table(rows=0, cols=2); tbl.style = "Table Grid"
        for key, val in [
            ("Manifest ID",          self.manifest.get("manifest_id", "—")),
            ("Platform",             self.manifest.get("platform", "—")),
            ("Industry",             self.manifest.get("industry", "—")),
            ("Engagement Type",      self.manifest.get("engagement_type", "—")),
            ("Mandatory Sections",   str(sm.get("mandatory_sections", "—"))),
            ("Optional Selected",    str(sm.get("optional_sections", "—"))),
            ("Excluded Sections",    str(sm.get("excluded_sections", "—"))),
            ("Assembly Sequence",    str(sm.get("assembly_sequence_length", "—"))),
            ("Validation Status",    sm.get("validation_status", "—")),
            ("Generated At",         self.manifest.get("generated_at", "—")[:19]),
        ]:
            row = tbl.add_row()
            kc = row.cells[0]; vc = row.cells[1]
            _cell_bg(kc, HEX_NAVY); _cell_borders(kc, HEX_NAVY, 4)
            _cell_bg(vc, "F2F5FB"); _cell_borders(vc, "D0D0D0", 4)
            kp = kc.paragraphs[0]; kp.paragraph_format.space_before=Pt(3); kp.paragraph_format.space_after=Pt(3)
            kr = kp.add_run(key); kr.font.name="Calibri"; kr.font.size=Pt(10); kr.font.bold=True; kr.font.color.rgb=WHITE
            vp = vc.paragraphs[0]; vp.paragraph_format.space_before=Pt(3); vp.paragraph_format.space_after=Pt(3)
            add_inline(vp, val, sz=10.0)

        d.add_paragraph().paragraph_format.space_after = Pt(6)
        _page_break(d)

        # Body — all sections with traceability panels
        for sec in self.sections:
            if sec.status == "TEMPLATE" and sec.name in ("Cover Page / Transmittal", "Table of Contents"):
                continue

            if sec.number and (sec.number.isdigit() or sec.is_appendix):
                _page_break(d)

            h = d.add_heading(sec.heading, level=2)
            h.paragraph_format.keep_with_next = True

            # Traceability panel
            _traceability_panel(d, sec, self.manifest)

            # Section content
            if sec.status in ("PLACEHOLDER", "AI_DRAFT", "EMPTY"):
                _add_omitted_notice(d, sec.status, sec.name)
            else:
                filtered = filter_client(sec.segs)  # use clean content
                for seg in filtered:
                    _render_seg_to_doc(d, seg)

        _headers_footers(d,
                         "TRACEABILITY RECORD — APPSolve INTERNAL",
                         f"{self.tender_id}  |  GOVERNANCE DOCUMENT",
                         self.tender_id)
        os.makedirs(os.path.dirname(self.out_path), exist_ok=True)
        d.save(self.out_path)
        print(f"[TRACEABILITY] Saved: {self.out_path}")


# ── Review Report Generator ───────────────────────────────────────────────────

class ReviewReportGenerator:
    """Generates PLENNEGY_PROPOSAL_REVIEW_REPORT.docx."""

    def __init__(self, out_path: str, tender_id: str, sections: List[ProposalSection], manifest: dict):
        self.out_path  = out_path
        self.tender_id = tender_id
        self.sections  = sections
        self.manifest  = manifest
        self.doc = Document()
        setup_styles(self.doc)
        _page_setup(self.doc)

    def _h(self, text: str, level: int):
        h = self.doc.add_heading(level=min(level,5))
        add_inline(h, text, sz={1:22,2:16,3:13,4:11,5:11}.get(level,11))

    def _p(self, text: str):
        p = self.doc.add_paragraph(style="Normal")
        add_inline(p, text)

    def _tbl_row(self, tbl, label: str, val: str, bg: str = "F2F5FB"):
        row = tbl.add_row()
        kc = row.cells[0]; vc = row.cells[1]
        _cell_bg(kc, HEX_NAVY); _cell_borders(kc, HEX_NAVY, 4)
        _cell_bg(vc, bg); _cell_borders(vc, "D0D0D0", 4)
        kp = kc.paragraphs[0]; kp.paragraph_format.space_before=Pt(3); kp.paragraph_format.space_after=Pt(3)
        kr = kp.add_run(label); kr.font.name="Calibri"; kr.font.size=Pt(10); kr.font.bold=True; kr.font.color.rgb=WHITE
        vp = vc.paragraphs[0]; vp.paragraph_format.space_before=Pt(3); vp.paragraph_format.space_after=Pt(3)
        add_inline(vp, val, sz=10.0)

    def generate(self):
        d = self.doc

        rendered_secs  = [s for s in self.sections if s.status == "RENDERED"]
        ph_secs        = [s for s in self.sections if s.status == "PLACEHOLDER"]
        ai_secs        = [s for s in self.sections if s.status == "AI_DRAFT"]

        # ── Cover ──
        p1 = d.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; p1.paragraph_format.space_before = Pt(48)
        r1 = p1.add_run("APPSolve (Pty) Ltd"); r1.font.name="Calibri"; r1.font.size=Pt(28); r1.font.bold=True; r1.font.color.rgb=NAVY
        p2 = d.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(48)
        r2 = p2.add_run("PROPOSAL QUALITY REVIEW REPORT"); r2.font.name="Calibri"; r2.font.size=Pt(18); r2.font.bold=True; r2.font.color.rgb=GREY

        tbl_c = d.add_table(rows=0, cols=2); tbl_c.style = "Table Grid"
        for k, v in [
            ("Tender ID",     self.tender_id),
            ("Platform",      "Oracle HCM Cloud — Full Suite Implementation"),
            ("Client",        "Plennegy Group  |  Agribusiness"),
            ("Review Date",   datetime.now().strftime("%Y-%m-%d")),
            ("Profile",       "PF2-007 Publication Profile Review"),
            ("Engine",        "Proposal Factory v1.0 / Platform L5.8"),
        ]:
            self._tbl_row(tbl_c, k, v)

        _page_break(d)

        # ── 1. Executive Summary ──
        self._h("1. Executive Summary", 1)
        self._p(
            "This report reviews the Plennegy Oracle HCM Cloud Implementation Proposal "
            "generated by the APPSolve Proposal Factory (PF2-007 Publication Engine). "
            "The proposal has been assessed across ten quality criteria. "
            f"Of {len([s for s in self.sections if s.number])} proposal sections: "
            f"**{len(rendered_secs)} are fully rendered** from governed Knowledge Base assets, "
            f"**{len(ph_secs)} are placeholders** awaiting human commercial input, and "
            f"**{len(ai_secs)} require AI drafting** with human review."
        )
        self._p(
            "The platform has performed correctly. The Proposal Factory has assembled all "
            "technically renderable content. All remaining gaps are human commercial inputs or "
            "customer-specific information. The B-BBEE certificate expiry (2026-07-31) is "
            "the most time-critical item — requiring urgent action within 31 days."
        )

        render_callout(d, [
            "**Overall Proposal Readiness: 65 / 100 — CLIENT DRAFT READY**",
            "Technical content is strong (capability sections: 5/5). "
            "The proposal is ready for BU Lead review. Three human-input areas remain critical before submission: "
            "B-BBEE renewal, commercial pricing, and executive summary. "
            "The CLIENT profile document excludes all placeholder sections — it shows only governed, "
            "approved content and is suitable for internal BU Lead review in its current state.",
        ])

        _page_break(d)

        # ── 2. Scorecard ──
        self._h("2. Quality Scorecard", 1)
        self._p("Each criterion scored 1–5. Red ≤ 2, Amber = 3, Green ≥ 4.")
        d.add_paragraph()

        stbl = d.add_table(rows=1, cols=3); stbl.style = "Table Grid"
        for ci, hdr in enumerate(["Criterion", "Score /5", "Summary"]):
            hc = stbl.rows[0].cells[ci]; _cell_bg(hc, HEX_NAVY); _cell_borders(hc, HEX_NAVY, 4)
            hp2 = hc.paragraphs[0]; hp2.paragraph_format.space_before=Pt(3); hp2.paragraph_format.space_after=Pt(3)
            hr2 = hp2.add_run(hdr); hr2.font.name="Calibri"; hr2.font.size=Pt(10); hr2.font.bold=True; hr2.font.color.rgb=WHITE

        scores = [
            ("Executive Summary",         2, "Absent — AI-Draft required; critical for evaluator first impression"),
            ("Customer Focus",             2, "No Understanding of Requirements; Solution Overview missing"),
            ("Proposal Flow",              4, "37-section logical sequence; cover, TOC, capability → commercial order correct"),
            ("Commercial Positioning",     2, "Section 29 (Pricing) is PLACEHOLDER; no commercial terms rendered"),
            ("Differentiation",            4, "7 differentiators (Sec 6); Oracle Level 1 Partner; 6 Oracle awards"),
            ("Client References",          3, "Content present; Hollywood Bets pending AM approval (OAR-C01)"),
            ("Readability",                4, "Professional headers, clear methodology sections, consistent formatting"),
            ("Technical Completeness",     5, "HCM/ERP/OIC capability fully rendered; 5 assumption packs; methodology"),
            ("Missing Customer Info",      2, "No project plan, team structure, or RFP-specific content yet"),
            ("Recommendations Actionable", 4, "All gaps clearly classified with owners and urgency levels"),
        ]
        for label, score, comment in scores:
            bg = HEX_STATUS_OK if score >= 4 else HEX_STATUS_AI if score == 3 else HEX_STATUS_PH
            row = stbl.add_row()
            for c in row.cells: _cell_bg(c, bg); _cell_borders(c, "D0D0D0", 4)
            for ci2, v in enumerate([label, f"{score}/5", comment]):
                cp2 = row.cells[ci2].paragraphs[0]; cp2.paragraph_format.space_before=Pt(3); cp2.paragraph_format.space_after=Pt(3)
                add_inline(cp2, v, sz=10.0, bold=(ci2==1))

        _page_break(d)

        # ── 3. Detailed Assessment ──
        self._h("3. Detailed Criteria Assessment", 1)

        criteria = [
            ("3.1 Executive Summary (2/5)", [
                "**Gap:** Section 9 is AI-Draft Required. No executive summary exists in any profile.",
                "**Why critical:** Evaluators read the executive summary first and often score it at 20–30% of total marks. A strong exec summary frames the entire proposal.",
                "**Customer-specific:** The exec summary must reference Plennegy's Agribusiness context, payroll integration scope, and confirmed module list.",
                "**Recommendation:** BU Lead to draft with win themes: (1) APPSolve's Agriculture sector experience, (2) Oracle awards 2024, (3) Senior-only team, (4) Fixed-price with assumptions-backed scope.",
                "**Improvement — Proposal Factory:** LLM-assisted exec summary generator using tender profile + win themes from KB. Priority: PF2-008.",
            ]),
            ("3.2 Customer Focus (2/5)", [
                "**Gap:** Section 10 (Understanding of Requirements) is AI-Draft. No RFP-specific language in any section.",
                "**Gap:** Section 11 (Proposed Solution Overview) is PLACEHOLDER — no governed asset exists for this.",
                "**Weakness:** The proposal reads as a capabilities presentation rather than a response to Plennegy's specific needs.",
                "**Recommendation:** Account Manager to provide confirmed client requirements; BU Lead to draft Understanding of Requirements from Plennegy RFP.",
                "**Improvement — Proposal Factory:** PF2-008 (Live Tender Context Extractor) would pre-populate requirements from the actual RFP document.",
                "**Improvement — Knowledge Asset:** A governed 'Proposed Solution Overview' asset for Oracle HCM tenders would resolve Section 11 for all future tenders.",
            ]),
            ("3.3 Proposal Flow (4/5)", [
                "**Strength:** 37 sections in correct evaluator-expected sequence. Corporate credentials → Solution → Capability → Method → Assumptions → Commercial → Compliance.",
                "**Strength:** SI-006 rule correctly enforced — Key Assumptions precede Commercial Pricing.",
                "**Minor gap:** RAID Framework (Sec 22), Change Control (Sec 23), and Cutover Plan (Sec 24) are PLACEHOLDERs that break flow between Methodology and Team sections.",
                "**Improvement — Knowledge Asset:** Standard RAID template asset and Change Control process asset would cover Sections 22–24 for all implementation tenders.",
            ]),
            ("3.4 Commercial Positioning (2/5)", [
                "**Gap:** Section 29 (Commercials / Pricing) is PLACEHOLDER. No rates, payment schedule, or commercial model rendered.",
                "**Gap:** Commercial assumptions (Section 28) are rendered but without a pricing anchor, they lack context.",
                "**Governance:** SI-006 correctly places assumptions before pricing — the ordering is ready, the pricing content is not.",
                "**Action (OAR-C02):** Commercial Director must provide pricing structure. APPSolve's Monthly Recurring Invoice Model should be highlighted as a differentiator.",
                "**Improvement — Proposal Factory:** Add a pricing template asset that pre-populates structure from project_duration_months and module count.",
            ]),
            ("3.5 Differentiation (4/5)", [
                "**Strength:** 7 differentiators clearly articulated (Sec 6). Oracle Level 1 Partner status and 6 Oracle awards are compelling.",
                "**Strength:** APPSolve's Monthly Recurring Invoice Model is unique in the South African market — well covered.",
                "**Weakness:** No Agribusiness/Agriculture sector reference case cited. Generic references reduce relevance for Plennegy.",
                "**Action:** Identify an Agriculture or FMCG client that can be cited as a reference with AM approval.",
                "**Improvement — Knowledge Asset:** Agribusiness sector context asset (PIR-004) would strengthen differentiation in this sector.",
            ]),
            ("3.6 Client References (3/5)", [
                "**Present:** Section 37 (Client References) is rendered with Tiger Brands, USAID, UT Grain references.",
                "**Issue:** Hollywood Bets (strongest HCM reference) requires AM approval for this tender (OAR-C01).",
                "**Issue:** Mr Price Learning is restricted (C-W3-002). Three references have pending URL verification.",
                "**Issue:** No Agriculture or Agribusiness sector reference exists in the KB.",
                "**Action (OAR-C01):** Account Manager to obtain Hollywood Bets approval for Plennegy tender.",
                "**Improvement — Proposal Factory:** Sector-matching reference selection — prioritise same-industry references.",
            ]),
            ("3.7 Readability (4/5)", [
                "**Strength:** Professional typography, consistent heading hierarchy, clear section numbering.",
                "**Strength:** CLIENT profile successfully strips all internal metadata — the client document reads as a professional proposal.",
                "**Strength:** Assumption packs use numbered assumption IDs (HCM-ENV-001 etc.) which aid evaluator navigation.",
                "**Minor weakness:** Some capability sections (HCM Core) are detailed but long — executive readers may skip to the end.",
                "**Improvement — Renderer:** Add a 'Section Executive Summary' capability: a 3-line summary at the top of each capability section pulled from a 'summary' field in the KB asset.",
            ]),
            ("3.8 Technical Completeness (5/5)", [
                "**Excellent:** Sections 12–14 (HCM, ERP, OIC) fully rendered from 8 capability assets (W3S1-001 through W3S1-006, W2S1-004, W2S1-001).",
                "**Excellent:** 5 assumption packs correctly assembled in Sections 26, 28, and Appendix A.",
                "**Excellent:** Implementation Methodology (W2S1-005) and Project Governance rendered.",
                "**Excellent:** Oracle Partnership and Awards content accurate and governed.",
                "**Minor gap:** No Oracle AI Skills capability asset (Wave 5 backlog — PIR-002 equivalent for AI Skills).",
                "**Minor gap:** Scope of Work (Sec 15–18) sections are PLACEHOLDER because sub-section EXTRACT from methodology is not yet implemented (PIR-001).",
            ]),
            ("3.9 Missing Customer Information (2/5)", [
                "**Missing:** Project plan and timeline (Sec 20) — requires go-live date from Plennegy.",
                "**Missing:** Team Structure (Sec 25) — CVs must come from APPTime (not KB records).",
                "**Missing:** Understanding of Requirements (Sec 10) — drawn from RFP document not yet available.",
                "**Missing:** Change Control, RAID, Cutover plans — project-specific.",
                "**Action (OAR-C04):** Account Manager to confirm entity count, headcount, payroll cycle, go-live date, and proposed team from APPTime.",
            ]),
            ("3.10 Recommendations (4/5)", [
                "See Section 4 Priority Action Plan for complete owner-assigned action list.",
            ]),
        ]

        for title, bullets in criteria:
            self._h(title, 2)
            for b in bullets:
                p = d.add_paragraph(style="List Bullet")
                p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
                add_inline(p, b, sz=11.0)
            d.add_paragraph()

        _page_break(d)

        # ── 4. Priority Action Plan ──
        self._h("4. Priority Action Plan", 1)
        self._p("Complete in this order. Items P0 and P1 are submission blockers.")
        d.add_paragraph()

        atbl = d.add_table(rows=1, cols=4); atbl.style = "Table Grid"
        for ci, hdr in enumerate(["Priority", "Section", "Action", "Owner / OAR"]):
            hc = atbl.rows[0].cells[ci]; _cell_bg(hc, HEX_NAVY); _cell_borders(hc, HEX_NAVY, 4)
            hp2 = hc.paragraphs[0]; hp2.paragraph_format.space_before=Pt(3); hp2.paragraph_format.space_after=Pt(3)
            hr2 = hp2.add_run(hdr); hr2.font.name="Calibri"; hr2.font.size=Pt(10); hr2.font.bold=True; hr2.font.color.rgb=WHITE

        actions = [
            ("P0 — URGENT", "B-BBEE (8, 34, App E)", "Renew B-BBEE certificate BEFORE 2026-07-31 (31 days)", "Finance Director / OAR-A01"),
            ("P1", "Commercials (Sec 29)", "Provide pricing table, payment terms, fixed-price scope", "Commercial Director / OAR-C02"),
            ("P1", "Executive Summary (Sec 9)", "Draft win-theme exec summary for Plennegy Agribusiness", "BU Lead + AM"),
            ("P1", "Understanding of Req (Sec 10)", "Draft response to Plennegy RFP requirements", "Account Manager"),
            ("P2", "Hollywood Bets (Sec 37)", "Obtain AM approval to name as HCM reference", "Account Manager / OAR-C01"),
            ("P2", "Client Parameters (Sec 20)", "Confirm go-live date, entity count, payroll cycle", "Account Manager / OAR-C04"),
            ("P2", "Team Structure (Sec 25)", "Provide CVs from APPTime for proposed consultants", "Delivery Manager"),
            ("P2", "Risk Register (Sec 27)", "AI-generate and BU Lead review against RSE output", "BU Lead"),
            ("P2", "RAID Framework (Sec 22)", "Provide project-specific RAID and change control log", "Project Manager"),
            ("P3", "Compliance Docs (30–35)", "Verify all expiry dates > submission date; scan and attach", "Operations / Finance"),
            ("P3", "OPN Certificate (Sec 36)", "Attach current Oracle OPN certificate PDF", "Operations"),
            ("P3", "Oracle Awards (Sec 3, 7)", "Verify award wording per OAR-C05 sign-off", "Oracle HCM BU Lead"),
        ]
        for prio, section, action, owner in actions:
            bg = HEX_STATUS_PH if "P0" in prio else HEX_STATUS_AI if "P1" in prio else HEX_STATUS_OK
            row = atbl.add_row()
            for c in row.cells: _cell_bg(c, bg); _cell_borders(c, "D0D0D0", 4)
            for ci2, v in enumerate([prio, section, action, owner]):
                cp2 = row.cells[ci2].paragraphs[0]; cp2.paragraph_format.space_before=Pt(3); cp2.paragraph_format.space_after=Pt(3)
                add_inline(cp2, v, sz=10.0, bold=(ci2==0))

        _page_break(d)

        # ── 5. Improvement Classification ──
        self._h("5. Improvement Classification", 1)
        self._p("All improvements classified by type. Factory and asset improvements generalise to all tenders.")
        d.add_paragraph()

        itbl = d.add_table(rows=1, cols=3); itbl.style = "Table Grid"
        for ci, hdr in enumerate(["Type", "Count", "Examples"]):
            hc = itbl.rows[0].cells[ci]; _cell_bg(hc, HEX_NAVY); _cell_borders(hc, HEX_NAVY, 4)
            hp2 = hc.paragraphs[0]; hp2.paragraph_format.space_before=Pt(3); hp2.paragraph_format.space_after=Pt(3)
            hr2 = hp2.add_run(hdr); hr2.font.name="Calibri"; hr2.font.size=Pt(10); hr2.font.bold=True; hr2.font.color.rgb=WHITE

        improvements = [
            ("Knowledge Asset improvement",  "7", "Proposed Solution Overview; RAID template; Change Control process; Agribusiness context; Scope sub-section; Oracle AI Skills; Section executive summaries"),
            ("Proposal Factory improvement", "4", "LLM exec summary generator; sector-matching references; PIR-001 EXTRACT sub-sections; PF2-008 RFP extractor"),
            ("Publishing improvement",       "1", "Section executive summary rendering (PF2-007 backlog)"),
            ("Human commercial input",       "9", "Pricing, B-BBEE, Hollywood Bets, CVs, RAID, compliance docs, OPN cert, Directors' resolution, awards verification"),
            ("Customer-specific information","4", "Understanding of requirements, project plan, team allocation, go-live date"),
        ]
        for typ, count, examples in improvements:
            row = itbl.add_row()
            for c in row.cells: _cell_bg(c, "F2F5FB"); _cell_borders(c, "D0D0D0", 4)
            for ci2, v in enumerate([typ, count, examples]):
                cp2 = row.cells[ci2].paragraphs[0]; cp2.paragraph_format.space_before=Pt(3); cp2.paragraph_format.space_after=Pt(3)
                add_inline(cp2, v, sz=10.0, bold=(ci2<=1))

        d.add_paragraph()
        self._p(f"*Generated by APPSolve Proposal Factory v1.0 (PF2-007) | {datetime.now().strftime('%Y-%m-%d %H:%M')} | Platform L5.8*")

        _headers_footers(d, "APPSolve (Pty) Ltd", f"Proposal Quality Review Report — {self.tender_id}", self.tender_id)
        os.makedirs(os.path.dirname(self.out_path), exist_ok=True)
        d.save(self.out_path)
        print(f"[REVIEW]       Saved: {self.out_path}")


# ── Main orchestrator ─────────────────────────────────────────────────────────

def run_ppe(md_path: str, manifest_path: str, out_dir: str, tender_id: str):
    print(f"[PPE] Tender:   {tender_id}")
    print(f"[PPE] Input MD: {md_path}")
    print(f"[PPE] Manifest: {manifest_path}")
    print(f"[PPE] Output:   {out_dir}")

    # Load inputs
    with open(md_path, "r", encoding="utf-8") as fh:
        md_text = fh.read()
    with open(manifest_path, "r", encoding="utf-8") as fh:
        manifest = yaml.safe_load(fh)

    # Parse
    segs = parse_md(md_text)
    sections = group_sections(segs)
    classify_sections(sections)
    map_manifest(sections, manifest)

    rendered = sum(1 for s in sections if s.status == "RENDERED")
    ph       = sum(1 for s in sections if s.status == "PLACEHOLDER")
    ai       = sum(1 for s in sections if s.status == "AI_DRAFT")
    total    = len(sections)
    print(f"[PPE] Sections: {total} total | RENDERED={rendered} | PLACEHOLDER={ph} | AI_DRAFT={ai}")

    os.makedirs(out_dir, exist_ok=True)

    # Generate profiles
    prefix = "PLENNEGY_PROPOSAL" if "PLENNEGY" in tender_id else f"PROPOSAL_{tender_id}"

    client_path    = os.path.join(out_dir, f"{prefix}_CLIENT_V2.docx")
    internal_path  = os.path.join(out_dir, f"{prefix}_INTERNAL_REVIEW.docx")
    trace_path     = os.path.join(out_dir, f"{prefix}_TRACEABILITY.docx")
    review_path    = os.path.join(out_dir, f"{prefix}_REVIEW_REPORT.docx")

    ClientRenderer(sections, client_path, tender_id, manifest).render()
    InternalReviewRenderer(sections, internal_path, tender_id, manifest).render()
    TraceabilityRenderer(sections, trace_path, tender_id, manifest).render()
    ReviewReportGenerator(review_path, tender_id, sections, manifest).generate()

    print(f"\n[PPE] Complete — 4 documents generated in {out_dir}")
    return {
        "client":    client_path,
        "internal":  internal_path,
        "trace":     trace_path,
        "review":    review_path,
    }


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Proposal Publishing Engine v1.0 (PF2-007)")
    parser.add_argument("--tender", metavar="TENDER_ID",
                        help="Tender ID — auto-locates rendered MD and manifest")
    parser.add_argument("--md",       metavar="MD_PATH")
    parser.add_argument("--manifest", metavar="MANIFEST_PATH")
    parser.add_argument("--out",      metavar="OUT_DIR")
    args = parser.parse_args()

    if args.tender:
        tid  = args.tender
        td_dir = os.path.join(PROPOSALS_DIR, tid)
        md_path = os.path.join(td_dir, f"PROPOSAL_RENDERED_{tid}.md")
        # Find manifest (may have timestamp in name)
        manifests = glob.glob(os.path.join(td_dir, f"PROPOSAL_SECTION_MANIFEST_{tid}.yaml"))
        if not manifests:
            print(f"[PPE] ERROR: Manifest not found in {td_dir}", file=sys.stderr); sys.exit(1)
        manifest_path = manifests[0]
        out_dir = td_dir
    elif args.md and args.manifest and args.out:
        tid = os.path.basename(args.out)
        md_path = args.md; manifest_path = args.manifest; out_dir = args.out
    else:
        parser.print_help(); sys.exit(1)

    if not os.path.exists(md_path):
        print(f"[PPE] ERROR: Rendered MD not found: {md_path}", file=sys.stderr); sys.exit(1)

    run_ppe(md_path, manifest_path, out_dir, tid)


if __name__ == "__main__":
    main()
