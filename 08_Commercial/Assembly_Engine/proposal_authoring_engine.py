#!/usr/bin/env python3
"""
proposal_authoring_engine.py — Proposal Authoring Engine v1.0 (PAE / PF2-009)

PAE is the final publishing step of the Proposal Factory.
It authors a client-facing proposal from scratch, treating all Knowledge Base
assets as source evidence — NOT as proposal sections.

PERMANENT PLATFORM RULE (GOV-PAE-001):
  Oracle Fusion Cloud proposals must use "Oracle's Modern Best Practices" and
  "Oracle Success Navigator". Oracle Unified Method (OUM) is prohibited for
  Fusion Cloud. OUM may only appear in Oracle E-Business Suite legacy contexts.

Two outputs:
  CLIENT_PROPOSAL_[Tender].docx
    — 25–40 pages, executive narrative, no internal content
  INTERNAL_REVIEW_PACK_[Tender].docx
    — Complete internal document: all V2 content, traceability, governance,
      source evidence, action register, approval records

Usage:
    python3 proposal_authoring_engine.py --tender PLENNEGY-HCM-001
"""

import argparse, glob, os, re, sys, yaml
from dataclasses import dataclass, field
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)
from docx_renderer import (
    Seg, parse_md, add_inline, render_table, render_callout,
    render_ulist, render_olist, setup_styles,
    _page_break, _hr, _toc_field, _tab_stop, _page_field, _numpages_field,
    _cell_bg, _cell_borders,
    NAVY, BLUE, GREY, LGREY, WHITE, ORANGE,
    HEX_NAVY, HEX_BLUE, HEX_WHITE, HEX_BG_PH, HEX_BG_AI,
    HEX_BD_PH, HEX_BD_AI, classify_comment, PROPOSALS_DIR
)
from proposal_shaper import (
    group_v2, classify_v2, v2_idx, filter_client,
    _ps, _ul, _tbl, _heading, _is_ai_sentinel,
    _page_setup, _headers_footers,
    HEX_STATUS_OK, HEX_STATUS_PH, HEX_STATUS_AI, LGREY2,
    HEX_BG_TRACE, HEX_BD_TRACE, HEX_BG_PENDING, HEX_BD_PENDING,
)
from pee import apply_budget

# ── Colours ────────────────────────────────────────────────────────────────────
HEX_INT_BANNER = "1A1A2E"
HEX_DIVIDER    = "E8E8E8"
DARK_BLUE      = RGBColor(0x1A, 0x1A, 0x2E)

# ── Authored content ──────────────────────────────────────────────────────────

def _authored_exec_summary() -> List[Seg]:
    return [
        Seg("comment", meta=(
            "PLACEHOLDER: Executive Summary — Authored by BU Lead + Account Manager. "
            "Include: Plennegy win theme, APPSolve positioning statement, "
            "2–3 key differentiators, call to action. Maximum 2 pages. "
            "Context: Plennegy is an agribusiness group seeking Oracle HCM Cloud. "
            "Key win themes: scale-up HR capability, single source of truth, "
            "integrated talent and recruiting."
        )),
    ]

def _authored_understanding() -> List[Seg]:
    return [
        Seg("comment", meta=(
            "PLACEHOLDER: Understanding of Plennegy's Requirements — Authored by Account Manager "
            "from RFP or initial qualification call. Demonstrate that APPSolve has understood "
            "Plennegy's specific context, pain points, and objectives — not a generic statement. "
            "Include: business context, HR technology objectives, key challenges, success criteria "
            "as understood by APPSolve. Maximum 2 pages."
        )),
    ]

def _authored_proposed_solution() -> List[Seg]:
    return [
        _ps("APPSolve proposes a phased Oracle Fusion HCM Cloud implementation for Plennegy Group, "
            "delivering core human capital management, talent acquisition, and integration services "
            "in Phase 1, followed by learning and talent management in Phase 2. The following "
            "scope is indicative, pending confirmation of final entity count, headcount, and "
            "go-live target from Plennegy."),
        _heading("Proposed Implementation Scope", 3),
        _tbl([
            "| Oracle Module | Business Capability Delivered | Phase |",
            "|---|---|---|",
            "| **Oracle HCM Core** | Single employee record — lifecycle, positions, absence, payroll data | 1 |",
            "| **Oracle Recruiting Cloud** | End-to-end recruitment — from vacancy to onboarding | 1 |",
            "| **Oracle Integration Cloud** | Payroll provider interface, third-party system connections | 1 + 2 |",
            "| **Oracle Learning Cloud** | Structured learning, certifications, compliance training | 2 |",
            "| **Oracle Talent Management** | Performance reviews, goals, succession planning | 2 |",
        ]),
        _heading("What Is Not in This Scope", 3),
        _ul([
            "Oracle Fusion Payroll (native) — APPSolve proposes OIC integration to Plennegy's existing payroll provider",
            "Oracle HR Help Desk — available as Phase 2 extension if required",
            "Oracle Analytics Cloud — separately licensed Oracle add-on, available post go-live",
            "Custom software development beyond Oracle's Low-Code/No-Code extension capability",
        ]),
        _ps("Final scope, commercial proposal, and payment model will be confirmed in a separate "
            "Commercial Proposal document issued by APPSolve's Commercial Director."),
    ]

def _authored_why_appsolve() -> List[Seg]:
    return [
        _ps("Oracle implementations succeed or fail on the quality of the implementation partner — "
            "not the technology. Over more than two decades, APPSolve has helped South African "
            "and international organisations realise the full value of Oracle technology by "
            "combining deep technical expertise with practical business implementation experience."),
        _ps("Oracle is at the core of what APPSolve does. Our team of more than 50 Senior "
            "Consultants brings Oracle HCM, ERP, Integration Cloud, Database Administration, "
            "and Cloud Infrastructure experience across hundreds of engagements in multiple "
            "sectors and geographies."),
        _heading("Our Oracle Recognition", 3),
        _ul([
            "**Oracle Level 1 Partner** — APPSolve holds Oracle Partner Network Level 1 status",
            "**Oracle Business Impact Award EMEA 2024** — recognising measurable business outcomes from our Oracle HCM implementations",
            "**Oracle ECEMEA Innovation Award 2024** — recognising innovation in Oracle Cloud delivery in our region",
        ]),
        _heading("Why Clients Choose APPSolve", 3),
        _ps("Four commitments distinguish APPSolve from other Oracle implementation partners in "
            "the South African market:"),
        _ul([
            "**We arrive prepared.** Our team designs the solution before the first workshop. "
            "We do not use your time to learn Oracle — we use it to understand your organisation.",
            "**We transfer knowledge.** Every engagement includes structured knowledge transfer "
            "so that your HR, IT, and project teams can operate and evolve the system independently by go-live.",
            "**We take personal accountability.** A named APPSolve Business Unit Lead is "
            "assigned to every engagement — a senior practitioner who takes personal responsibility "
            "for delivery quality and client satisfaction.",
            "**We stay.** APPSolve's Application Managed Services practice provides post-go-live "
            "support, Oracle quarterly update management, and continuous optimisation. "
            "The majority of our clients continue with APPSolve for years after go-live.",
        ]),
        _heading("Our Delivery Footprint", 3),
        _ps("APPSolve delivers Oracle implementations across South Africa (Gauteng and "
            "Western Cape), Sub-Saharan Africa (Botswana, Namibia, Mozambique, Zambia, Tanzania), "
            "and internationally (USA, France, Abu Dhabi). Our Oracle HCM portfolio spans "
            "gaming, retail, agribusiness, financial services, and professional services."),
        _ps("Reference project details are available on request, subject to individual client "
            "approval. Contact your APPSolve Account Manager."),
    ]

def _authored_how_we_deliver() -> List[Seg]:
    return [
        _ps("APPSolve's delivery approach for Oracle Fusion HCM Cloud is built on Oracle's "
            "Modern Best Practices and Oracle Success Navigator — the adoption framework "
            "Oracle has designed specifically for cloud SaaS implementations. We do not apply "
            "a legacy ERP methodology to Oracle Fusion Cloud."),
        _ps("Oracle Fusion HCM is a continuously evolving cloud platform, not a customised "
            "on-premises system. Our delivery philosophy is to configure Oracle's built-in "
            "best-practice processes to fit your organisation — rather than rebuilding your "
            "current HR processes in a new technology at higher risk and cost."),
        _heading("Our Delivery Philosophy", 3),
        _ul([
            "**Configure, not customise.** Oracle Fusion HCM is rich out of the box. "
            "Custom code adds cost, risk, and maintenance burden. We work within Oracle's "
            "standard configuration and Low-Code/No-Code extension framework.",
            "**Decisions enable delivery.** Oracle Fusion HCM requires your organisation to "
            "make configuration choices. APPSolve structures workshops to enable clear "
            "decisions at the right time — delays in decisions extend timelines.",
            "**Data quality is shared.** The quality of your Oracle HCM outcomes depends "
            "on the quality of the data you bring in. APPSolve provides templates, tools, "
            "and guidance — your team owns the data quality process.",
            "**Go-live is the beginning.** Every engagement includes structured hypercare and "
            "handover to our Application Managed Services practice. We plan for your "
            "long-term success from day one.",
        ]),
        _heading("The Six Delivery Phases", 3),
        _tbl([
            "| Phase | Name | Indicative Duration | What Happens for You |",
            "|---|---|---|---|",
            "| 1 | Mobilise | 4 weeks | Your project team is established, governance is in place, Oracle environments are provisioned, and the project is formally kicked off |",
            "| 2 | Scope & Design | 8–10 weeks | APPSolve runs Solution Design Workshops with your HR and IT teams to configure Oracle to your organisation — policies, structures, workflows |",
            "| 3 | Prototype | 6–8 weeks | You see your configured Oracle system for the first time in CRP-2 — a full prototype demonstration. Integration development begins. |",
            "| 4 | Validate | 6–8 weeks | Your team tests the configured system against real scenarios. Data migration is rehearsed. Go-live readiness is formally assessed. |",
            "| 5 | Deploy | 2–4 weeks | Go-live is executed with APPSolve on-site support. Hypercare begins immediately. |",
            "| 6 | Evolve | 12+ months | APPSolve manages Oracle quarterly cloud updates, optimises your configuration, and provides Application Managed Services |",
        ]),
        _heading("Quality and Change Management", 3),
        _ps("APPSolve assigns a dedicated Quality Assurance review at the end of each phase. "
            "Design decisions are formally documented and signed off before configuration begins. "
            "Change management — communicating the change, training your team, and supporting "
            "adoption — is built into our delivery model from Scope & Design onwards, not added "
            "as an afterthought before go-live."),
    ]

def _authored_hcm_capability() -> List[Seg]:
    return [
        _ps("Oracle Fusion HCM Cloud gives Plennegy Group a single, unified human capital "
            "management platform — from a candidate's first application to an employee's "
            "last day. The following describes what APPSolve will deliver for you "
            "and the outcomes you can expect."),
        _heading("Core Human Resource Management", 3),
        _ps("Your HR team will manage the complete employee lifecycle in one system. "
            "Organisation structures, position management, employment terms, personal data, "
            "and document management are all maintained in Oracle — with configurable "
            "approval workflows, full audit history, and manager and employee self-service "
            "available from day one."),
        _ps("Absence management is a frequent compliance and payroll risk for multi-entity "
            "organisations. Oracle HCM manages all absence types — annual leave, sick leave, "
            "parental leave, and statutory absences — with accurate accrual calculations, "
            "approval routing, and manager visibility. Absence data flows automatically "
            "to payroll processing, eliminating manual reconciliation."),
        _heading("Talent Acquisition", 3),
        _ps("Oracle Recruiting Cloud replaces disconnected, email-based recruitment with "
            "a structured, trackable hiring process. From vacancy creation to offer "
            "acceptance, every step is in one system. Candidates experience a professional, "
            "mobile-friendly application process. Hiring managers see their pipeline in "
            "real time. HR reports on time-to-hire, source effectiveness, and offer "
            "outcomes from go-live."),
        _ps("Approved candidates move directly from Oracle Recruiting into Oracle HCM Core — "
            "eliminating duplicate data entry, onboarding errors, and the delay between "
            "offer acceptance and system record creation. APPSolve has implemented Oracle "
            "Recruiting Cloud in large South African organisations and brings practical "
            "knowledge of the configuration, integrations, and adoption practices that "
            "make it work."),
        _heading("Learning and Development", 3),
        _ps("Oracle Learning Cloud gives Plennegy a structured learning environment without "
            "additional third-party tools. Course catalogues, learning paths, certifications, "
            "mandatory compliance training, and instructor-led training scheduling are "
            "all managed in Oracle. Employees access their learning through Oracle's "
            "self-service interface on desktop or mobile. Completion records and "
            "certification status are maintained against each employee's profile and "
            "visible to managers and HR in real time."),
        _heading("Talent Management", 3),
        _ps("Oracle's performance management, goal-setting, and succession planning modules "
            "give Plennegy's managers and HR teams a structured approach to employee "
            "development. Managers set goals with employees in Oracle, track progress "
            "through the year, and complete performance reviews in a single, consistent "
            "workflow. HR and leadership gain cross-functional talent visibility through "
            "talent profiles, succession dashboards, and calibration tools."),
        _heading("Payroll Integration and Integration Services", 3),
        _ps("Oracle Integration Cloud (OIC) connects Oracle Fusion HCM to Plennegy's "
            "payroll provider. APPSolve's proven bidirectional payroll integration moves "
            "employee and position data from Oracle HCM into payroll processing, and "
            "returns confirmed pay results and statutory data to Oracle. The integration "
            "includes exception handling, real-time monitoring, and alerting — so that "
            "payroll accuracy does not depend on manual reconciliation."),
        _ps("APPSolve has implemented Oracle Integration Cloud payroll interfaces in "
            "multiple South African production environments and brings proven integration "
            "patterns, testing frameworks, and operational runbooks to every engagement."),
        _heading("Workforce Reporting and Intelligence", 3),
        _ps("Oracle OTBI (Transactional Business Intelligence) and BI Publisher are "
            "included in every Oracle HCM implementation. APPSolve delivers a standard "
            "set of HR operational reports — headcount, absence analysis, recruitment "
            "pipeline, and learning completion — configured from go-live. Additional "
            "reports can be developed during the project or as part of post-go-live AMS."),
    ]

def _authored_governance() -> List[Seg]:
    return [
        _ps("APPSolve's project governance model provides clear accountability, rapid "
            "decision-making, and transparent risk management throughout the engagement. "
            "The framework operates at three levels: strategic, operational, and workstream."),
        _tbl([
            "| Governance Level | Frequency | Led By | Purpose |",
            "|---|---|---|---|",
            "| Steering Committee | Monthly | Executive Sponsors (client + APPSolve BU Lead) | Strategic decisions, scope authority, escalation |",
            "| Project Management | Weekly | Client PM + APPSolve PM | Delivery tracking, risk management, action ownership |",
            "| Workstream Reviews | Bi-weekly | Functional Leads | Module progress, design decisions, CRP planning |",
        ]),
        _ps("APPSolve issues a weekly project status report to all stakeholders covering "
            "RAG status by workstream, upcoming milestones, open risks, and pending client "
            "decisions. All scope changes are processed through a formal Change Request — "
            "assessed for commercial and timeline impact before any approval is given."),
        _ps("A project-specific RAID Register (Risks, Assumptions, Issues, Dependencies) "
            "is maintained throughout the engagement and reviewed at each Project "
            "Management meeting. Specific RAID content, Change Control procedures, "
            "and the Cutover Plan will be developed during the Mobilise phase with "
            "Plennegy's project sponsor and APPSolve's Project Manager."),
    ]

def _authored_roadmap() -> List[Seg]:
    return [
        _ps("The following roadmap provides indicative timelines for Plennegy's Oracle "
            "Fusion HCM Cloud implementation. Durations are subject to scope confirmation, "
            "Oracle environment provisioning timelines, and Plennegy resource availability."),
        _tbl([
            "| Phase | Indicative Duration | Key Client Milestones |",
            "|---|---|---|",
            "| Mobilise | 4 weeks | Governance sign-off; project team confirmed; kick-off |",
            "| Scope & Design | 8–10 weeks | Design documents approved; CRP-1 complete |",
            "| Prototype | 6–8 weeks | CRP-2 prototype sign-off; integration approved |",
            "| Validate | 6–8 weeks | UAT sign-off; data migration approved; go-live confirmed |",
            "| Deploy | 2–4 weeks | Go-live; hypercare active |",
            "| Evolve | 12+ months | AMS active; quarterly updates managed |",
        ]),
        _ul([
            "**Kick-off** target: within 2 weeks of contract signature",
            "**Go-live** target: to be confirmed with Plennegy (scope and entity count required first)",
            "Detailed project plan with named milestones and dates is developed during Mobilise",
            "Oracle environment provisioning is managed between Oracle Corporation and Plennegy — not within APPSolve's control",
        ]),
    ]

def _authored_key_assumptions() -> List[Seg]:
    return [
        _ps("APPSolve's proposal is governed by a comprehensive set of commercial assumptions "
            "that define the contractual basis for scope, cost, and timeline. The following "
            "categories apply to this engagement. Full detail is in Appendix A."),
        _tbl([
            "| Assumption Category | Count | Most Material to Scope and Cost |",
            "|---|---|---|",
            "| Oracle HCM Core (Core HR, Absence, Payroll Data) | 115 | Entity count, headcount, payroll system, data migration scope |",
            "| Oracle Recruiting Cloud | 54 | Vacancy volume, pre-employment check requirements, onboarding scope |",
            "| Oracle Learning Cloud | 37 | Course catalogue size, delivery methods, certification requirements |",
            "| Oracle Talent Management | 31 | Review cycle design, succession scope, performance rating model |",
            "| Oracle Integration Cloud | 104 | Number of integrations, payroll provider, third-party systems |",
        ]),
        _ul([
            "Client is responsible for data extraction, cleansing, and validation. APPSolve provides templates and guidance.",
            "Oracle Fusion HCM SaaS environments are provisioned and managed by Oracle Corporation — provisioning timelines are outside APPSolve's control.",
            "Scope and pricing are based on the entity count, headcount, and module list confirmed during Mobilise. Scope changes require a formal Change Request.",
            "Go-live readiness is a joint decision. APPSolve may withhold go-live recommendation if agreed readiness criteria are not met.",
        ]),
        _ps("The complete Assumption Schedule is in **Appendix A**. All assumptions govern "
            "this proposal and the subsequent Statement of Work."),
    ]

def _authored_commercial() -> List[Seg]:
    return [
        _ps("APPSolve's formal commercial proposal will be issued by our Commercial Director "
            "once the scope inputs below are confirmed. The commercial document constitutes "
            "the binding offer and must be counter-signed by both parties before work commences."),
        _heading("Information Required Before Commercial Proposal", 3),
        _ul([
            "**Confirmed scope:** modules, legal entities, employee headcount, payroll provider, integration points",
            "**Preferred commercial model:** Fixed-price, Time and Materials, or Monthly Recurring Invoice",
            "**Target go-live date:** required to resource-plan the engagement",
            "**Budget range:** enables APPSolve to phase or prioritise the solution where beneficial",
        ]),
        _heading("APPSolve's Commercial Models", 3),
        _tbl([
            "| Model | Description | Best Suited To |",
            "|---|---|---|",
            "| Fixed Price | Agreed scope at a fixed cost; Change Requests separately priced | Defined HCM Cloud scope with low likelihood of scope change |",
            "| Time and Materials | Billing at agreed rates per day; no fixed cost ceiling | Scope expected to evolve during delivery |",
            "| Monthly Recurring Invoice | Fixed monthly cost; capacity transfers between months | Predictable cost management; APPSolve's preferred model for sustained engagements |",
        ]),
        _ps("_Pricing is provided exclusively in APPSolve's formal Commercial Proposal "
            "document. No rates, estimates, or pricing terms in this document are binding "
            "until the Commercial Proposal is issued and counter-signed._"),
    ]

def _authored_next_steps() -> List[Seg]:
    return [
        _ps("APPSolve is ready to progress this engagement. The following actions are "
            "Plennegy's to complete before APPSolve can issue the final commercial proposal:"),
        _tbl([
            "| Priority | Your Action | Timeline |",
            "|---|---|---|",
            "| 1 | **Confirm Oracle HCM scope** — which modules, how many legal entities, employee headcount, payroll provider, go-live target | Before commercial proposal |",
            "| 2 | **Nominate your Project Sponsor** — a Plennegy executive with authority to approve scope, design decisions, and commercial terms | Before commercial proposal |",
            "| 3 | **Confirm your preferred commercial model** — Fixed Price, Time and Materials, or Monthly Recurring Invoice | Before commercial proposal |",
            "| 4 | **Approve reference client usage** — advise your APPSolve Account Manager which reference clients may be named in this submission | Before submission |",
            "| 5 | **Confirm any RFP or scoring requirements** — if this submission must address specific evaluation criteria, advise your Account Manager | Before submission |",
            "| 6 | **Schedule a scope confirmation call** — APPSolve's Account Manager and Business Unit Lead will attend | This week |",
        ]),
        _ps("Once scope is confirmed, APPSolve will produce the final commercial proposal "
            "within 5–10 business days. The final proposal, APPSolve's formal Statement "
            "of Work, and relevant compliance certificates will be submitted together "
            "as a complete submission package."),
    ]

# ── Architecture standard docs ────────────────────────────────────────────────
_AUTHORED_SECTIONS_MAP = {
    "exec_summary":       _authored_exec_summary,
    "understanding":      _authored_understanding,
    "proposed_solution":  _authored_proposed_solution,
    "why_appsolve":       _authored_why_appsolve,
    "how_we_deliver":     _authored_how_we_deliver,
    "hcm_capability":     _authored_hcm_capability,
    "governance":         _authored_governance,
    "roadmap":            _authored_roadmap,
    "key_assumptions":    _authored_key_assumptions,
    "commercial":         _authored_commercial,
    "next_steps":         _authored_next_steps,
}

# ── PAE section definition ────────────────────────────────────────────────────
PAE_SECTIONS = [
    {"id": "1",  "heading": "1. Executive Summary",             "fn": "exec_summary",     "app": False},
    {"id": "2",  "heading": "2. Understanding Plennegy's Requirements", "fn": "understanding", "app": False},
    {"id": "3",  "heading": "3. The Solution We Propose",       "fn": "proposed_solution","app": False},
    {"id": "4",  "heading": "4. Why APPSolve",                  "fn": "why_appsolve",     "app": False},
    {"id": "5",  "heading": "5. How We Deliver",                "fn": "how_we_deliver",   "app": False},
    {"id": "6",  "heading": "6. What We Deliver for You — Oracle Fusion HCM", "fn": "hcm_capability","app": False},
    {"id": "7",  "heading": "7. Project Governance",            "fn": "governance",       "app": False},
    {"id": "8",  "heading": "8. Your Implementation Roadmap",   "fn": "roadmap",          "app": False},
    {"id": "9",  "heading": "9. Key Assumptions",               "fn": "key_assumptions",  "app": False},
    {"id": "10", "heading": "10. Commercial",                   "fn": "commercial",       "app": False},
    {"id": "11", "heading": "11. Your Next Steps",              "fn": "next_steps",       "app": False},
    {"id": "A",  "heading": "Appendix A — Assumption Schedule", "fn": None, "v2_name": "Complete Assumption Schedule",
     "app": True, "budget": 120},
    {"id": "B",  "heading": "Appendix B — Reference Summary",   "fn": None,
     "v2_name": "Client References", "app": True, "budget": 24},
]

# ── Shared DOCX helpers ───────────────────────────────────────────────────────

def _count_content(segs: List[Seg]) -> int:
    return sum(1 for s in segs if s.kind in ("paragraph","table","ulist","olist") and not _is_ai_sentinel(s))

def _render_seg(doc, seg: Seg):
    if seg.kind == "heading":
        h = doc.add_heading(level=min(seg.level, 5)); h.paragraph_format.keep_with_next=True
        add_inline(h, seg.lines[0] if seg.lines else "", sz={1:22,2:16,3:13,4:11,5:11}.get(seg.level,11))
    elif seg.kind == "table":   render_table(doc, seg.lines)
    elif seg.kind == "ulist":   render_ulist(doc, seg.lines)
    elif seg.kind == "olist":   render_olist(doc, seg.lines)
    elif seg.kind == "blockquote": render_callout(doc, seg.lines)
    elif seg.kind == "hr":      _hr(doc)
    elif seg.kind == "paragraph":
        text = " ".join(seg.lines).strip()
        if text:
            p = doc.add_paragraph(style="Normal"); add_inline(p, text)
    elif seg.kind == "comment":
        kind = classify_comment(seg.meta)
        if kind in ("placeholder","ai_draft"):
            render_callout(doc, [f"**{'Action Required' if kind=='placeholder' else 'AI-Draft Required'}:** {seg.meta}"])
        elif kind == "governance_note":
            p = doc.add_paragraph()
            p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(3)
            r = p.add_run(f"⚙  {seg.meta}"); r.font.name="Calibri"; r.font.size=Pt(8.5); r.font.italic=True; r.font.color.rgb=LGREY2

def _pending_client(doc, label: str, note: str = ""):
    tbl = doc.add_table(rows=1, cols=1); tbl.style="Table Grid"
    c = tbl.cell(0,0); _cell_bg(c, "F7F7F7"); _cell_borders(c,"AAAAAA",6)
    p = c.paragraphs[0]; p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
    r = p.add_run(f"[ {label} — content to be completed before submission ]")
    r.font.name="Calibri"; r.font.size=Pt(10); r.font.italic=True; r.font.color.rgb=LGREY2
    if note:
        p2 = c.add_paragraph()
        r2 = p2.add_run(note); r2.font.name="Calibri"; r2.font.size=Pt(9); r2.font.color.rgb=LGREY2
    doc.add_paragraph().paragraph_format.space_after=Pt(6)

def _internal_action_callout(doc, label: str, action: str, owner: str = ""):
    tbl = doc.add_table(rows=1, cols=1); tbl.style="Table Grid"
    c = tbl.cell(0,0); _cell_bg(c,"F8D7DA"); _cell_borders(c,"C00000",8)
    p = c.paragraphs[0]; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(2)
    r = p.add_run(f"⚠  {label}"); r.font.name="Calibri"; r.font.size=Pt(10); r.font.bold=True
    if action:
        p2 = c.add_paragraph()
        r2 = p2.add_run(action); r2.font.name="Calibri"; r2.font.size=Pt(9.5)
    if owner:
        p3 = c.add_paragraph()
        r3 = p3.add_run(f"Owner: {owner}"); r3.font.name="Calibri"; r3.font.size=Pt(9); r3.font.italic=True; r3.font.color.rgb=LGREY2
    doc.add_paragraph().paragraph_format.space_after=Pt(6)

# ── CLIENT PROPOSAL renderer ──────────────────────────────────────────────────

class ClientProposalRenderer:
    def __init__(self, v2_secs, out_path: str, tender_id: str):
        self.v2_secs = v2_secs
        self.v2_idx  = {s.name.lower(): s for s in v2_secs}
        self.out_path= out_path; self.tid = tender_id
        self.doc = Document(); setup_styles(self.doc); _page_setup(self.doc)
        self.stats = {"authored": 0, "pending": 0, "appendix": 0}

    def render(self):
        d = self.doc
        self._cover()
        h = d.add_heading("Table of Contents", level=1); h.paragraph_format.space_before=Pt(0)
        _toc_field(d); _page_break(d)
        # Body sections (1–11)
        for defn in PAE_SECTIONS:
            if defn.get("app"): continue
            _page_break(d)
            h = d.add_heading(defn["heading"], level=2); h.paragraph_format.keep_with_next=True
            fn = _AUTHORED_SECTIONS_MAP.get(defn["fn"])
            segs = fn() if fn else []
            has_ph = any(s.kind == "comment" and classify_comment(s.meta) in ("placeholder","ai_draft") for s in segs)
            if has_ph:
                lbl = re.sub(r"^\d+\.\s+", "", defn["heading"])
                _pending_client(d, lbl)
                self.stats["pending"] += 1
            else:
                for seg in segs: _render_seg(d, seg)
                self.stats["authored"] += 1
        # Appendices
        for defn in PAE_SECTIONS:
            if not defn.get("app"): continue
            _page_break(d)
            d.add_heading(defn["heading"], level=2).paragraph_format.keep_with_next=True
            v2_name = defn.get("v2_name","")
            v2_sec  = self.v2_idx.get(v2_name.lower()) if v2_name else None
            if v2_sec and v2_sec.status == "RENDERED":
                segs = filter_client(v2_sec.segs)
                kept, _, _ = apply_budget(segs, defn.get("budget", 120), "")
                for seg in kept: _render_seg(d, seg)
            else:
                _pending_client(d, re.sub(r"^Appendix\s+\w+\s+[—-]\s+","",defn["heading"]))
            self.stats["appendix"] += 1
        _headers_footers(d, "APPSolve (Pty) Ltd",
                         "Oracle HCM Cloud — Implementation Proposal  |  COMMERCIAL IN CONFIDENCE",
                         self.tid)
        os.makedirs(os.path.dirname(self.out_path), exist_ok=True)
        d.save(self.out_path)
        sz = os.path.getsize(self.out_path)
        print(f"[CLIENT PAE]  {self.out_path.split('/')[-1]}  ({sz//1024} KB)  "
              f"Authored={self.stats['authored']} Pending={self.stats['pending']} App={self.stats['appendix']}")

    def _cover(self):
        d = self.doc
        d.add_paragraph().paragraph_format.space_before=Pt(60)
        for txt, sz, bold, colour in [
            ("APPSolve (Pty) Ltd",                           32, True,  NAVY),
            ("Enterprise Technology Services  |  Oracle  |  Acumatica  |  BeBanking", 11, False, BLUE),
        ]:
            p = d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            if "Enterprise" in txt: p.paragraph_format.space_after=Pt(48)
            r = p.add_run(txt); r.font.name="Calibri"; r.font.size=Pt(sz); r.font.bold=bold; r.font.color.rgb=colour
        for txt, sz, bold, colour in [
            ("ORACLE FUSION HCM CLOUD",                      18, True,  GREY),
            ("IMPLEMENTATION PROPOSAL",                       14, True,  GREY),
        ]:
            p = d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(4)
            r = p.add_run(txt); r.font.name="Calibri"; r.font.size=Pt(sz); r.font.bold=bold; r.font.color.rgb=colour
        p_cl = d.add_paragraph(); p_cl.alignment=WD_ALIGN_PARAGRAPH.CENTER; p_cl.paragraph_format.space_after=Pt(40)
        r_cl = p_cl.add_run("Plennegy Group"); r_cl.font.name="Calibri"; r_cl.font.size=Pt(16); r_cl.font.bold=True; r_cl.font.color.rgb=NAVY
        tbl = d.add_table(rows=0, cols=2); tbl.style="Table Grid"
        for k, v in [("Prepared For","Plennegy Group"),("Platform","Oracle Fusion HCM Cloud"),
                     ("Engagement","Full Suite HCM Implementation"),("Prepared By","APPSolve (Pty) Ltd"),
                     ("Document","Client Proposal — PAE v1.0"),
                     ("Date", datetime.now().strftime("%Y-%m-%d")),("Tender ID", self.tid)]:
            row=tbl.add_row(); kc=row.cells[0]; vc=row.cells[1]
            _cell_bg(kc,HEX_NAVY); _cell_borders(kc,HEX_NAVY,4)
            _cell_bg(vc,"F2F5FB"); _cell_borders(vc,"D0D0D0",4)
            kp=kc.paragraphs[0]; kp.paragraph_format.space_before=Pt(4); kp.paragraph_format.space_after=Pt(4)
            kr=kp.add_run(k); kr.font.name="Calibri"; kr.font.size=Pt(11); kr.font.bold=True; kr.font.color.rgb=WHITE
            vp=vc.paragraphs[0]; vp.paragraph_format.space_before=Pt(4); vp.paragraph_format.space_after=Pt(4)
            add_inline(vp, v, sz=11.0)
        fp=d.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; fp.paragraph_format.space_before=Pt(48)
        fr=fp.add_run("COMMERCIAL IN CONFIDENCE  |  Proposal Authoring Engine v1.0")
        fr.font.name="Calibri"; fr.font.size=Pt(10); fr.font.bold=True; fr.font.color.rgb=LGREY2
        _page_break(d)

# ── INTERNAL REVIEW PACK renderer ────────────────────────────────────────────

class InternalReviewPackRenderer:
    def __init__(self, v2_secs, out_path: str, tender_id: str):
        self.v2_secs  = v2_secs
        self.out_path = out_path; self.tid = tender_id
        self.doc = Document(); setup_styles(self.doc); _page_setup(self.doc)

    def render(self):
        d = self.doc
        # Banner
        p1=d.add_paragraph(); p1.alignment=WD_ALIGN_PARAGRAPH.CENTER; p1.paragraph_format.space_before=Pt(24)
        r1=p1.add_run("APPSolve — Internal Review Pack")
        r1.font.name="Calibri"; r1.font.size=Pt(22); r1.font.bold=True; r1.font.color.rgb=NAVY
        p2=d.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after=Pt(8)
        r2=p2.add_run(f"{self.tid}  |  {datetime.now().strftime('%Y-%m-%d')}  |  PAE v1.0  |  Platform L6.2")
        r2.font.name="Calibri"; r2.font.size=Pt(11); r2.font.color.rgb=BLUE
        render_callout(d, [
            "**INTERNAL USE ONLY — APPSolve Confidential.** "
            "This pack contains the complete internal knowledge base content, source evidence, "
            "governance records, and traceability. Do NOT distribute to clients. "
            "All content removed from the Client Proposal is in this document.",
        ])
        _page_break(d)
        # Action register
        self._action_register()
        _page_break(d)
        # Published sections (show what went to client)
        h = d.add_heading("Part 1 — Sections Published to Client", level=1); h.paragraph_format.space_before=Pt(0)
        for defn in PAE_SECTIONS:
            if defn.get("app"): continue
            _page_break(d)
            hdr = d.add_heading(defn["heading"], level=2); hdr.paragraph_format.keep_with_next=True
            fn = _AUTHORED_SECTIONS_MAP.get(defn.get("fn",""))
            segs = fn() if fn else []
            has_ph = any(s.kind=="comment" and classify_comment(s.meta) in ("placeholder","ai_draft") for s in segs)
            if has_ph:
                _internal_action_callout(d, f"PENDING — {defn['heading']}", "Section not authored yet — see Action Register.", "AM / BU Lead")
            else:
                render_callout(d, ["**PAE Authored Content** — Published to Client. Review for accuracy before submission."])
                for seg in segs: _render_seg(d, seg)
        _page_break(d)
        # KB assets not in client
        h2 = d.add_heading("Part 2 — Knowledge Base Content (Not Published to Client)", level=1); h2.paragraph_format.space_before=Pt(0)
        d.add_paragraph(style="Normal").runs and None
        p_note = d.add_paragraph(style="Normal")
        add_inline(p_note, "The following Knowledge Base sections were removed from the Client Proposal. "
                   "They are retained here for traceability, governance, and future editorial decisions.")
        for sec in self.v2_secs:
            if not sec.segs: continue
            _page_break(d)
            h_s = d.add_heading(sec.name, level=2); h_s.paragraph_format.keep_with_next=True
            badge = {"RENDERED": "✓ RENDERED", "PLACEHOLDER": "⚠ PLACEHOLDER",
                     "AI_DRAFT": "✏ AI-DRAFT", "EMPTY": "— EMPTY"}.get(sec.status, sec.status)
            p_b = d.add_paragraph(style="Normal")
            add_inline(p_b, f"Status: {badge}  |  V2 section — not published to Client Proposal  |  Source content only", sz=9.0)
            for seg in sec.segs:
                _render_seg(d, seg)
        _headers_footers(d, "APPSolve INTERNAL REVIEW PACK",
                         f"{self.tid}  |  PAE v1.0  |  INTERNAL ONLY", self.tid)
        os.makedirs(os.path.dirname(self.out_path), exist_ok=True)
        d.save(self.out_path)
        sz = os.path.getsize(self.out_path)
        print(f"[INTERNAL PAE]{self.out_path.split('/')[-1]}  ({sz//1024} KB)  V2 sections={len(self.v2_secs)}")

    def _action_register(self):
        d = self.doc
        h = d.add_heading("Action Register — Pre-Submission Requirements", level=1); h.paragraph_format.space_before=Pt(0)
        tbl = d.add_table(rows=1, cols=4); tbl.style="Table Grid"
        for ci, hdr in enumerate(["Priority", "Action", "Owner", "Status"]):
            hc=tbl.rows[0].cells[ci]; _cell_bg(hc,HEX_NAVY); _cell_borders(hc,HEX_NAVY,4)
            hp=hc.paragraphs[0]; hp.paragraph_format.space_before=Pt(3); hp.paragraph_format.space_after=Pt(3)
            hr2=hp.add_run(hdr); hr2.font.name="Calibri"; hr2.font.size=Pt(10); hr2.font.bold=True; hr2.font.color.rgb=WHITE
        actions = [
            ("P0 URGENT", "Renew B-BBEE Level 3 certificate", "Finance Director", "Due 2026-07-31"),
            ("P1", "Author Executive Summary (Sec 1) with win themes", "BU Lead", "Pending"),
            ("P1", "Author Understanding of Requirements (Sec 2) from RFP", "Account Manager", "Pending"),
            ("P1", "Confirm scope: modules, entities, headcount, payroll provider, go-live date", "AM (OAR-C04)", "Pending"),
            ("P1", "Issue Commercial Proposal once scope confirmed", "Commercial Director (OAR-C02)", "Pending"),
            ("P2", "Approve Hollywood Bets reference for this tender", "Account Manager (OAR-C01)", "Pending"),
            ("P2", "BU Lead review: Why APPSolve, How We Deliver, Capability, Governance sections", "BU Lead", "Pending"),
            ("P2", "Provide proposed team structure and consultant CVs", "Delivery Manager", "Pending"),
            ("P3", "Verify compliance document expiry dates (B-BBEE, OPN cert, Directors Resolution)", "Operations", "Pending"),
            ("P3", "Confirm Oracle award claims are current for this tender", "BU Lead", "Pending"),
        ]
        for prio, action, owner, status in actions:
            bg = "F8D7DA" if "P0" in prio or "P1" in prio else "FFF3CD" if "P2" in prio else "F2F5FB"
            row = tbl.add_row()
            for c in row.cells: _cell_bg(c, bg); _cell_borders(c,"D0D0D0",4)
            for ci2, v in enumerate([prio, action, owner, status]):
                cp=row.cells[ci2].paragraphs[0]; cp.paragraph_format.space_before=Pt(3); cp.paragraph_format.space_after=Pt(3)
                add_inline(cp, v, sz=9.5)

# ── Orchestrator ──────────────────────────────────────────────────────────────

def run_pae(md_path: str, manifest_path: str, out_dir: str, tender_id: str) -> dict:
    print(f"[PAE] Tender:   {tender_id}")
    print(f"[PAE] Input MD: {md_path}")
    with open(md_path, "r", encoding="utf-8") as fh:
        md_text = fh.read()
    all_segs = parse_md(md_text)
    v2_secs  = group_v2(all_segs)
    classify_v2(v2_secs)
    print(f"[PAE] V2: {len(v2_secs)} secs | "
          f"RENDERED={sum(1 for s in v2_secs if s.status=='RENDERED')} "
          f"PH={sum(1 for s in v2_secs if s.status=='PLACEHOLDER')} "
          f"AI={sum(1 for s in v2_secs if s.status=='AI_DRAFT')}")

    os.makedirs(out_dir, exist_ok=True)
    prefix = "PLENNEGY" if "PLENNEGY" in tender_id.upper() else tender_id

    client_path   = os.path.join(out_dir, f"CLIENT_PROPOSAL_{prefix}.docx")
    internal_path = os.path.join(out_dir, f"INTERNAL_REVIEW_PACK_{prefix}.docx")

    ClientProposalRenderer(v2_secs, client_path, tender_id).render()
    InternalReviewPackRenderer(v2_secs, internal_path, tender_id).render()

    print(f"\n[PAE] Complete — 2 outputs in {out_dir}")
    return {"client": client_path, "internal": internal_path}

# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Proposal Authoring Engine v1.0 (PAE)")
    parser.add_argument("--tender", metavar="TENDER_ID")
    parser.add_argument("--md",       metavar="MD_PATH")
    parser.add_argument("--manifest", metavar="MANIFEST_PATH")
    parser.add_argument("--out",      metavar="OUT_DIR")
    args = parser.parse_args()
    if args.tender:
        tid = args.tender
        td_dir = os.path.join(PROPOSALS_DIR, tid)
        md_path = os.path.join(td_dir, f"PROPOSAL_RENDERED_{tid}.md")
        mfiles  = glob.glob(os.path.join(td_dir, f"PROPOSAL_SECTION_MANIFEST_{tid}.yaml"))
        if not mfiles:
            print(f"[PAE] ERROR: Manifest not found in {td_dir}", file=sys.stderr); sys.exit(1)
        manifest_path = mfiles[0]; out_dir = td_dir
    elif args.md and args.manifest and args.out:
        tid = os.path.basename(args.out); md_path = args.md
        manifest_path = args.manifest; out_dir = args.out
    else:
        parser.print_help(); sys.exit(1)
    if not os.path.exists(md_path):
        print(f"[PAE] ERROR: Rendered MD not found: {md_path}", file=sys.stderr); sys.exit(1)
    run_pae(md_path, manifest_path, out_dir, tid)

if __name__ == "__main__":
    main()
